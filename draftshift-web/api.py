"""
DraftShift Web API — FastAPI backend for pleading generation.

Runs on Replit at draftshift.replit.dev
"""

from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pathlib import Path
import json
import io
import logging
import base64

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# DraftShift is optional - API works without it (generates mock DOCX)
generate_pleading = None
SUPPORTED_TYPES = ["motion", "opposition", "reply", "declaration"]

logger.info("✅ DraftShift Web API initialized (light mode - no heavy dependencies)")

# Initialize FastAPI
app = FastAPI(
    title="DraftShift API",
    description="Generate California civil pleadings from JSON input",
    version="0.1.0"
)

# Enable CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================================================
# API ENDPOINTS
# ============================================================

@app.get("/api/health")
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "draftshift_ready": generate_pleading is not None,
        "supported_types": SUPPORTED_TYPES
    }


@app.post("/api/build")
async def build_pleading(data: dict):
    """
    Build a pleading from JSON data.
    
    Request body:
    {
        "type": "motion|opposition|reply|declaration",
        "attorney": {...},
        "case": {...},
        "title": "...",
        "arguments": [...],
        "position": {...}
    }
    
    Returns:
    {
        "success": true,
        "filename": "Motion.docx",
        "data": "base64_encoded_docx"
    }
    """
    try:
        # Validate input
        if "type" not in data:
            raise ValueError("Missing required field: 'type'")
        
        pleading_type = data["type"].lower()
        if pleading_type not in SUPPORTED_TYPES:
            raise ValueError(f"Invalid type: {pleading_type}. Supported: {SUPPORTED_TYPES}")
        
        logger.info(f"Building {pleading_type} pleading")
        
        # If generate_pleading is available, use it
        if generate_pleading:
            result = generate_pleading(data)
            docx_bytes = result
        else:
            # Fallback: create a simple DOCX with the JSON data
            from docx import Document
            doc = Document()
            doc.add_heading(data.get("title", pleading_type.capitalize()), 0)
            doc.add_paragraph(json.dumps(data, indent=2))
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            docx_bytes = output.getvalue()
        
        # Return DOCX as base64
        filename = f"{pleading_type.capitalize()}.docx"
        docx_base64 = base64.b64encode(docx_bytes).decode()
        
        logger.info(f"Successfully built {filename}")
        
        return {
            "success": True,
            "filename": filename,
            "data": docx_base64
        }
        
    except ValueError as e:
        logger.error(f"Validation error: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Build error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Build failed: {str(e)}")


@app.get("/api/fixtures/{fixture_name}")
async def get_fixture(fixture_name: str):
    """
    Get a fixture JSON file for testing.
    
    Available: motion.json, opposition.json, reply.json, declaration.json
    """
    # Try multiple paths
    possible_paths = [
        Path("DraftShift/tests/fixtures") / fixture_name,
        Path("draftshift/tests/fixtures") / fixture_name,
        Path("./fixtures") / fixture_name,
    ]
    
    fixture_path = None
    for path in possible_paths:
        if path.exists():
            fixture_path = path
            break
    
    if not fixture_path:
        logger.warning(f"Fixture not found: {fixture_name}")
        # Return mock data instead
        return {
            "type": fixture_name.replace(".json", ""),
            "attorney": {"name": "John Doe", "firm": "Doe & Associates"},
            "case": {"number": "2024-001", "court": "California Superior Court"},
            "title": f"{fixture_name.replace('.json', '').capitalize()} Title",
            "arguments": ["First argument", "Second argument"],
            "position": {}
        }
    
    try:
        with open(fixture_path, "r") as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"Error loading fixture: {e}")
        raise HTTPException(status_code=500, detail="Failed to load fixture")


# ============================================================
# STATIC FILES
# ============================================================

# Serve React frontend from dist folder
from pathlib import Path
dist_path = Path(__file__).parent / "dist"
if dist_path.exists():
    try:
        app.mount("/", StaticFiles(directory=str(dist_path), html=True), name="static")
        logger.info(f"✅ Static files mounted from {dist_path}")
    except Exception as e:
        logger.warning(f"⚠️ Could not mount static files: {e}")
else:
    logger.warning(f"⚠️ Dist directory not found at {dist_path}. Run 'npm run build' first.")


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":
    import uvicorn
    
    # Run on 0.0.0.0 for Replit
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8000,
        log_level="info"
    )
