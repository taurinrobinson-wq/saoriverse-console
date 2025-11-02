from io import BytesIO

from docx import Document


def generate_doc(date, time, event, mood, reflections, insights):
    doc = Document()
    doc.add_heading("Personal Log Entry", level=1)
    doc.add_paragraph(f"Date: {date}")
    doc.add_paragraph(f"Time: {time}")
    doc.add_paragraph(f"Event: {event}")
    doc.add_paragraph(f"Mood: {mood}")
    doc.add_paragraph("Reflections:")
    doc.add_paragraph(reflections)
    doc.add_paragraph("Insights:")
    doc.add_paragraph(insights)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
