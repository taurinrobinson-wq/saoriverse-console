from docx import Document
from io import BytesIO
import json
import requests
import time
import streamlit as st
import datetime

# --- Generate Word Document from Personal Log ---


def generate_doc(date, time, event, mood, reflections, insights):
    doc = Document()
    doc.add_heading("Personal Log Entry", level=1)

    doc.add_paragraph(f"Date: {date}")
    doc.add_paragraph(f"Time: {time}")
    doc.add_paragraph(f"Event: {event}")
    doc.add_paragraph(f"Mood: {mood}")
    doc.add_paragraph("Reflections:")
    doc.add_paragraph(reflections)
    doc.add_paragraph("Insights:")
    doc.add_paragraph(insights)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
# --- End Generate Word Document ---


# Page configuration
st.set_page_config(
    page_title="FirstPerson - Personal AI Companion",
    page_icon="/static/graphics/FirstPerson-Logo-normalized.svg",
    layout="wide",
    initial_sidebar_state="collapsed"
)


class SaoynxAuthentication:
    """Clean authentication system matching SAOYNX design"""

    def __init__(self):
        # Use Streamlit secrets
        try:
            self.supabase_url = st.secrets["supabase"]["url"]
            self.supabase_key = st.secrets["supabase"]["key"]
        except KeyError:
            st.error("❌ Supabase configuration not found in secrets")
            st.stop()

        self.init_session_state()

    def create_session_token(self, username: str, user_id: str) -> str:
        """Create a secure session token"""
        import base64
        session_data = {
            "username": username,
            "user_id": user_id,
            "created": datetime.datetime.now().isoformat(),
            "expires": (datetime.datetime.now() + datetime.timedelta(days=2)).isoformat()
        }
        # Simple base64 encoding (in production, use proper JWT)
        token = base64.b64encode(json.dumps(session_data).encode()).decode()
        return token

    def validate_session_token(self, token: str) -> dict:
        """Validate and decode session token"""
        try:
            import base64
            # Decode the token
            decoded_data = base64.b64decode(token.encode()).decode()
            session_data = json.loads(decoded_data)

            # Check required fields
            required_fields = ["username", "user_id", "expires"]
            for field in required_fields:
                if field not in session_data:
                    return {"valid": False, "error": f"Missing field: {field}"}

            # Check expiration
            try:
                expires = datetime.datetime.fromisoformat(
                    session_data["expires"])
                if datetime.datetime.now() < expires:
                    return {"valid": True, "data": session_data}
                else:
                    return {"valid": False, "error": "Session expired"}
            except ValueError as e:
                return {"valid": False, "error": f"Invalid date format: {e}"}

        except Exception as e:
            return {"valid": False, "error": f"Token validation error: {str(e)}"}

    def init_session_state(self):
        """Initialize authentication session state with persistence"""
        # Initialize defaults first
        if 'authenticated' not in st.session_state:
            st.session_state.authenticated = False
        if 'user_id' not in st.session_state:
            st.session_state.user_id = None
        if 'username' not in st.session_state:
            st.session_state.username = None
        if 'session_expires' not in st.session_state:
            st.session_state.session_expires = None

        # Only try to restore session if not already authenticated
        if not st.session_state.authenticated:
            # Check for existing session token in query params
            query_params = st.query_params
            session_token = query_params.get("session_token")

            if session_token:
                # Try to restore session from token
                session_result = self.validate_session_token(session_token)
                if session_result["valid"]:
                    data = session_result["data"]
                    st.session_state.authenticated = True
                    st.session_state.user_id = data["user_id"]
                    st.session_state.username = data["username"]
                    st.session_state.session_expires = data["expires"]
                else:
                    # Invalid token, remove it
                    if "session_token" in st.query_params:
                        del st.query_params["session_token"]

    def authenticate_user(self, username: str, password: str) -> dict:
        """Authenticate user login"""
        try:
            auth_url = st.secrets.get("supabase", {}).get(
                "auth_function_url", f"{self.supabase_url}/functions/v1/auth-manager")
            response = requests.post(
                auth_url,
                headers={
                    "Authorization": f"Bearer {self.supabase_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "action": "authenticate",
                    "username": username,
                    "password": password
                },
                timeout=10
            )

            if response.status_code == 200:
                data = response.json()
                if data.get("authenticated"):
                    # Set session state
                    st.session_state.authenticated = True
                    st.session_state.user_id = data.get("user_id")
                    st.session_state.username = username
                    st.session_state.session_expires = (
                        datetime.datetime.now() + datetime.timedelta(days=2)).isoformat()

                    # Create persistent session token
                    session_token = self.create_session_token(
                        username, data.get("user_id"))

                    # Add session token to URL for persistence
                    st.query_params["session_token"] = session_token

                    return {"success": True, "message": "Login successful"}
                else:
                    error_msg = data.get(
                        "error", "Invalid username or password")
                    return {"success": False, "message": f"Login failed: {error_msg}"}
            else:
                return {"success": False, "message": f"Authentication service error (HTTP {response.status_code})"}

        except Exception as e:
            return {"success": False, "message": f"Login error: {str(e)}"}

    def create_user(self, username: str, password: str) -> dict:
        """Create new user account"""
        try:
            auth_url = st.secrets.get("supabase", {}).get(
                "auth_function_url", f"{self.supabase_url}/functions/v1/auth-manager")
            response = requests.post(
                auth_url,
                headers={
                    "Authorization": f"Bearer {self.supabase_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "action": "create_user",
                    "username": username,
                    "password": password,
                    "created_at": datetime.datetime.now().isoformat()
                },
                timeout=10
            )

            if response.status_code == 200:
                data = response.json()
                if data.get("success"):
                    return {"success": True, "message": "Account created successfully"}
                else:
                    return {"success": False, "message": data.get("error", "Failed to create account")}
            else:
                error_data = response.json() if response.headers.get(
                    'content-type', '').startswith('application/json') else {}
                return {"success": False, "message": error_data.get("error", "Failed to create account")}

        except Exception as e:
            return {"success": False, "message": f"Registration error: {str(e)}"}

    def quick_login_bypass(self):
        """Quick demo access"""
        import uuid
        user_id = str(uuid.uuid4())
        st.session_state.authenticated = True
        st.session_state.user_id = user_id
        st.session_state.username = "demo_user"
        st.session_state.session_expires = (
            datetime.datetime.now() + datetime.timedelta(days=2)).isoformat()
        # Create persistent session token for demo user too
        session_token = self.create_session_token("demo_user", user_id)
        st.query_params["session_token"] = session_token
        st.rerun()

    def logout(self):
        """Logout user and clear session"""
        st.session_state.authenticated = False
        st.session_state.user_id = None
        st.session_state.username = None
        st.session_state.session_expires = None

        # Clear session token from URL
        if "session_token" in st.query_params:
            del st.query_params["session_token"]

        st.rerun()

    def render_splash_interface(self):
        """Render the clean SAOYNX splash interface"""

        # Custom CSS matching your design
        st.markdown("""
        <style>
        /* Hide Streamlit default elements */
        .stDeployButton {display: none;}
        header[data-testid="stHeader"] {display: none;}
        .stMainBlockContainer {padding-top: 1rem;}
        .main .block-container {padding-top: 1rem; padding-bottom: 1rem;}
        
        /* Main container - professional centering */
        .main-splash {
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            width: 400px;
            background: white;
            padding: 2rem;
            text-align: center;
            border-radius: 8px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
        }
        
        /* Logo section */
        .logo-section {
            margin-bottom: 1rem;
            width: 100%;
        }
        
        .logo-text {
            font-size: 2rem;
            font-weight: 300;
            letter-spacing: 4px;
            color: #2E2E2E;
            margin: 0.5rem 0 0.2rem 0;
            font-family: 'Arial', sans-serif;
            text-align: center;
        }
        
        .logo-subtitle {
            font-size: 1rem;
            color: #666;
            letter-spacing: 2px;
            font-weight: 300;
            text-transform: uppercase;
            line-height: 1.2;
            text-align: center;
        }
        
        /* Auth sections - tight compact layout */
        .auth-container {
            margin: 0.5rem 0;
        }
        
        .auth-question {
            font-size: 0.9rem;
            color: #666;
            margin-bottom: 0.3rem;
            text-align: center;
        }
        
        .auth-title {
            font-size: 1.1rem;
            color: #2E2E2E;
            margin-bottom: 1rem;
            font-weight: 300;
        }
        
        /* Form container centering - compact */
        .form-container {
            max-width: 350px;
            margin: 1rem auto;
            padding: 1rem;
            text-align: center;
        }
        
        /* Form styling - ultra tight spacing */
        .stTextInput > div > div > input {
            border-radius: 30px;
            border: 2px solid #ddd;
            padding: 6px 12px;
            font-size: 0.85rem;
            text-align: center;
            height: 30px;
            line-height: 0.6;
        }
        
        /* Minimal form spacing */
        .stTextInput {
            margin-bottom: 0.1rem;
            line-height: 0.6;
        }
        
        /* Compact form layout */
        .stTextInput > label {
            margin-bottom: 0.1rem !important;
            line-height: 0.6;
        }
        
        .stTextInput > div > div > input:focus {
            border-color: #666;
            box-shadow: 0 0 0 1px #666;
        }
        
        /* Professional button styling */
        .stButton > button {
            width: 100%;
            max-width: 140px;
            height: 35px;
            border: 2px solid #2E2E2E;
            background-color: #2E2E2E;
            border-radius: 8px;
            font-size: 0.8rem;
            color: white;
            font-weight: 600;
            cursor: pointer;
            outline: none;
            transition: all 0.3s ease;
            text-transform: uppercase;
            letter-spacing: 1px;
        }
        
        .stButton > button:hover {
            background-color: white;
            color: #2E2E2E;
            border-color: #2E2E2E;
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(46, 46, 46, 0.2);
        }
        
        /* Quick access - clean divider */
        .quick-access {
            margin-top: 2rem;
            padding-top: 1rem;
            border-top: 1px solid #ddd;
            text-align: center;
        }
        
        .quick-title {
            color: #999;
            font-size: 0.8rem;
            margin-bottom: 0.8rem;
            text-transform: uppercase;
            letter-spacing: 1px;
        }
        </style>
        """, unsafe_allow_html=True)

        # Complete logo with text - perfectly centered
        st.markdown(
            '<div style="text-align: center; margin-bottom: 1rem;">', unsafe_allow_html=True)

        col1, col2, col3 = st.columns([1, 1, 1])
        with col2:
            try:
                st.image(
                    "/static/graphics/FirstPerson-Logo-normalized.svg", width=120)
            except Exception:
                st.markdown('''
                <div style="font-size: 4rem;">🧠</div>
                <div style="font-size: 2rem; font-weight: 300; letter-spacing: 4px; color: #2E2E2E; margin: 0.5rem 0 0.2rem 0;">FirstPerson</div>
                <div style="font-size: 1rem; color: #666; letter-spacing: 2px; font-weight: 300; text-transform: uppercase;">Personal AI<br>Companion</div>
                ''', unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)

        # Handle different states
        if st.session_state.get('show_login', False):
            self.render_login_form()
        elif st.session_state.get('show_register', False):
            self.render_register_form()
        else:
            # Main splash screen - professional centered layout
            st.markdown('<div class="auth-container">', unsafe_allow_html=True)

            # Tight, compact button layout
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                # Questions closer together
                qcol1, qcol2 = st.columns([1, 1], gap="small")
                with qcol1:
                    st.markdown(
                        '<div class="auth-question">existing user?</div>', unsafe_allow_html=True)
                with qcol2:
                    st.markdown(
                        '<div class="auth-question">new user?</div>', unsafe_allow_html=True)

                # Buttons closer together
                bcol1, bcol2 = st.columns([1, 1], gap="small")
                with bcol1:
                    if st.button("Sign In", key="existing_btn"):
                        st.session_state.show_login = True
                        st.rerun()
                with bcol2:
                    if st.button("Register Now", key="new_btn"):
                        st.session_state.show_register = True
                        st.rerun()

            st.markdown('</div>', unsafe_allow_html=True)

            # Quick access section
            st.markdown('<div class="quick-access">', unsafe_allow_html=True)
            st.markdown('<div class="quick-title">Quick Access</div>',
                        unsafe_allow_html=True)

            col1, col2, col3 = st.columns([1, 1, 1])
            with col2:
                if st.button("⚡ Demo Mode", help="Try the system instantly"):
                    self.quick_login_bypass()

            st.markdown('</div>', unsafe_allow_html=True)

    def render_login_form(self):
        """Clean login form matching your design"""

        st.markdown('<div class="form-container">', unsafe_allow_html=True)

        # Back button
        if st.button("← Back", key="back_login"):
            st.session_state.show_login = False
            st.rerun()

        # No duplicate logo - just the title
        st.markdown(
            '<div class="auth-title" style="text-align: center; margin: 1rem 0;">Sign In</div>', unsafe_allow_html=True)

        # Login form - ultra compact
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            with st.form("login_form"):
                # Tight spacing labels
                st.markdown(
                    '<div style="text-align: left; margin-bottom: 0.1rem; color: #666; font-size: 0.8rem;">login:</div>', unsafe_allow_html=True)
                username = st.text_input(
                    "login", label_visibility="collapsed", placeholder="username")

                st.markdown(
                    '<div style="text-align: left; margin-bottom: 0.1rem; margin-top: 0.1rem; color: #666; font-size: 0.8rem;">password:</div>', unsafe_allow_html=True)
                password = st.text_input(
                    "password", label_visibility="collapsed", type="password", placeholder="password")

                st.markdown(
                    '<div style="margin: 0.2rem 0 0.1rem 0;"></div>', unsafe_allow_html=True)
                login_submitted = st.form_submit_button(
                    "Sign In", use_container_width=True)

                if login_submitted:
                    if not username or not password:
                        st.error("Please enter both username and password")
                    else:
                        with st.spinner("Signing in..."):
                            result = self.authenticate_user(username, password)

                        if result["success"]:
                            st.success("Welcome back!")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error(result["message"])

        st.markdown('</div>', unsafe_allow_html=True)

    def render_register_form(self):
        """Clean registration form"""

        st.markdown('<div class="form-container">', unsafe_allow_html=True)

        # Back button
        if st.button("← Back", key="back_register"):
            st.session_state.show_register = False
            st.rerun()

        # No duplicate logo - just the title
        st.markdown(
            '<div class="auth-title" style="text-align: center; margin: 1rem 0;">Create Account</div>', unsafe_allow_html=True)

        # Register form - ultra compact
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            with st.form("register_form"):
                # Tight spacing labels
                st.markdown(
                    '<div style="text-align: left; margin-bottom: 0.1rem; color: #666; font-size: 0.8rem;">login:</div>', unsafe_allow_html=True)
                username = st.text_input(
                    "login", label_visibility="collapsed", placeholder="choose username")

                st.markdown(
                    '<div style="text-align: left; margin-bottom: 0.1rem; margin-top: 0.1rem; color: #666; font-size: 0.8rem;">password:</div>', unsafe_allow_html=True)
                password = st.text_input(
                    "password", label_visibility="collapsed", type="password", placeholder="choose password")

                st.markdown(
                    '<div style="text-align: left; margin-bottom: 0.1rem; margin-top: 0.1rem; color: #666; font-size: 0.8rem;">confirm:</div>', unsafe_allow_html=True)
                confirm_password = st.text_input(
                    "confirm", label_visibility="collapsed", type="password", placeholder="confirm password")

                st.markdown(
                    '<div style="margin: 0.2rem 0 0.1rem 0;"></div>', unsafe_allow_html=True)
                register_submitted = st.form_submit_button(
                    "Register Now", use_container_width=True)

                if register_submitted:
                    if not username or not password:
                        st.error("Username and password are required")
                    elif password != confirm_password:
                        st.error("Passwords do not match")
                    elif len(password) < 6:
                        st.error("Password must be at least 6 characters")
                    else:
                        with st.spinner("Creating your account..."):
                            result = self.create_user(username, password)

                        if result["success"]:
                            st.success("Account created! Please sign in.")
                            st.session_state.show_register = False
                            st.session_state.show_login = True
                            time.sleep(2)
                            st.rerun()
                        else:
                            st.error(result["message"])

        st.markdown('</div>', unsafe_allow_html=True)


def render_main_app():
    """Main app interface for authenticated users - Full Emotional OS"""

    # Header with logo and title very close together
    col1, col2 = st.columns([0.5, 8], gap="small")
    with col1:
        try:
            st.image("/static/graphics/FirstPerson-Logo-normalized.svg", width=24)
        except Exception:
            st.markdown(
                '<div style="font-size: 2.5rem; margin: 0; line-height: 1;">🧠</div>', unsafe_allow_html=True)
    with col2:
        st.markdown('<h1 style="margin: 0; margin-left: -35px; padding-top: 10px; color: #2E2E2E; font-weight: 300; letter-spacing: 2px; font-size: 2.2rem;">FirstPerson - Personal AI Companion</h1>', unsafe_allow_html=True)

    st.markdown("*Your private space for emotional processing and growth*")

    # User header with logout
    col1, col2, col3 = st.columns([2, 1, 1])

    with col1:
        st.write(f"Welcome back, **{st.session_state.username}**! 👋")

    with col2:
        if st.button("⚙️ Settings", help="User settings and preferences"):
            st.info("Settings panel coming soon!")

    with col3:
        if st.button("🚪 Logout", help="Sign out of your account"):
            auth = SaoynxAuthentication()
            auth.logout()

    # User-specific conversation history
    conversation_key = f"conversation_history_{st.session_state.user_id}"
    if conversation_key not in st.session_state:
        st.session_state[conversation_key] = []

    # Main interface
    with st.container():
        # Processing mode selection
        col1, col2 = st.columns([2, 1])

    with col1:
        # Set default processing mode if not already set
        if 'processing_mode' not in st.session_state:
            st.session_state.processing_mode = "hybrid"  # or "regular" if that's your label

        # Render selectbox using session state
        processing_mode = st.selectbox(
            "Processing Mode",
            ["hybrid", "local", "ai_preferred"],
            index=["hybrid", "local", "ai_preferred"].index(
                st.session_state.processing_mode),
            help="Hybrid: Best performance, Local: Maximum privacy, AI: Most sophisticated responses"
        )

        # Update session state with selected mode
        st.session_state.processing_mode = processing_mode

    with col2:
        if st.button("Clear History", type="secondary"):
            st.session_state[conversation_key] = []
            st.rerun()

    # Chat interface
    chat_container = st.container()

    # Display conversation history
    with chat_container:
        for i, exchange in enumerate(st.session_state[conversation_key]):
            with st.chat_message("user"):
                st.write(exchange["user"])
            with st.chat_message("assistant"):
                st.write(exchange["assistant"])
                if "processing_time" in exchange:
                    st.caption(
                        f"Processed in {exchange['processing_time']} • Mode: {exchange.get('mode', 'unknown')}")

    # ✅ Process uploaded document if present
    if "uploaded_text" in st.session_state:
        with chat_container:
            with st.chat_message("user"):
                st.write("📄 Document uploaded for processing.")
            with st.chat_message("assistant"):
                with st.spinner("Metabolizing document..."):
                    response = f"Processed document content:\n\n{st.session_state['uploaded_text'][:500]}..."
                    st.write(response)

    # ✅ Document upload for emotional processing
    uploaded_file = st.file_uploader(
        "📄 Upload a document", type=["txt", "docx", "pdf"])
    if uploaded_file:
        file_content = uploaded_file.read().decode(
            "utf-8", errors="ignore")  # Adjust for docx/pdf later
        st.session_state["uploaded_text"] = file_content
        st.success("Document uploaded successfully!")

    # ✅ Input area
    user_input = st.chat_input("Share what you're feeling...")
    if user_input:
        with chat_container:
            with st.chat_message("user"):
                st.write(user_input)
            with st.chat_message("assistant"):
                with st.spinner("Processing your emotional input..."):
                    start_time = time.time()
                    try:
                        saori_url = st.secrets["supabase"]["saori_function_url"]
                        # Continue with your Saori API call...
                        response_data = requests.post(
                            saori_url,
                            headers={
                                "Authorization": f"Bearer {st.secrets['supabase']['key']}",
                                "Content-Type": "application/json"
                            },
                            json={
                                "message": user_input,
                                "mode": processing_mode,
                                "user_id": st.session_state.user_id
                            },
                            timeout=15
                        )
                        if response_data.status_code == 200:
                            result = response_data.json()
                            response = result.get(
                                "reply", "I'm here to listen.")
                            glyph_info = result.get("glyph", {})
                            processing_details = result.get("log", {})  # noqa: F841  # optional debug/logging info
                        else:
                            response = "I'm experiencing some technical difficulties, but I'm still here for you."
                            glyph_info = {}
                            processing_details = {"error": f"HTTP {response_data.status_code}"}  # noqa: F841  # optional debug/logging info
                    except Exception as e:
                        response = "I'm having trouble connecting right now, but your feelings are still valid and important."
                        glyph_info = {}
                        processing_details = {"error": str(e)}  # noqa: F841  # optional debug/logging info
                    processing_time = time.time() - start_time
                    st.write(response)
                    # Show processing details
                    if glyph_info and glyph_info.get("tag_name"):
                        st.caption(
                            f"✨ Emotional resonance: {glyph_info.get('tag_name')} • Processed in {processing_time:.2f}s • Mode: {processing_mode}")
                    else:
                        st.caption(
                            f"Processed in {processing_time:.2f}s • Mode: {processing_mode}")
        # Add to conversation history
        st.session_state[conversation_key].append({
            "user": user_input,
            "assistant": response,
            "processing_time": f"{processing_time:.2f}s",
            "mode": processing_mode,
            "glyph": glyph_info.get("tag_name") if glyph_info else None,
            "timestamp": datetime.datetime.now().isoformat()
        })
        st.rerun()

    # Personal Log Button
    if "show_personal_log" not in st.session_state:
        st.session_state.show_personal_log = False
    if st.button("Start Personal Log", help="Begin a structured emotional entry—date, event, mood, reflections, and insights."):
        st.session_state.show_personal_log = True
        st.rerun()

    if st.session_state.show_personal_log:
        import datetime as dt
        st.markdown("### 📘 Start Personal Log")
        st.markdown(
            "Use this space to record a structured emotional entry. You'll log the date, event, mood, reflections, and insights.")

        # Auto-filled date and time
        date = st.date_input("Date", value=dt.date.today())
        log_time = st.time_input("Time", value=dt.datetime.now().time())

        # Event description
        event = st.text_area("Event", placeholder="What happened?")

        # Mood (user-defined or AI-suggested)
        mood = st.text_input("Mood", placeholder="How did it feel?")

        # Reflections (freeform or guided)
        reflections = st.text_area(
            "Reflections", placeholder="What’s emerging emotionally?")

        # Insights (symbolic clarity)
        insights = st.text_area(
            "Insights", placeholder="What truth or clarity surfaced?")

        # Completion prompt
        def generate_doc(date, log_time, event, mood, reflections, insights):
            from docx import Document
            import io
            doc = Document()
            doc.add_heading("Personal Log", 0)
            doc.add_paragraph(f"Date: {date}")
            doc.add_paragraph(f"Time: {log_time}")
            doc.add_paragraph(f"Event: {event}")
            doc.add_paragraph(f"Mood: {mood}")
            doc.add_paragraph(f"Reflections: {reflections}")
            doc.add_paragraph(f"Insights: {insights}")
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()

        if st.button("Conclude Log"):
            st.success("Your personal log has been saved.")
            try:
                st.download_button(
                    label="Download as Word Doc",
                    data=generate_doc(date, log_time, event,
                                      mood, reflections, insights),
                    file_name="personal_log.docx"
                )
            except Exception as e:
                st.error(f"Error generating Word document: {e}")
        else:
            st.info("You can continue adding to this log.")

    # Sidebar with user stats and settings
    with st.sidebar:
        st.subheader("Your Emotional Journey")

        # User-specific stats
        conversation_count = len(st.session_state[conversation_key])
        st.metric("Conversations", conversation_count)

        if conversation_count > 0:
            recent_conversation = st.session_state[conversation_key][-1]
            st.metric("Last Session", recent_conversation["timestamp"][:10])

        st.subheader("Privacy Settings")
        st.write("🔒 Your data is completely isolated")
        st.write("🧠 Learning happens only from your conversations")
        st.write("⚡ Optimized for 2-3 second responses")

        if st.button("Download My Data", type="secondary"):
            # Option to export user's conversation data
            user_data = {
                "user_id": st.session_state.user_id,
                "username": st.session_state.username,
                "conversations": st.session_state[conversation_key],
                "export_date": datetime.datetime.now().isoformat()
            }
            st.download_button(
                "Download JSON",
                json.dumps(user_data, indent=2),
                file_name=f"emotional_os_data_{st.session_state.username}_{datetime.datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json"
            )


def main():
    """Main application entry point"""
    auth = SaoynxAuthentication()

    # Check if user is authenticated
    if st.session_state.authenticated:
        render_main_app()
    else:
        auth.render_splash_interface()


if __name__ == "__main__":
    main()
