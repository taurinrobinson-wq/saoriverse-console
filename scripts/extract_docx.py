#!/usr/bin/env python3
"""Simple DOCX -> UTF-8 text extractor.
Usage: python scripts/extract_docx.py input.docx output.txt
"""
import sys
from pathlib import Path

try:
    import docx
except ImportError:
    print("python-docx not installed", file=sys.stderr)
    raise


def extract(path: Path) -> str:
    doc = docx.Document(path)
    out_lines = []
    # Extract paragraphs
    for p in doc.paragraphs:
        text = p.text
        if text is None:
            continue
        out_lines.append(text)
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            out_lines.append('\t'.join(cells))
    return "\n".join(out_lines)


def main():
    if len(sys.argv) < 3:
        print("Usage: extract_docx.py input.docx output.txt", file=sys.stderr)
        sys.exit(2)
    inp = Path(sys.argv[1])
    out = Path(sys.argv[2])
    if not inp.exists():
        print(f"Input not found: {inp}", file=sys.stderr)
        sys.exit(3)
    text = extract(inp)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8")
    print(f"Wrote {out} ({len(text.splitlines())} lines)")

if __name__ == '__main__':
    main()
