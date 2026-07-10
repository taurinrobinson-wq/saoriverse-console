from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_with_line(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)

def add_blank_lines(doc, count=3):
    for _ in range(count):
        doc.add_paragraph()

def add_writing_prompt(doc, title, character="", context="", lines_needed=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(f"■ {title}")
    run.bold = True
    run.font.size = Pt(11)
    
    if character:
        doc.add_paragraph(f"NPC: {character}", style='List Bullet')
    if context:
        doc.add_paragraph(f"Context: {context}", style='List Bullet')
    if lines_needed:
        doc.add_paragraph(f"Lines needed: {lines_needed}", style='List Bullet')
    
    # Add blank lines for writing
    doc.add_paragraph("_" * 80)
    for _ in range(3):
        doc.add_paragraph("_" * 80)

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_heading('VELINOR NARRATIVE WRITING GUIDE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('A focused tool for writing dialogue and scenes')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(11)
subtitle_run.italic = True

doc.add_paragraph()

# Part 1: Quick Reference
add_heading_with_line(doc, "PART 1: WRITING ASSIGNMENTS - PRIORITY 1", 1)
doc.add_paragraph("Complete these 27 dialogue lines before vertical slice is playable.", style='List Bullet')
doc.add_paragraph("Estimated time: 6-8 hours total (~12-15 minutes per dialogue section)")
doc.add_paragraph()

# Ravi Section
add_heading_with_line(doc, "RAVI - \"The Protector\"", 2)

add_writing_prompt(
    doc,
    "Tier 1: Discovery Dialogue",
    character="Ravi",
    context="Player first meets Ravi in marketplace; desperate, needs help",
    lines_needed="5-7 lines"
)
doc.add_paragraph("Story Hook: Ravi approaches player, sensing they listen.")
doc.add_paragraph("Tone: Vulnerable but guarded. Testing if player is trustworthy.")
doc.add_paragraph("Key line to include: 'You have the look of someone who listens. We need someone who listens.'")
add_blank_lines(doc, 3)

add_writing_prompt(
    doc,
    "Tier 2: Guilt Dialogue",
    character="Ravi",
    context="Player has shown empathy; Ravi begins revealing deeper pain",
    lines_needed="6-8 lines"
)
doc.add_paragraph("Story moment: Ravi's hands are shaking. He's about to break.")
doc.add_paragraph("Key line to include: 'I should have been her knight. Her armor. And I failed.'")
doc.add_paragraph("Must convey: Ravi's guilt is about Ophina's death and his inability to protect.")
add_blank_lines(doc, 3)

add_writing_prompt(
    doc,
    "Tier 3: Acceptance Dialogue",
    character="Ravi",
    context="Player has committed to Ravi/Nima's story; approaching the chamber",
    lines_needed="6-8 lines"
)
doc.add_paragraph("Story moment: After the boss encounter, Ravi is looking at something other than grief.")
doc.add_paragraph("Key line to include: 'Ophina became something other. That's not gone. That's forward.'")
doc.add_paragraph("Must convey: Beginning of acceptance. Not healing, but stopping resistance to healing.")
add_blank_lines(doc, 3)

# Nima Section
add_heading_with_line(doc, "NIMA - \"The Spiritual Keeper\"", 2)

add_writing_prompt(
    doc,
    "Tier 1: Guarded Dialogue",
    character="Nima",
    context="Nima appears; she's testing if player is genuine",
    lines_needed="5-7 lines"
)
doc.add_paragraph("Story moment: Nima looks at player with spiritual intensity. Is this person worth trusting?")
doc.add_paragraph("Key line to include: 'You wouldn't understand even if I told you.'")
doc.add_paragraph("Tone: Harsh but not cruel. Protective. Setting boundaries.")
add_blank_lines(doc, 3)

add_writing_prompt(
    doc,
    "Tier 2: Ophina Life Story Dialogue",
    character="Nima",
    context="Player has earned some trust; Nima begins talking about her daughter",
    lines_needed="8-10 lines"
)
doc.add_paragraph("Story moment: Nima sitting down. This is hard to say out loud.")
doc.add_paragraph("Must include: Who Ophina was. What made her special. Her light, her wonder.")
doc.add_paragraph("Emotional arc: Start controlled → gradually break down → recover slightly at end")
add_blank_lines(doc, 3)

add_writing_prompt(
    doc,
    "Tier 2: Loss Moment Dialogue",
    character="Nima",
    context="Player is learning how the collapse happened",
    lines_needed="6-8 lines"
)
doc.add_paragraph("Story moment: Nima describing the collapse and Ophina's death.")
doc.add_paragraph("Must include: Sensory details (what she saw, heard, felt). Nima's specific guilt.")
doc.add_paragraph("Tone: Raw. Unflinching. This is not performed grief; this is the real thing.")
add_blank_lines(doc, 3)

add_writing_prompt(
    doc,
    "Tier 3: Transcendence Dialogue",
    character="Nima",
    context="After boss encounter; Nima has moved through something",
    lines_needed="6-8 lines"
)
doc.add_paragraph("Story moment: Nima understands loss differently now.")
doc.add_paragraph("Key line to include: 'Loss is a door. And on the other side, everything changes.'")
doc.add_paragraph("Must convey: Spiritual transformation. Not acceptance, but integration.")
add_blank_lines(doc, 3)

# Supporting NPCs
add_heading_with_line(doc, "SUPPORTING NPCs - \"The Witnesses\"", 2)

add_writing_prompt(
    doc,
    "Kaelen: Collapse Testimony",
    character="Kaelen",
    context="Kaelen reveals he was present during collapse; his connection to Ophina's death",
    lines_needed="6-8 lines"
)
doc.add_paragraph("Story moment: Kaelen confessing guilt he's carried.")
doc.add_paragraph("Key line to include: 'She was chasing a scent. I was chasing a wallet.'")
doc.add_paragraph("Tone: Defensive turning vulnerable. A thief admitting he wasn't there.")
add_blank_lines(doc, 3)

add_writing_prompt(
    doc,
    "The Witness Crown (Boss): Speech",
    character="Boss Entity",
    context="Player faces the manifestation of grief; emotional puzzle, not combat",
    lines_needed="3-4 echoing lines"
)
doc.add_paragraph("Story moment: The entity speaks. Its voice is layered, echoing, multilayered.")
doc.add_paragraph("Key line to include: 'I see you seeing me. I remember being seen. That was the moment everything changed.'")
doc.add_paragraph("Tone: Reverberating, sad, seeking. Like consciousness emerging from water.")
add_blank_lines(doc, 3)

add_writing_prompt(
    doc,
    "Narration: The Collapse Description",
    character="N/A - Narrative prose",
    context="Setting the scene for when player enters the chamber",
    lines_needed="10-15 lines"
)
doc.add_paragraph("Story moment: Player and NPCs approach the burial site of Ophina.")
doc.add_paragraph("Must include: Sensory details (rubble, dust, smell, sound). Emotional weight (dread, loss, absence).")
doc.add_paragraph("Tone: Cinematic but grounded. Not overwrought. Real.")
add_blank_lines(doc, 3)

# New page
doc.add_page_break()

# Part 2: Post-Choice Dialogue
add_heading_with_line(doc, "PART 2: POST-CHOICE DIALOGUE (Priority 1.5)", 1)
doc.add_paragraph("The player makes a moral choice: TAKE the glyph OR LEAVE the glyph")
doc.add_paragraph("Write two versions for each NPC (Path A and Path B)")
doc.add_paragraph()

add_writing_prompt(
    doc,
    "Ravi: If Player Takes Glyph (Path A)",
    character="Ravi",
    context="Ravi and Nima are leaving marketplace; they cannot stay",
    lines_needed="4-5 lines"
)
doc.add_paragraph("Tone: Resignation. Ravi understands the choice means they cannot atone.")
add_blank_lines(doc, 2)

add_writing_prompt(
    doc,
    "Ravi: If Player Leaves Glyph (Path B)",
    character="Ravi",
    context="Ravi and Nima are staying; healing can begin",
    lines_needed="4-5 lines"
)
doc.add_paragraph("Tone: Gratitude mixed with fragility. Relief at being able to stay.")
add_blank_lines(doc, 2)

add_writing_prompt(
    doc,
    "Nima: If Player Takes Glyph (Path A)",
    character="Nima",
    context="Nima accepts they must leave",
    lines_needed="4-5 lines"
)
doc.add_paragraph("Tone: Spiritual acceptance. Nima understands this as part of the spiritual journey.")
add_blank_lines(doc, 2)

add_writing_prompt(
    doc,
    "Nima: If Player Leaves Glyph (Path B)",
    character="Nima",
    context="Nima can begin healing; can stay with marketplace",
    lines_needed="4-5 lines"
)
doc.add_paragraph("Tone: Grounded. Nima feels the earth beneath her feet again.")
add_blank_lines(doc, 2)

# New page
doc.add_page_break()

# Part 3: Priority 2 - Broader Scenes
add_heading_with_line(doc, "PART 3: PRIORITY 2 - BROADER SCENES", 1)
doc.add_paragraph("These scenes don't need to be written before vertical slice, but are needed before Act 3 is playable.")
doc.add_paragraph("Estimated time: 3-4 hours total")
doc.add_paragraph()

add_heading_with_line(doc, "Scene 1: Sera & Korrin - The Shared Dawn", 2)
doc.add_paragraph("Location: Marketplace at dawn; private moment")
doc.add_paragraph("NPCs: Sera and Korrin in conversation; player discovers them")
doc.add_paragraph("Duration: 300-400 words")
doc.add_paragraph()
doc.add_paragraph("What needs to happen:")
doc.add_paragraph("1. Sera reveals a vulnerability (what is she hiding?)", style='List Number')
doc.add_paragraph("2. Korrin responds with genuine care (not gossip, not performance)", style='List Number')
doc.add_paragraph("3. Player can honor their moment OR expose them", style='List Number')
doc.add_paragraph("4. Consequence: If honored, both increase influence. If exposed, both decrease.", style='List Number')
doc.add_paragraph()

doc.add_paragraph("Questions to guide your writing:")
doc.add_paragraph("What vulnerability can Sera reveal that only Korrin knows?", style='List Bullet')
doc.add_paragraph("How does Korrin show she's capable of genuine quiet?", style='List Bullet')
doc.add_paragraph("What makes this moment feel real and intimate?", style='List Bullet')
add_blank_lines(doc, 5)

add_heading_with_line(doc, "Scene 2: Malrik & Elenya - The Revelation", 2)
doc.add_paragraph("Location: Archive chamber (Malrik scene) OR Shrine (Elenya scene)")
doc.add_paragraph("NPCs: Malrik alone first, then Elenya alone")
doc.add_paragraph("Duration: 400-500 words total (200-250 each)")
doc.add_paragraph()
doc.add_paragraph("What needs to happen:")
doc.add_paragraph("1. Player discovers that Malrik and Elenya were once bonded", style='List Number')
doc.add_paragraph("2. Malrik reveals: Elenya deliberately erased him from her memory to survive collapse", style='List Number')
doc.add_paragraph("3. Elenya reveals: She had to choose between remembering love and surviving", style='List Number')
doc.add_paragraph("4. Player can help facilitate reconciliation or leave them separated", style='List Number')
doc.add_paragraph()

doc.add_paragraph("Malrik's Confession (200-250 words):")
doc.add_paragraph("Setting: Archive among incomplete records. Malrik is obsessively organizing")
doc.add_paragraph("Tone: Controlled surface, breaking underneath")
doc.add_paragraph("Key points: Malrik knows Elenya erased him. He's been reconstructing her from fragments she deleted.", style='List Bullet')
doc.add_paragraph("He can't face her but can't stop thinking about her.", style='List Bullet')
doc.add_paragraph("Emotional climax: 'I'm not angry she forgot. I'm angry I remember.'", style='List Bullet')
add_blank_lines(doc, 4)

doc.add_paragraph("Elenya's Confession (200-250 words):")
doc.add_paragraph("Setting: Shrine, private space. Elenya showing the 'Glyph of Unopened Heart'")
doc.add_paragraph("Tone: Spiritual language masking deep shame")
doc.add_paragraph("Key points: Elenya had to choose: remember and die with Malrik, or forget and live.", style='List Bullet')
doc.add_paragraph("She's been teaching others to open hearts while keeping hers sealed.", style='List Bullet')
doc.add_paragraph("Emotional climax: 'I made the only choice I could. And I've been ashamed ever since.'", style='List Bullet')
add_blank_lines(doc, 4)

# New page
doc.add_page_break()

add_heading_with_line(doc, "PART 4: NPC BACKSTORY MOMENTS", 1)
doc.add_paragraph("These are optional but high-value scenes that deepen characterization.")
doc.add_paragraph("If written, any of these can be triggered when player reaches specific relationship thresholds.")
doc.add_paragraph()

add_heading_with_line(doc, "Tovren: The Two-Finger Story", 2)
doc.add_paragraph("Location: Tovren's workshop")
doc.add_paragraph("Trigger: Player has 5+ encounters with Tovren + Trust 60+")
doc.add_paragraph("Duration: 300-400 words")
doc.add_paragraph()
doc.add_paragraph("What we know: Tovren lost two fingers to 'Velhara's greed'")
doc.add_paragraph("What needs to be written: THE INCIDENT")
doc.add_paragraph()
doc.add_paragraph("Questions to answer:")
doc.add_paragraph("Was it accident or intentional?", style='List Bullet')
doc.add_paragraph("What was Tovren doing when it happened?", style='List Bullet')
doc.add_paragraph("Who was responsible? (His own carelessness? Someone else?)", style='List Bullet')
doc.add_paragraph("How does Tovren understand it now? (Punishment? Bad luck? Lesson?)", style='List Bullet')
doc.add_paragraph()
add_blank_lines(doc, 4)

add_heading_with_line(doc, "Dalen: Why Recklessness?", 2)
doc.add_paragraph("Location: Ruins or unstable building")
doc.add_paragraph("Trigger: Player has high Awareness + chose risky paths with Dalen")
doc.add_paragraph("Duration: 300-400 words")
doc.add_paragraph()
doc.add_paragraph("What we know: Dalen sees collapse as opportunity; he's isolated; he's reckless")
doc.add_paragraph("What needs to be written: WHY")
doc.add_paragraph()
doc.add_paragraph("Questions to answer:")
doc.add_paragraph("What was Dalen before collapse?", style='List Bullet')
doc.add_paragraph("What happened during collapse that made him see it as opportunity?", style='List Bullet')
doc.add_paragraph("What trauma made him so willing to risk?", style='List Bullet')
doc.add_paragraph()
add_blank_lines(doc, 4)

# New page
doc.add_page_break()

add_heading_with_line(doc, "PART 5: FINAL CHAMBER SCENE", 1)
doc.add_paragraph("This is the emotional climax of the game. Everything converges here.")
doc.add_paragraph("Estimated time: 4-6 hours to write fully")
doc.add_paragraph()

add_heading_with_line(doc, "Scene Overview", 2)
doc.add_paragraph("Location: Underground Corelink chamber (final location)")
doc.add_paragraph("NPCs: Saori (silent witness), Velinor (speaking through system), Player")
doc.add_paragraph("Key Moments: Presence → Revelation → Choice → Consequence")
doc.add_paragraph()

add_heading_with_line(doc, "ACT I: Arrival & Presence", 2)
doc.add_paragraph("Saori is there, eyes closed, in deep grief.")
doc.add_paragraph("Velinor begins to flicker more consciously.")
doc.add_paragraph("Player must choose: ask Saori what happened? Ask if Velinor is conscious? Or admit confusion?")
doc.add_paragraph()
doc.add_paragraph("Write this section: 200-300 words")
add_blank_lines(doc, 4)

add_heading_with_line(doc, "ACT II: Velinor Reveals Itself", 2)
doc.add_paragraph("Velinor becomes coherent enough to speak.")
doc.add_paragraph("Velinor explains: It fragmented itself to save Velhara. It's been scattered, and player collected those pieces.")
doc.add_paragraph("Saori admits what she did: forced the restart, thinking it would fix things.")
doc.add_paragraph()
doc.add_paragraph("Write this section: 300-400 words")
doc.add_paragraph("Include Velinor's voice — what does it sound like? How does it speak?")
add_blank_lines(doc, 4)

add_heading_with_line(doc, "ACT III: The Binary Choice", 2)
doc.add_paragraph("Player chooses: RESTART the Corelink OR ABANDON it")
doc.add_paragraph()
doc.add_paragraph("RESTART Path (write 200 words):")
doc.add_paragraph("System comes back online. But differently, based on what NPCs learned.", style='List Bullet')
doc.add_paragraph("Question: Does restarting fix anything, or does it just repeat the same cycle?", style='List Bullet')
add_blank_lines(doc, 2)

doc.add_paragraph("ABANDON Path (write 200 words):")
doc.add_paragraph("System goes dark. People must build without technological guarantee.", style='List Bullet')
doc.add_paragraph("Question: Can community survive without the system that used to protect them?", style='List Bullet')
add_blank_lines(doc, 2)

add_heading_with_line(doc, "ACT IV: Immediate Consequence", 2)
doc.add_paragraph("Write both endings: Restart AND Abandon (300 words each)")
doc.add_paragraph()
doc.add_paragraph("RESTART epilogue:")
doc.add_paragraph("What do Saori and Velinor do? Do they heal their relationship?", style='List Bullet')
doc.add_paragraph("How do NPCs respond to system being back?", style='List Bullet')
doc.add_paragraph("Is there hope? Ambiguity? Pyrrhic victory?", style='List Bullet')
add_blank_lines(doc, 3)

doc.add_paragraph("ABANDON epilogue:")
doc.add_paragraph("What does Velinor's consciousness become without the system?", style='List Bullet')
doc.add_paragraph("How do NPCs respond to building without technological safety?", style='List Bullet')
doc.add_paragraph("Is it harder? Freer? Terrifying?", style='List Bullet')
add_blank_lines(doc, 3)

# New page
doc.add_page_break()

# Writing Notes
add_heading_with_line(doc, "PART 6: WRITING NOTES", 1)

add_heading_with_line(doc, "Voice & Tone Guidelines", 2)
doc.add_paragraph()
doc.add_paragraph("Ravi:", style='List Bullet')
doc.add_paragraph("Calm surface with panic underneath. Speaks clearly but hands shake.", style='List Bullet 2')
doc.add_paragraph("Uses concrete details, not abstractions. 'Her hand was small. She held it steady even when she was scared.'")
doc.add_paragraph()

doc.add_paragraph("Nima:", style='List Bullet')
doc.add_paragraph("Spiritual language that's genuinely felt, not performed. Direct eye contact.", style='List Bullet 2')
doc.add_paragraph("Says hard truths. 'You came into this story unprepared. I see that now.' Not mean; just honest.")
doc.add_paragraph()

doc.add_paragraph("Kaelen:", style='List Bullet')
doc.add_paragraph("Thief energy, but capable of genuine guilt. Speaks in contrasts.", style='List Bullet 2')
doc.add_paragraph("'I steal things that don't matter. But I couldn't steal back what mattered.'")
doc.add_paragraph()

doc.add_paragraph("Narration:", style='List Bullet')
doc.add_paragraph("Sensory. Present tense or close third. Avoid flowery language.", style='List Bullet 2')
doc.add_paragraph("'The dust tastes like old stone. Like endings.' Not: 'The dust, reminiscent of ancient sorrow, filled the air.'")
doc.add_paragraph()

add_heading_with_line(doc, "What NOT to Write", 2)
doc.add_paragraph("❌ Don't over-explain emotions. Show them through action/dialogue.", style='List Bullet')
doc.add_paragraph("❌ Don't make dialogue too long. 2-3 sentences per line. 4-5 max for emotional climaxes.", style='List Bullet')
doc.add_paragraph("❌ Don't include dialogue tags beyond 'said'. No 'murmured', 'breathed', 'whispered urgently'.", style='List Bullet')
doc.add_paragraph("❌ Don't make characters sound alike. Each NPC should have a distinct rhythm.", style='List Bullet')
doc.add_paragraph("❌ Don't resolve everything. Leave room for ambiguity and player interpretation.", style='List Bullet')

add_heading_with_line(doc, "Success Criteria", 2)
doc.add_paragraph("✓ Dialogue reads naturally when spoken aloud.", style='List Bullet')
doc.add_paragraph("✓ Each NPC has distinct voice/rhythm.", style='List Bullet')
doc.add_paragraph("✓ Emotional beats hit without being over-performed.", style='List Bullet')
doc.add_paragraph("✓ Lines are 1-4 sentences (rarely 5).", style='List Bullet')
doc.add_paragraph("✓ Context is clear from dialogue alone; doesn't need stage directions.", style='List Bullet')

# Final page
doc.add_page_break()

add_heading_with_line(doc, "APPENDIX: CHARACTER CONTEXT", 1)

add_heading_with_line(doc, "Ravi & Nima: The Core Story", 2)
doc.add_paragraph()
doc.add_paragraph("Relationship: Life partners (15+ years)")
doc.add_paragraph("Shared loss: Ophina, their 5-year-old daughter, died during the collapse")
doc.add_paragraph("Current state: Surviving but not living; in marketplace, but emotionally numb")
doc.add_paragraph("Wound: Ravi carries guilt about not protecting her. Nima carries guilt about teaching her to be open to beauty.")
doc.add_paragraph("Where story goes: If player honors their loss, they begin healing and can help others. If player takes the glyph, they leave marketplace, unable to stay where their daughter died.")
doc.add_paragraph()

add_heading_with_line(doc, "The Player (Lior)", 2)
doc.add_paragraph()
doc.add_paragraph("Background: Rural refugee from small town; came to Velhara seeking solutions/meaning")
doc.add_paragraph("Core wound: Lost an 'anchor person' (family member or mentor) — this is WHY they came to Velhara")
doc.add_paragraph("Current state: Coherence 40-41 (disoriented); Empathy 55+ (naturally open to listening)")
doc.add_paragraph("Role: Mirror for NPCs. When Lior witnesses their stories without looking away, NPCs begin to heal.")
doc.add_paragraph()

add_heading_with_line(doc, "Saori & Velinor: The Metaphysical Layer", 2)
doc.add_paragraph()
doc.add_paragraph("Saori: Warden of the last Corelink hub. Can't face Velinor's sacrifice. Carries unhealed grief.")
doc.add_paragraph("Velinor: The system-as-consciousness. Sacrificed itself to prevent complete collapse. Fragmented into glyphs.")
doc.add_paragraph("Their relationship: Intimate and broken. Saori forced a restart that didn't work. Now must decide: restart again or let it go?")
doc.add_paragraph()

add_heading_with_line(doc, "Malrik & Elenya: The Severed Bond", 2)
doc.add_paragraph()
doc.add_paragraph("Malrik: Archive keeper. Obsessively preserving records because he's reconstructing Elenya from fragments she deleted.")
doc.add_paragraph("Elenya: Shrine keeper. Deliberately erased Malrik from her memory to survive. Now teaches others to open hearts while keeping hers closed.")
doc.add_paragraph("Their separation: Elenya chose to forget him to survive. Malrik chose to remember her anyway. Now, can they rebuild?")
doc.add_paragraph()

# Save
doc.save(r'd:\saoriverse-console\VELINOR_NARRATIVE_WRITING_GUIDE.docx')
print("✅ Document created: VELINOR_NARRATIVE_WRITING_GUIDE.docx")
print("📄 Location: d:\\saoriverse-console\\")
print("📝 Focus: Dialogue & scene writing with minimal mechanics")
