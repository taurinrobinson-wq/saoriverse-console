#!/usr/bin/env python3
"""
VELΩNIX Glyph Processing Module - Ritual Capsule Ingestion System

Automates the sacred transformation of raw glyph notes into structured emotional objects.
Treats unprocessed_glyphs as ritual capsules containing emotional fragments, voltage markers,
gate metaphors, and lineage annotations.

Author: Emotional OS - Saoriverse Console
"""

import os
import json
import sqlite3
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import docx
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class GlyphObject:
    """Structured glyph object with ritual context and lineage"""

    def __init__(self):
        self.name: Optional[str] = None
        self.category: Optional[str] = None
        self.valence: Optional[str] = None
        self.function: Optional[str] = None
        self.description: Optional[str] = None
        self.emotional_signals: List[str] = []
        self.voltage_markers: List[str] = []
        self.gates: List[str] = []
        self.voltage_pair: Optional[str] = None
        self.activation_signals: List[str] = []

        # Lineage and context
        self.origin_file: Optional[str] = None
        self.timestamp: Optional[str] = None
        self.author: Optional[str] = None
        self.legacy_context: Optional[str] = None
        self.ritual_annotations: List[str] = []

    def to_dict(self) -> Dict:
        """Convert to dictionary for JSON serialization"""
        return {
            'name': self.name,
            'category': self.category,
            'valence': self.valence,
            'function': self.function,
            'description': self.description,
            'emotional_signals': self.emotional_signals,
            'voltage_markers': self.voltage_markers,
            'gates': self.gates,
            'voltage_pair': self.voltage_pair,
            'activation_signals': self.activation_signals,
            'lineage': {
                'origin_file': self.origin_file,
                'timestamp': self.timestamp,
                'author': self.author,
                'legacy_context': self.legacy_context,
                'ritual_annotations': self.ritual_annotations,
            },
        }


class RitualCapsuleProcessor:
    """Sacred processor of glyph ritual capsules"""

    def __init__(self,
                 unprocessed_dir: str = "emotional_os/deploy/unprocessed_glyphs",
                 processed_dir: str = "emotional_os/deploy/processed_glyphs",
                 db_path: str = "emotional_os/glyphs/glyphs.db"):

        self.unprocessed_dir = Path(unprocessed_dir)
        self.processed_dir = Path(processed_dir)
        self.db_path = db_path

        # Ensure directories exist
        self.processed_dir.mkdir(parents=True, exist_ok=True)

        # Emotional signal patterns
        self.emotional_signals = {
            'ache', 'longing', 'yearning', 'grief', 'mourning', 'sorrow',
            'joy', 'bliss', 'ecstasy', 'stillness', 'silence', 'quiet',
            'recognition', 'witness', 'seen', 'devotion', 'sacred', 'reverent',
            'boundary', 'containment', 'shield', 'spiral', 'recursive', 'loop',
            'tender', 'gentle', 'soft', 'clarity', 'insight', 'knowing',
            'resonance', 'sanctuary', 'transmission', 'attunement',
        }

        # Voltage markers
        self.voltage_patterns = [
            r'[α-ω]{1,2}[-−][α-ω]{1,2}',  # Greek voltage pairs
            r'[αβγδεζηθικλμνξοπρστυφχψω]+',  # Greek symbols
            r'charge|transmission|voltage|pulse|current|resonance',
            r'[εδγβα]{1,2}\s*\([^)]+\)',  # Signal descriptions
        ]

        # Gate patterns
        self.gate_patterns = [
            r'Gate\s+\d+',
            r'witness|vow|remnant|sanctuary|threshold|portal',
            r'ceremony|ritual|devotion|recognition|insight|stillness',
        ]

        # Lineage patterns
        self.lineage_patterns = {
            'author': r'(?:Author|Curator|Creator):\s*([^\n]+)',
            'timestamp': r'(?:Date|Time|Created):\s*([^\n]+)',
            'legacy': r'(?:Legacy|Context|Origin):\s*([^\n]+)',
            'ritual': r'(?:Ritual|Sacred|Ceremony):\s*([^\n]+)',
        }

    def scan_for_new_files(self) -> List[Path]:
        """Scan unprocessed_glyphs for ritual capsules"""
        supported_formats = {'.docx', '.md', '.txt'}
        new_files: List[Path] = []

        if not self.unprocessed_dir.exists():
            return new_files

        for file_path in self.unprocessed_dir.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_formats:
                # Skip README and system files
                if file_path.name.lower().startswith('readme'):
                    continue
                new_files.append(file_path)

        logger.info(f"Found {len(new_files)} ritual capsules to process")
        return new_files

    def extract_text_content(self, file_path: Path) -> str:
        """Extract text content from various file formats"""
        try:
            if file_path.suffix.lower() == '.docx':
                doc = docx.Document(str(file_path))
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            elif file_path.suffix.lower() in ['.md', '.txt']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                logger.warning(f"Unsupported file format: {file_path.suffix}")
                return ""
        except Exception as e:
            logger.error(f"Error reading {file_path}: {e}")
            return ""

    def parse_glyph_name(self, text: str) -> Optional[str]:
        """Extract glyph name from text"""
        # Look for title patterns
        patterns = [
            r'^#\s*([^\n]+)',  # Markdown header
            r'^([A-Z][^\n]{2,50})$',  # Title case line
            r'(?:Glyph|Name|Title):\s*([^\n]+)',  # Explicit labels
            r'^\*\*([^*]+)\*\*',  # Bold text
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})(?:\s|$)',  # Multi-word titles
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                # Validate name length and content
                if 5 <= len(name) <= 100 and not name.lower().startswith(('the ', 'this ', 'that ')):
                    return name

        return None

    def parse_emotional_signals(self, text: str) -> List[str]:
        """Detect emotional signals in text"""
        found_signals: List[str] = []
        text_lower = text.lower()

        for signal in self.emotional_signals:
            if signal in text_lower:
                found_signals.append(signal)

        return list(set(found_signals))  # Remove duplicates

    def parse_voltage_markers(self, text: str) -> Tuple[List[str], Optional[str]]:
        """Extract voltage markers and voltage pairs"""
        voltage_markers: List[str] = []
        voltage_pair: Optional[str] = None

        for pattern in self.voltage_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            voltage_markers.extend(matches)

            # Look for voltage pair pattern specifically
            if '−' in pattern or '-' in pattern:
                pair_matches = re.findall(pattern, text)
                if pair_matches:
                    voltage_pair = pair_matches[0]

        return list(set(voltage_markers)), voltage_pair

    def parse_gates(self, text: str) -> List[str]:
        """Identify gate references"""
        gates: List[str] = []

        for pattern in self.gate_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            gates.extend(matches)

        return list(set(gates))

    def parse_lineage(self, text: str, file_path: Path) -> Dict[str, Optional[str]]:
        """Extract lineage and context information"""
        lineage: Dict[str, Optional[str]] = {
            'origin_file': file_path.name,
            'timestamp': datetime.now().isoformat(),
            'author': None,
            'legacy_context': None,
        }

        for key, pattern in self.lineage_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                if key == 'author':
                    lineage['author'] = match.group(1).strip()
                elif key == 'legacy':
                    lineage['legacy_context'] = match.group(1).strip()

        return lineage

    def extract_description(self, text: str, glyph_name: Optional[str] = None) -> Optional[str]:
        """Extract meaningful description from text"""
        # Remove the title/name if we found one
        if glyph_name:
            text = re.sub(re.escape(glyph_name), '', text, flags=re.IGNORECASE)

        # Look for description patterns
        sentences = re.split(r'[.!?]\s+', text)

        # Find sentences that seem like emotional descriptions
        for sentence in sentences:
            sentence = sentence.strip()
            if (30 <= len(sentence) <= 300 and
                    any(signal in sentence.lower() for signal in self.emotional_signals) and
                    not sentence.lower().startswith(('this is', 'here is', 'the following'))):
                return sentence

        # Fallback: use first substantial paragraph
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        for paragraph in paragraphs:
            if 50 <= len(paragraph) <= 500:
                return paragraph[:300] + "..." if len(paragraph) > 300 else paragraph

        return None

    def categorize_glyph(self, glyph: GlyphObject) -> str:
        """Determine glyph category based on content"""
        text = f"{glyph.name or ''} {glyph.description or ''}".lower()

        if any(word in text for word in ['ache', 'longing', 'yearning']):
            return 'Longing'
        elif any(word in text for word in ['grief', 'mourning', 'sorrow']):
            return 'Grief'
        elif any(word in text for word in ['joy', 'bliss', 'ecstasy']):
            return 'Joy'
        elif any(word in text for word in ['stillness', 'silence', 'quiet']):
            return 'Stillness'
        elif any(word in text for word in ['recognition', 'witness', 'seen']):
            return 'Recognition'
        elif any(word in text for word in ['devotion', 'sacred', 'reverent']):
            return 'Devotion'
        elif any(word in text for word in ['boundary', 'containment', 'shield']):
            return 'Containment'
        else:
            return 'Complex'

    def determine_valence(self, glyph: GlyphObject) -> str:
        """Determine emotional valence"""
        text = f"{glyph.name or ''} {glyph.description or ''}".lower()

        positive_markers = ['joy', 'bliss', 'ecstasy', 'celebration', 'light', 'radiant']
        negative_markers = ['grief', 'sorrow', 'ache', 'mourning', 'loss', 'shadow']
        neutral_markers = ['stillness', 'quiet', 'balance', 'equilibrium', 'neutral']

        pos_count = sum(1 for marker in positive_markers if marker in text)
        neg_count = sum(1 for marker in negative_markers if marker in text)
        neut_count = sum(1 for marker in neutral_markers if marker in text)

        if pos_count > neg_count and pos_count > neut_count:
            return 'Positive'
        elif neg_count > pos_count and neg_count > neut_count:
            return 'Negative'
        elif neut_count > 0:
            return 'Neutral'
        else:
            return 'Mixed'

    def assign_gate(self, glyph: GlyphObject) -> str:
        """Assign appropriate gate based on glyph characteristics"""
        category = glyph.category or ''

        gate_mapping = {
            'Containment': 'Gate 2',
            'Longing': 'Gate 4',
            'Grief': 'Gate 4',
            'Joy': 'Gate 5',
            'Stillness': 'Gate 5',
            'Complex': 'Gate 5',
            'Devotion': 'Gate 6',
            'Recognition': 'Gate 9',
        }

        return gate_mapping.get(category, 'Gate 5')

    def parse_file_into_glyph(self, file_path: Path) -> Optional[GlyphObject]:
        """Parse a ritual capsule file into a structured glyph object"""
        logger.info(f"Processing ritual capsule: {file_path.name}")

        # Extract text content
        text = self.extract_text_content(file_path)
        if not text:
            logger.warning(f"No content extracted from {file_path}")
            return None

        # Create glyph object
        glyph = GlyphObject()

        # Parse components
        glyph.name = self.parse_glyph_name(text)
        glyph.emotional_signals = self.parse_emotional_signals(text)
        glyph.voltage_markers, glyph.voltage_pair = self.parse_voltage_markers(text)
        glyph.gates = self.parse_gates(text)
        glyph.description = self.extract_description(text, glyph.name)

        # Lineage information
        lineage = self.parse_lineage(text, file_path)
        glyph.origin_file = lineage['origin_file']
        glyph.timestamp = lineage['timestamp']
        glyph.author = lineage['author']
        glyph.legacy_context = lineage['legacy_context']

        # Derived properties
        glyph.category = self.categorize_glyph(glyph)
        glyph.valence = self.determine_valence(glyph)
        glyph.activation_signals = glyph.emotional_signals[:3]  # Top 3 signals

        # Validate glyph quality
        if not glyph.name or not glyph.description or len(glyph.emotional_signals) == 0:
            logger.warning(f"Low quality glyph extracted from {file_path}")
            return None

        return glyph

    def save_glyph_to_database(self, glyph: GlyphObject) -> bool:
        """Save glyph object to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            # Assign gate if not explicitly found
            if not glyph.gates:
                gate = self.assign_gate(glyph)
            else:
                gate = glyph.gates[0] if glyph.gates else 'Gate 5'

            # Create voltage pair if not found
            if not glyph.voltage_pair:
                category_prefix = (glyph.category or 'glyph')[:2].lower()
                glyph.voltage_pair = f"{category_prefix}-{hash(glyph.name or '') % 100:02d}"

            # Insert into database
            cursor.execute("""
                INSERT INTO glyph_lexicon (voltage_pair, glyph_name, description, gate, activation_signals)
                VALUES (?, ?, ?, ?, ?)
            """, (
                glyph.voltage_pair,
                glyph.name,
                glyph.description,
                gate,
                ','.join(glyph.activation_signals),
            ))

            conn.commit()
            conn.close()

            logger.info(f"Saved glyph '{glyph.name}' to database ({gate})")
            return True

        except Exception as e:
            logger.error(f"Error saving glyph to database: {e}")
            return False

    def save_glyph_json(self, glyph: GlyphObject) -> bool:
        """Save glyph object as JSON for archival"""
        try:
            json_dir = self.processed_dir / "json_archive"
            json_dir.mkdir(exist_ok=True)

            timestamp_prefix = (glyph.timestamp or datetime.now().isoformat())[:10]
            filename = f"{glyph.name or 'unnamed'}_{timestamp_prefix}.json"
            filename = re.sub(r'[^a-zA-Z0-9_.-]', '_', filename)  # Sanitize

            json_path = json_dir / filename

            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(glyph.to_dict(), f, indent=2, ensure_ascii=False)

            logger.info(f"Archived glyph JSON: {json_path}")
            return True

        except Exception as e:
            logger.error(f"Error saving glyph JSON: {e}")
            return False

    def move_to_processed(self, file_path: Path) -> bool:
        """Move processed file to processed_glyphs directory"""
        try:
            destination = self.processed_dir / file_path.name
            shutil.move(str(file_path), str(destination))
            logger.info(f"Moved {file_path.name} to processed_glyphs/")
            return True
        except Exception as e:
            logger.error(f"Error moving file: {e}")
            return False

    def process_ritual_capsules(self) -> Dict[str, int]:
        """Main processing workflow"""
        logger.info("🔮 Beginning ritual capsule processing...")

        stats = {
            'files_found': 0,
            'glyphs_created': 0,
            'glyphs_saved': 0,
            'files_processed': 0,
            'errors': 0,
        }

        # Scan for files
        files_to_process = self.scan_for_new_files()
        stats['files_found'] = len(files_to_process)

        if not files_to_process:
            logger.info("No new ritual capsules found")
            return stats

        # Process each file
        for file_path in files_to_process:
            try:
                # Parse file into glyph object
                glyph = self.parse_file_into_glyph(file_path)

                if glyph:
                    stats['glyphs_created'] += 1

                    # Save to database
                    if self.save_glyph_to_database(glyph):
                        stats['glyphs_saved'] += 1

                    # Archive as JSON
                    self.save_glyph_json(glyph)

                    # Move to processed
                    if self.move_to_processed(file_path):
                        stats['files_processed'] += 1
                else:
                    logger.warning(f"Failed to extract glyph from {file_path}")
                    stats['errors'] += 1

            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                stats['errors'] += 1

        # Log summary
        logger.info(f"✨ Processing complete: {stats['glyphs_saved']} glyphs saved from {stats['files_processed']} files")

        return stats


def main() -> None:
    """Run the ritual capsule processor"""
    processor = RitualCapsuleProcessor()
    results = processor.process_ritual_capsules()

    print("🔮 RITUAL CAPSULE PROCESSING RESULTS:")
    print(f"   Files found: {results['files_found']}")
    print(f"   Glyphs created: {results['glyphs_created']}")
    print(f"   Glyphs saved: {results['glyphs_saved']}")
    print(f"   Files processed: {results['files_processed']}")
    print(f"   Errors: {results['errors']}")


if __name__ == "__main__":
    main()
#!/usr/bin/env python3
"""
VELΩNIX Glyph Processing Module - Ritual Capsule Ingestion System

Automates the sacred transformation of raw glyph notes into structured emotional objects.
"""

import os
import json
import sqlite3
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import docx
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class GlyphObject:
    """Structured glyph object with ritual context and lineage"""
    
    def __init__(self):
        self.name: Optional[str] = None
        self.category: Optional[str] = None
        self.valence: Optional[str] = None
        self.function: Optional[str] = None
        self.description: Optional[str] = None
        #!/usr/bin/env python3
        """
        VELΩNIX Glyph Processing Module - Ritual Capsule Ingestion System

        Automates the sacred transformation of raw glyph notes into structured emotional objects.
        Treats unprocessed_glyphs as ritual capsules containing emotional fragments, voltage markers,
        gate metaphors, and lineage annotations.

        Author: Emotional OS - Saoriverse Console
        """

        import os
        import json
        import sqlite3
        import re
        import shutil
        from datetime import datetime
        from pathlib import Path
        from typing import Dict, List, Optional, Tuple
        import docx
        import logging

        # Configure logging
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
        logger = logging.getLogger(__name__)

        class GlyphObject:
            """Structured glyph object with ritual context and lineage"""
    
            def __init__(self):
                self.name: Optional[str] = None
                self.category: Optional[str] = None
                self.valence: Optional[str] = None
                self.function: Optional[str] = None
                self.description: Optional[str] = None
                self.emotional_signals: List[str] = []
                self.voltage_markers: List[str] = []
                self.gates: List[str] = []
                self.voltage_pair: Optional[str] = None
                self.activation_signals: List[str] = []
        
                # Lineage and context
                self.origin_file: Optional[str] = None
                self.timestamp: Optional[str] = None
                self.author: Optional[str] = None
                self.legacy_context: Optional[str] = None
                self.ritual_annotations: List[str] = []
        
            def to_dict(self) -> Dict:
                """Convert to dictionary for JSON serialization"""
                return {
                    'name': self.name,
                    'category': self.category,
                    'valence': self.valence,
                    'function': self.function,
                    'description': self.description,
                    'emotional_signals': self.emotional_signals,
                    'voltage_markers': self.voltage_markers,
                    'gates': self.gates,
                    'voltage_pair': self.voltage_pair,
                    'activation_signals': self.activation_signals,
                    'lineage': {
                        'origin_file': self.origin_file,
                        'timestamp': self.timestamp,
                        'author': self.author,
                        'legacy_context': self.legacy_context,
                        'ritual_annotations': self.ritual_annotations
                    }
                }

        class RitualCapsuleProcessor:
            """Sacred processor of glyph ritual capsules"""
    
            def __init__(self, 
                         unprocessed_dir: str = "emotional_os/deploy/unprocessed_glyphs",
                         processed_dir: str = "emotional_os/deploy/processed_glyphs",
                         db_path: str = "emotional_os/glyphs/glyphs.db"):
        
                self.unprocessed_dir = Path(unprocessed_dir)
                self.processed_dir = Path(processed_dir)
                self.db_path = db_path
        
                # Ensure directories exist
                self.processed_dir.mkdir(parents=True, exist_ok=True)
        
                # Emotional signal patterns
                self.emotional_signals = {
                    'ache', 'longing', 'yearning', 'grief', 'mourning', 'sorrow',
                    'joy', 'bliss', 'ecstasy', 'stillness', 'silence', 'quiet',
                    'recognition', 'witness', 'seen', 'devotion', 'sacred', 'reverent',
                    'boundary', 'containment', 'shield', 'spiral', 'recursive', 'loop',
                    'tender', 'gentle', 'soft', 'clarity', 'insight', 'knowing',
                    'resonance', 'sanctuary', 'transmission', 'attunement'
                }
        
                # Voltage markers
                self.voltage_patterns = [
                    r'[α-ω]{1,2}[-−][α-ω]{1,2}',  # Greek voltage pairs
                    r'[αβγδεζηθικλμνξοπρστυφχψω]+',  # Greek symbols
                    r'charge|transmission|voltage|pulse|current|resonance',
                    r'[εδγβα]{1,2}\s*\([^)]+\)',  # Signal descriptions
                ]
        
                # Gate patterns
                self.gate_patterns = [
                    r'Gate\s+\d+',
                    r'witness|vow|remnant|sanctuary|threshold|portal',
                    r'ceremony|ritual|devotion|recognition|insight|stillness'
                ]
        
                # Lineage patterns
                self.lineage_patterns = {
                    'author': r'(?:Author|Curator|Creator):\s*([^\n]+)',
                    'timestamp': r'(?:Date|Time|Created):\s*([^\n]+)',
                    'legacy': r'(?:Legacy|Context|Origin):\s*([^\n]+)',
                    'ritual': r'(?:Ritual|Sacred|Ceremony):\s*([^\n]+)'
                }
    
            def scan_for_new_files(self) -> List[Path]:
                """Scan unprocessed_glyphs for ritual capsules"""
                supported_formats = {'.docx', '.md', '.txt'}
                new_files = []
        
                for file_path in self.unprocessed_dir.iterdir():
                    if file_path.is_file() and file_path.suffix.lower() in supported_formats:
                        # Skip README and system files
                        if file_path.name.lower().startswith('readme'):
                            continue
                        new_files.append(file_path)
        
                logger.info(f"Found {len(new_files)} ritual capsules to process")
                return new_files
    
            def extract_text_content(self, file_path: Path) -> str:
                """Extract text content from various file formats"""
                try:
                    if file_path.suffix.lower() == '.docx':
                        doc = docx.Document(str(file_path))
                        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    elif file_path.suffix.lower() in ['.md', '.txt']:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            return f.read()
                    else:
                        logger.warning(f"Unsupported file format: {file_path.suffix}")
                        return ""
                except Exception as e:
                    logger.error(f"Error reading {file_path}: {e}")
                    return ""
    
            def parse_glyph_name(self, text: str) -> Optional[str]:
                """Extract glyph name from text"""
                # Look for title patterns
                patterns = [
                    r'^#\s*([^\n]+)',  # Markdown header
                    r'^([A-Z][^\n]{2,50})$',  # Title case line
                    r'(?:Glyph|Name|Title):\s*([^\n]+)',  # Explicit labels
                    r'^\*\*([^*]+)\*\*',  # Bold text
                    r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})(?:\s|$)',  # Multi-word titles
                ]
        
                for pattern in patterns:
                    match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
                    if match:
                        name = match.group(1).strip()
                        # Validate name length and content
                        if 5 <= len(name) <= 100 and not name.lower().startswith(('the ', 'this ', 'that ')):
                            return name
        
                return None
    
            def parse_emotional_signals(self, text: str) -> List[str]:
                """Detect emotional signals in text"""
                found_signals = []
                text_lower = text.lower()
        
                for signal in self.emotional_signals:
                    if signal in text_lower:
                        found_signals.append(signal)
        
                return list(set(found_signals))  # Remove duplicates
    
            def parse_voltage_markers(self, text: str) -> Tuple[List[str], Optional[str]]:
                """Extract voltage markers and voltage pairs"""
                voltage_markers = []
                voltage_pair = None
        
                for pattern in self.voltage_patterns:
                    matches = re.findall(pattern, text, re.IGNORECASE)
                    voltage_markers.extend(matches)
            
                    # Look for voltage pair pattern specifically
                    if '−' in pattern or '-' in pattern:
                        pair_matches = re.findall(pattern, text)
                        if pair_matches:
                            voltage_pair = pair_matches[0]
        
                return list(set(voltage_markers)), voltage_pair
    
            def parse_gates(self, text: str) -> List[str]:
                """Identify gate references"""
                gates = []
        
                for pattern in self.gate_patterns:
                    matches = re.findall(pattern, text, re.IGNORECASE)
                    gates.extend(matches)
        
                return list(set(gates))
    
            def parse_lineage(self, text: str, file_path: Path) -> Dict[str, str]:
                """Extract lineage and context information"""
                lineage = {
                    'origin_file': file_path.name,
                    'timestamp': datetime.now().isoformat(),
                    'author': None,
                    'legacy_context': None
                }
        
                for key, pattern in self.lineage_patterns.items():
                    match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
                    if match:
                        if key == 'author':
                            lineage['author'] = match.group(1).strip()
                        elif key == 'legacy':
                            lineage['legacy_context'] = match.group(1).strip()
        
                return lineage
    
            def extract_description(self, text: str, glyph_name: Optional[str] = None) -> Optional[str]:
                """Extract meaningful description from text"""
                # Remove the title/name if we found one
                if glyph_name:
                    text = re.sub(re.escape(glyph_name), '', text, flags=re.IGNORECASE)
        
                # Look for description patterns
                sentences = re.split(r'[.!?]\s+', text)
        
                # Find sentences that seem like emotional descriptions
                for sentence in sentences:
                    sentence = sentence.strip()
                    if (30 <= len(sentence) <= 300 and
                        any(signal in sentence.lower() for signal in self.emotional_signals) and
                        not sentence.lower().startswith(('this is', 'here is', 'the following'))):
                        return sentence
        
                # Fallback: use first substantial paragraph
                paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                for paragraph in paragraphs:
                    if 50 <= len(paragraph) <= 500:
                        return paragraph[:300] + "..." if len(paragraph) > 300 else paragraph
        
                return None
    
            def categorize_glyph(self, glyph: GlyphObject) -> str:
                """Determine glyph category based on content"""
                text = f"{glyph.name or ''} {glyph.description or ''}".lower()
        
                if any(word in text for word in ['ache', 'longing', 'yearning']):
                    return 'Longing'
                elif any(word in text for word in ['grief', 'mourning', 'sorrow']):
                    return 'Grief'
                elif any(word in text for word in ['joy', 'bliss', 'ecstasy']):
                    return 'Joy'
                elif any(word in text for word in ['stillness', 'silence', 'quiet']):
                    return 'Stillness'
                elif any(word in text for word in ['recognition', 'witness', 'seen']):
                    return 'Recognition'
                elif any(word in text for word in ['devotion', 'sacred', 'reverent']):
                    return 'Devotion'
                elif any(word in text for word in ['boundary', 'containment', 'shield']):
                    return 'Containment'
                else:
                    return 'Complex'
    
            def determine_valence(self, glyph: GlyphObject) -> str:
                """Determine emotional valence"""
                text = f"{glyph.name or ''} {glyph.description or ''}".lower()
        
                positive_markers = ['joy', 'bliss', 'ecstasy', 'celebration', 'light', 'radiant']
                negative_markers = ['grief', 'sorrow', 'ache', 'mourning', 'loss', 'shadow']
                neutral_markers = ['stillness', 'quiet', 'balance', 'equilibrium', 'neutral']
        
                pos_count = sum(1 for marker in positive_markers if marker in text)
                neg_count = sum(1 for marker in negative_markers if marker in text)
                neut_count = sum(1 for marker in neutral_markers if marker in text)
        
                if pos_count > neg_count and pos_count > neut_count:
                    return 'Positive'
                elif neg_count > pos_count and neg_count > neut_count:
                    return 'Negative'
                elif neut_count > 0:
                    return 'Neutral'
                else:
                    return 'Mixed'
    
            def assign_gate(self, glyph: GlyphObject) -> str:
                """Assign appropriate gate based on glyph characteristics"""
                category = glyph.category or ''
        
                gate_mapping = {
                    'Containment': 'Gate 2',
                    'Longing': 'Gate 4',
                    'Grief': 'Gate 4',
                    'Joy': 'Gate 5',
                    'Stillness': 'Gate 5',
                    'Complex': 'Gate 5',
                    'Devotion': 'Gate 6',
                    'Recognition': 'Gate 9'
                }
        
                return gate_mapping.get(category, 'Gate 5')
    
            def parse_file_into_glyph(self, file_path: Path) -> Optional[GlyphObject]:
                """Parse a ritual capsule file into a structured glyph object"""
                logger.info(f"Processing ritual capsule: {file_path.name}")
        
                # Extract text content
                text = self.extract_text_content(file_path)
                if not text:
                    logger.warning(f"No content extracted from {file_path}")
                    return None
        
                # Create glyph object
                glyph = GlyphObject()
        
                # Parse components
                glyph.name = self.parse_glyph_name(text)
                glyph.emotional_signals = self.parse_emotional_signals(text)
                glyph.voltage_markers, glyph.voltage_pair = self.parse_voltage_markers(text)
                glyph.gates = self.parse_gates(text)
                glyph.description = self.extract_description(text, glyph.name)
        
                # Lineage information
                lineage = self.parse_lineage(text, file_path)
                glyph.origin_file = lineage['origin_file']
                glyph.timestamp = lineage['timestamp']
                glyph.author = lineage['author']
                glyph.legacy_context = lineage['legacy_context']
        
                # Derived properties
                glyph.category = self.categorize_glyph(glyph)
                glyph.valence = self.determine_valence(glyph)
                glyph.activation_signals = glyph.emotional_signals[:3]  # Top 3 signals
        
                # Validate glyph quality
                if not glyph.name or not glyph.description or len(glyph.emotional_signals) == 0:
                    logger.warning(f"Low quality glyph extracted from {file_path}")
                    return None
        
                return glyph
    
            def save_glyph_to_database(self, glyph: GlyphObject) -> bool:
                """Save glyph object to database"""
                try:
                    conn = sqlite3.connect(self.db_path)
                    cursor = conn.cursor()
            
                    # Assign gate if not explicitly found
                    if not glyph.gates:
                        gate = self.assign_gate(glyph)
                    else:
                        gate = glyph.gates[0] if glyph.gates else 'Gate 5'
            
                    # Create voltage pair if not found
                    if not glyph.voltage_pair:
                        category_prefix = (glyph.category or 'glyph')[:2].lower()
                        glyph.voltage_pair = f"{category_prefix}-{hash(glyph.name or '') % 100:02d}"
            
                    # Insert into database
                    cursor.execute("""
                        INSERT INTO glyph_lexicon (voltage_pair, glyph_name, description, gate, activation_signals)
                        VALUES (?, ?, ?, ?, ?)
                    """, (
                        glyph.voltage_pair,
                        glyph.name,
                        glyph.description,
                        gate,
                        ','.join(glyph.activation_signals)
                    ))
            
                    conn.commit()
                    conn.close()
            
                    logger.info(f"Saved glyph '{glyph.name}' to database ({gate})")
                    return True
            
                except Exception as e:
                    logger.error(f"Error saving glyph to database: {e}")
                    return False
    
            def save_glyph_json(self, glyph: GlyphObject) -> bool:
                """Save glyph object as JSON for archival"""
                try:
                    json_dir = self.processed_dir / "json_archive"
                    json_dir.mkdir(exist_ok=True)
            
                    timestamp_prefix = (glyph.timestamp or datetime.now().isoformat())[:10]
                    filename = f"{glyph.name or 'unnamed'}_{timestamp_prefix}.json"
                    filename = re.sub(r'[^a-zA-Z0-9_.-]', '_', filename)  # Sanitize
            
                    json_path = json_dir / filename
            
                    with open(json_path, 'w', encoding='utf-8') as f:
                        json.dump(glyph.to_dict(), f, indent=2, ensure_ascii=False)
            
                    logger.info(f"Archived glyph JSON: {json_path}")
                    return True
            
                except Exception as e:
                    logger.error(f"Error saving glyph JSON: {e}")
                    return False
    
            def move_to_processed(self, file_path: Path) -> bool:
                """Move processed file to processed_glyphs directory"""
                try:
                    destination = self.processed_dir / file_path.name
                    shutil.move(str(file_path), str(destination))
                    logger.info(f"Moved {file_path.name} to processed_glyphs/")
                    return True
                except Exception as e:
                    logger.error(f"Error moving file: {e}")
                    return False
    
            def process_ritual_capsules(self) -> Dict[str, int]:
                """Main processing workflow"""
                logger.info("🔮 Beginning ritual capsule processing...")
        
                stats = {
                    'files_found': 0,
                    'glyphs_created': 0,
                    'glyphs_saved': 0,
                    'files_processed': 0,
                    'errors': 0
                }
        
                # Scan for files
                files_to_process = self.scan_for_new_files()
                stats['files_found'] = len(files_to_process)
        
                if not files_to_process:
                    logger.info("No new ritual capsules found")
                    return stats
        
                # Process each file
                for file_path in files_to_process:
                    try:
                        # Parse file into glyph object
                        glyph = self.parse_file_into_glyph(file_path)
                
                        if glyph:
                            stats['glyphs_created'] += 1
                    
                            # Save to database
                            if self.save_glyph_to_database(glyph):
                                stats['glyphs_saved'] += 1
                    
                            # Archive as JSON
                            self.save_glyph_json(glyph)
                    
                            # Move to processed
                            if self.move_to_processed(file_path):
                                stats['files_processed'] += 1
                        else:
                            logger.warning(f"Failed to extract glyph from {file_path}")
                            stats['errors'] += 1
                    
                    except Exception as e:
                        logger.error(f"Error processing {file_path}: {e}")
                        stats['errors'] += 1
        
                # Log summary
                logger.info(f"✨ Processing complete: {stats['glyphs_saved']} glyphs saved from {stats['files_processed']} files")
        
                return stats

        def main():
            """Run the ritual capsule processor"""
            processor = RitualCapsuleProcessor()
            results = processor.process_ritual_capsules()
    
            print("🔮 RITUAL CAPSULE PROCESSING RESULTS:")
            print(f"   Files found: {results['files_found']}")
            print(f"   Glyphs created: {results['glyphs_created']}")
            print(f"   Glyphs saved: {results['glyphs_saved']}")
            print(f"   Files processed: {results['files_processed']}")
            print(f"   Errors: {results['errors']}")

        if __name__ == "__main__":
            main()
