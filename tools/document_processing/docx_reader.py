#!/usr/bin/env python3
"""
DOCX Reader Utility
Provides functions to read and extract content from DOCX files.
"""

import json
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.shared import Inches, Pt


def read_docx(filepath: str) -> Dict[str, Any]:
    """
    Read a DOCX file and extract all content.

    Args:
        filepath: Path to the DOCX file

    Returns:
        Dictionary containing paragraphs, tables, and metadata
    """
    try:
        doc = Document(filepath)

        result = {
            "filepath": str(filepath),
            "filename": Path(filepath).name,
            "paragraphs": [],
            "tables": [],
            "sections": [],
            "core_properties": {
                "title": doc.core_properties.title or "N/A",
                "author": doc.core_properties.author or "N/A",
                "subject": doc.core_properties.subject or "N/A",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "N/A",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "N/A",
            },
        }

        # Extract paragraphs
        for i, para in enumerate(doc.paragraphs):
            result["paragraphs"].append(
                {
                    "index": i,
                    "text": para.text,
                    "level": para.style.name if para.style else "Normal",
                    "style": para.style.name if para.style else "Normal",
                    "bold": any(run.bold for run in para.runs) if para.runs else False,
                    "italic": any(run.italic for run in para.runs) if para.runs else False,
                }
            )

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "index": table_idx,
                "rows": [],
                "row_count": len(table.rows),
                "col_count": len(table.columns),
            }

            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    row_data.append(
                        {
                            "content": cell.text,
                            "col_index": cell_idx,
                        }
                    )
                table_data["rows"].append(row_data)

            result["tables"].append(table_data)

        return result

    except Exception as e:
        return {
            "error": str(e),
            "filepath": str(filepath),
        }


def print_docx_content(filepath: str) -> None:
    """
    Print DOCX content in a readable format to stdout.

    Args:
        filepath: Path to the DOCX file
    """
    data = read_docx(filepath)

    if "error" in data:
        print(f"Error reading file: {data['error']}")
        return

    print(f"\n{'='*60}")
    print(f"Document: {data['filename']}")
    print(f"{'='*60}\n")

    # Print metadata
    print("📋 DOCUMENT PROPERTIES")
    print("-" * 40)
    for key, value in data["core_properties"].items():
        print(f"  {key.title()}: {value}")

    # Print paragraphs
    if data["paragraphs"]:
        print(f"\n📝 CONTENT ({len(data['paragraphs'])} paragraphs)")
        print("-" * 40)
        for para in data["paragraphs"]:
            if para["text"].strip():  # Skip empty paragraphs
                indent = "  " if para["level"] != "Normal" else ""
                marker = "📌 " if para["bold"] else "   "
                print(f"{marker}{indent}{para['text']}")

    # Print tables
    if data["tables"]:
        print(f"\n📊 TABLES ({len(data['tables'])} table(s))")
        print("-" * 40)
        for table in data["tables"]:
            print(f"\nTable {table['index'] + 1} ({table['row_count']}×{table['col_count']}):")
            for row in table["rows"]:
                row_text = " | ".join([cell["content"] for cell in row])
                print(f"  {row_text}")

    print(f"\n{'='*60}\n")


def docx_to_text(filepath: str) -> str:
    """
    Extract plain text from DOCX file.

    Args:
        filepath: Path to the DOCX file

    Returns:
        Plain text content
    """
    data = read_docx(filepath)

    if "error" in data:
        return f"Error reading file: {data['error']}"

    text_parts = []

    # Add paragraphs
    for para in data["paragraphs"]:
        if para["text"].strip():
            text_parts.append(para["text"])

    # Add tables
    for table in data["tables"]:
        text_parts.append("\n[TABLE]")
        for row in table["rows"]:
            row_text = " | ".join([cell["content"] for cell in row])
            text_parts.append(row_text)
        text_parts.append("[/TABLE]\n")

    return "\n".join(text_parts)


def export_docx_json(filepath: str, output_filepath: str = None) -> str:
    """
    Export DOCX content as JSON.

    Args:
        filepath: Path to the DOCX file
        output_filepath: Optional output JSON file path

    Returns:
        JSON string of DOCX content
    """
    data = read_docx(filepath)

    json_str = json.dumps(data, indent=2)

    if output_filepath:
        with open(output_filepath, "w") as f:
            f.write(json_str)
        print(f"✅ Exported to {output_filepath}")

    return json_str


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python3 docx_reader.py <docx_file> [command]")
        print("\nCommands:")
        print("  view    - Display formatted content (default)")
        print("  text    - Extract plain text only")
        print("  json    - Export as JSON")
        sys.exit(1)

    filepath = sys.argv[1]
    command = sys.argv[2] if len(sys.argv) > 2 else "view"

    if not Path(filepath).exists():
        print(f"❌ File not found: {filepath}")
        sys.exit(1)

    if command == "view":
        print_docx_content(filepath)
    elif command == "text":
        print(docx_to_text(filepath))
    elif command == "json":
        output = sys.argv[3] if len(sys.argv) > 3 else filepath.replace(".docx", ".json")
        export_docx_json(filepath, output)
    else:
        print(f"Unknown command: {command}")
        sys.exit(1)
