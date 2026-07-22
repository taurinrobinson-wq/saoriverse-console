from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

file_path = 'd:/saoriverse-console/Darwin/docs/Cho,KyungMobilitasRespRFA.docx'

print(f"Analyzing: {file_path}\n")

doc = Document(file_path)

# Analyze document structure
print("=" * 100)
print("DISCOVERY DOCUMENT TEMPLATE ANALYSIS - CHO v. MOBILITAS")
print("=" * 100)

# Check sections and margins
print("\n### SECTIONS & MARGINS ###")
for i, section in enumerate(doc.sections):
    print(f"Section {i}:")
    print(f"  Top margin: {section.top_margin.inches:.2f}\" ({section.top_margin.twips} twips)")
    print(f"  Bottom margin: {section.bottom_margin.inches:.2f}\" ({section.bottom_margin.twips} twips)")
    print(f"  Left margin: {section.left_margin.inches:.2f}\" ({section.left_margin.twips} twips)")
    print(f"  Right margin: {section.right_margin.inches:.2f}\" ({section.right_margin.twips} twips)")

# Analyze paragraphs
print("\n### PARAGRAPHS ###")
print(f"Total paragraphs: {len(doc.paragraphs)}\n")

for i, para in enumerate(doc.paragraphs[:45]):  # First 45 paragraphs to get complete picture
    text = para.text.strip()
    if text or i < 15:  # Show empty paragraphs at start
        print(f"\nPara {i}: {text[:70]}")
        print(f"  Alignment: {para.alignment}")
        
        pf = para.paragraph_format
        print(f"  Line spacing: {pf.line_spacing}")
        print(f"  Space before: {pf.space_before} | Space after: {pf.space_after}")
        print(f"  First line indent: {pf.first_line_indent}")
        print(f"  Left indent: {pf.left_indent} | Right indent: {pf.right_indent}")
        
        if para.runs and text:
            for run in para.runs:
                if run.text.strip():
                    print(f"  Run: {run.text[:50]}")
                    print(f"    Font: {run.font.name} | Size: {run.font.size} | Bold: {run.font.bold} | Underline: {run.font.underline}")

# Analyze tables
print(f"\n\n### TABLES ###")
print(f"Total tables: {len(doc.tables)}\n")

for i, table in enumerate(doc.tables):
    print(f"\nTable {i}:")
    print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")
    print(f"  Alignment: {table.alignment}")
    if table.rows:
        for row_idx, row in enumerate(table.rows[:8]):
            print(f"\n  Row {row_idx}:")
            for j, cell in enumerate(row.cells):
                cell_text = cell.text.strip()[:60]
                print(f"    Cell {j}: {cell_text}")
                if cell.paragraphs and cell.text.strip():
                    para = cell.paragraphs[0]
                    if para.runs:
                        run = para.runs[0]
                        print(f"      Font: {run.font.name} | Size: {run.font.size} | Bold: {run.font.bold}")

print("\n" + "=" * 100)
