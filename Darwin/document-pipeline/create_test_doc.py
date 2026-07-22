from docx import Document
from docx.shared import Pt

# Create document
doc = Document()

# Add some content
doc.add_paragraph('Interrogatory No. 1').runs[0].font.size = Pt(12)
doc.add_paragraph('Please identify the date when the contract was signed.')

doc.add_paragraph('Interrogatory No. 2')
doc.add_paragraph('Describe in detail all communications regarding the matter.')

doc.add_paragraph('Request for Production No. 1')
doc.add_paragraph('Produce all emails related to this transaction.')

doc.add_paragraph('Topic 1: Background Facts')
doc.add_paragraph('Provide background information about the incident.')

# Save
doc.save('test-document.docx')
print('Test document created: test-document.docx')
