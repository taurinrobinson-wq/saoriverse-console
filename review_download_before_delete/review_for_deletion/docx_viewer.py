#!/usr/bin/env python3
"""
DOCX Viewer - Streamlit App
Interactive viewer for DOCX files with formatting preservation.
"""

import streamlit as st
from pathlib import Path
from docx import Document
import json


def get_docx_files(directory="."):
    """Get all DOCX files in the current directory."""
    return sorted([f for f in Path(directory).glob("*.docx")])


def extract_docx_content(filepath):
    """Extract all content from DOCX file."""
    doc = Document(filepath)
    
    content = {
        'metadata': {
            'title': doc.core_properties.title or 'N/A',
            'author': doc.core_properties.author or 'N/A',
            'subject': doc.core_properties.subject or 'N/A',
            'created': str(doc.core_properties.created) if doc.core_properties.created else 'N/A',
            'modified': str(doc.core_properties.modified) if doc.core_properties.modified else 'N/A',
        },
        'paragraphs': [],
        'tables': [],
    }
    
    # Extract paragraphs with styling
    for para in doc.paragraphs:
        if para.text.strip():
            para_info = {
                'text': para.text,
                'style': para.style.name if para.style else 'Normal',
                'alignment': str(para.alignment),
            }
            content['paragraphs'].append(para_info)
    
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content['tables'].append(table_data)
    
    return content


def render_paragraph(text, style):
    """Render paragraph with appropriate Streamlit formatting."""
    if 'Heading' in style:
        level = 1
        if 'Heading 2' in style:
            level = 2
        elif 'Heading 3' in style:
            level = 3
        
        if level == 1:
            st.header(text)
        elif level == 2:
            st.subheader(text)
        else:
            st.markdown(f"**{text}**")
    else:
        st.write(text)


def main():
    st.set_page_config(page_title="DOCX Viewer", layout="wide")
    
    st.title("📄 DOCX Viewer")
    st.markdown("View and explore Word documents in your workspace")
    
    # Sidebar options
    with st.sidebar:
        st.header("Options")
        view_mode = st.radio(
            "View Mode:",
            ["Document View", "Raw Content", "Metadata", "Export"]
        )
        
        st.divider()
        
        # File selection
        docx_files = get_docx_files()
        if not docx_files:
            st.warning("No DOCX files found in the current directory")
            return
        
        selected_file = st.selectbox(
            "Select a DOCX file:",
            docx_files,
            format_func=lambda x: x.name
        )
    
    if not selected_file:
        return
    
    # Load document
    try:
        content = extract_docx_content(selected_file)
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return
    
    # Display header
    st.markdown(f"## {selected_file.name}")
    
    if view_mode == "Document View":
        # Display document as it appears
        col1, col2 = st.columns([3, 1])
        
        with col2:
            st.subheader("Document Info")
            for key, value in content['metadata'].items():
                st.caption(f"**{key}:** {value}")
        
        with col1:
            st.subheader("Content")
            
            # Paragraphs
            if content['paragraphs']:
                for para in content['paragraphs']:
                    render_paragraph(para['text'], para['style'])
            
            # Tables
            if content['tables']:
                st.subheader("Tables")
                for i, table in enumerate(content['tables']):
                    with st.expander(f"Table {i + 1}"):
                        st.dataframe(table, use_container_width=True)
    
    elif view_mode == "Raw Content":
        st.subheader("Raw Document Structure")
        
        with st.expander("Metadata", expanded=True):
            st.json(content['metadata'])
        
        with st.expander(f"Paragraphs ({len(content['paragraphs'])})"):
            st.json(content['paragraphs'])
        
        with st.expander(f"Tables ({len(content['tables'])})"):
            for i, table in enumerate(content['tables']):
                st.write(f"**Table {i + 1}:**")
                st.dataframe(table)
    
    elif view_mode == "Metadata":
        st.subheader("Document Metadata")
        
        col1, col2 = st.columns(2)
        with col1:
            st.write("**Title:**", content['metadata']['title'])
            st.write("**Author:**", content['metadata']['author'])
            st.write("**Subject:**", content['metadata']['subject'])
        
        with col2:
            st.write("**Created:**", content['metadata']['created'])
            st.write("**Modified:**", content['metadata']['modified'])
        
        st.divider()
        st.write(f"**Paragraphs:** {len(content['paragraphs'])}")
        st.write(f"**Tables:** {len(content['tables'])}")
    
    elif view_mode == "Export":
        st.subheader("Export Options")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📋 Copy as JSON"):
                json_str = json.dumps(content, indent=2)
                st.code(json_str, language="json")
                st.success("JSON copied to clipboard!")
        
        with col2:
            json_str = json.dumps(content, indent=2)
            st.download_button(
                label="⬇️ Download as JSON",
                data=json_str,
                file_name=f"{selected_file.stem}.json",
                mime="application/json"
            )
        
        st.divider()
        
        # Export as text
        text_content = "\n".join([p['text'] for p in content['paragraphs']])
        st.download_button(
            label="⬇️ Download as TXT",
            data=text_content,
            file_name=f"{selected_file.stem}.txt",
            mime="text/plain"
        )


if __name__ == "__main__":
    main()
