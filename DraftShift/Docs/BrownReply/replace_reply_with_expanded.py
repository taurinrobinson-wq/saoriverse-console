import json
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parent
PINCITES = ROOT / 'pincites.json'
IN_DOCX = ROOT / 'PakBrownReplyMtnNewTrial_with_reply.docx'
OUT_DOCX = ROOT / 'PakBrownReplyMtnNewTrial_rev1.docx'

with open(PINCITES, 'r', encoding='utf-8') as f:
    pincites = json.load(f)

# build expanded reply (Option 1, credibility-focused)
lines = []
lines.append('Reply Section — Credibility and Record (Short, Judge‑Friendly)')
lines.append('')
lines.append('Introduction: Plaintiffs’ opposition asks the Court to reweigh the record in their favor. It cannot. The Final Statement of Decision (SOD) squarely resolves the disputed factual questions and credibly narrows liability to Ms. Pak in limited respects; it rejects broader theories against Mr. Norris and others. (See SOD citations below.)')
lines.append('')
# Credibility
lines.append('I. Credibility — the controlling facts')
bt = pincites.get('brown_testimony_weight')
if bt:
    lines.append(f"- The SOD found Mr. Brown’s testimony of limited weight and relied instead on corroborated evidence where present (SOD p.{bt['page']}:{bt['start_line']}).")
else:
    lines.append('- The SOD found Mr. Brown’s testimony of limited weight.')
pn = pincites.get('pak_not_credible')
if pn:
    lines.append(f"- The court specifically questioned Ms. Pak’s credibility (SOD p.{pn['page']}:{pn['start_line']}).")

lines.append('')
# Liability narrowness
lines.append('II. Liability — what the SOD actually holds')
fp = pincites.get('failed_prove_norris_pak')
if fp:
    lines.append(f"- Plaintiffs failed to prove that Norris or Pak wrongfully took or retained plaintiffs' property; the court found consensual transfers in the relevant period (SOD p.{fp['page']}:{fp['start_line']}).")

p188 = pincites.get('pak_188500')
if p188:
    lines.append(f"- The SOD found Pak breached fiduciary duties with respect to lead purchases totaling $188,500 and some misapplied firm funds (SOD p.{p188['page']}:{p188['start_line']}-{p188['end_line']}).")

l24 = pincites.get('lease_24000')
if l24:
    lines.append(f"- The SOD awarded narrow lease/occupancy damages (SOD p.{l24['page']}:{l24['start_line']}).")

lines.append('')
# Remedies
lines.append('III. Remedies and scope')
af = pincites.get('attorney_fees_15657')
if af:
    lines.append(f"- Because the court found elder financial abuse as to Pak, the SOD contemplates attorney fees under Welfare & Inst. Code §15657.5 (SOD p.{af['page']}:{af['start_line']}).")

lines.append('')
lines.append('Conclusion: The Court should enter judgment consistent with the SOD: deny Plaintiffs any relief beyond the limited remedies the court awarded and preserve only those narrow claims the SOD sustained against Ms. Pak.')

# now open docx, remove previous appended section if present, append expanded reply
if not IN_DOCX.exists():
    doc = Document()
else:
    doc = Document(str(IN_DOCX))

# remove any existing paragraph containing the exact heading used earlier
heading_text = 'Reply Section — Credibility and Record (Short, Judge‑Friendly)'
paras = list(doc.paragraphs)
found_idx = None
for i, p in enumerate(paras):
    if heading_text in p.text:
        found_idx = i
        break

if found_idx is not None:
    # remove from found_idx to end
    # python-docx doesn't support deleting paragraphs directly; recreate doc
    new_doc = Document()
    for p in paras[:found_idx]:
        new_doc._body._element.append(p._p)
    # append expanded content
    new_doc.add_heading(heading_text, level=2)
    for ln in lines[1:]:
        new_doc.add_paragraph(ln)
    new_doc.save(str(OUT_DOCX))
    print('WROTE', OUT_DOCX)
else:
    # just append
    doc.add_heading(heading_text, level=2)
    for ln in lines[1:]:
        doc.add_paragraph(ln)
    doc.save(str(OUT_DOCX))
    print('WROTE', OUT_DOCX)
