import json
import os
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent
SOD_PAGES = ROOT / 'sod_pages.json'
REPLY_TEXT = ROOT / 'ready_reply_with_pincites.txt'
DOCX_PATH = Path(r'D:\saoriverse-console\DraftShift\Docs\BrownReply\PakBrownReplyMtnNewTrial.docx')

# target phrases to find (verbatim where possible)
targets = [
    {
        'key': 'brown_testimony_weight',
        'text': "Accordingly, the court could give little weight to his testimony, except where it was corroborated by other testimony or evidence.",
        'fallback': 'could give little weight to his testimony'
    },
    {
        'key': 'pak_not_credible',
        'text': "The court does not find Pak to be a credible witness",
        'fallback': 'does not find Pak to be a credible'
    },
    {
        'key': 'no_conversion_norris',
        'text': "Plaintiffs did not establish that the Norris defendants took or retained plaintiffs' property.",
        'fallback': 'Plaintiffs did not establish that the Norris'
    },
    {
        'key': 'pak_188500',
        'text': "In sum, the court agrees with plaintiffs that they have established that Pak breached her fiduciary duty to plaintiffs by using $188,500 of money belonging to the firm and/or Brown to purchase leads for the Norris firm.",
        'fallback': 'Pak breached her fiduciary duty to plaintiffs by using $188,500'
    },
    {
        'key': 'lease_24000',
        'text': "are entitled to damages in the amount of $24,000",
        'fallback': 'entitled to damages in the amount of $24,000'
    },
    {
        'key': 'attorney_fees_15657',
        'text': "the court shall award to the plaintiff reasonable attorney's fees and costs.",
        'fallback': "award to the plaintiff reasonable attorney's fees"
    },
    {
        'key': 'failed_prove_norris_pak',
        'text': "Plaintiffs failed to prove that Norris or Pak wrongfully obtained any clients or retained funds received in settlement for those clients which rightfully belonged to plaintiffs.",
        'fallback': 'failed to prove that Norris or Pak'
    }
]

# load sod pages
with open(SOD_PAGES, 'r', encoding='utf-8') as f:
    pages = json.load(f)

# build mapping page->lines
page_lines = {}
for p in pages:
    pg = p['page']
    txt = p['text']
    # normalize newlines
    lines = txt.splitlines()
    page_lines[pg] = lines

results = {}

# helper: find substring within page lines; return (start_line, end_line) 1-based
def find_in_page_lines(lines, substring):
    sub = substring.strip()
    # try exact line match
    for i, line in enumerate(lines):
        if sub in line:
            return i+1, i+1
    # try spanning lines: join up to 5 lines
    N = len(lines)
    for i in range(N):
        accum = lines[i]
        for j in range(i+1, min(i+8, N)):
            accum += '\n' + lines[j]
            if sub in accum:
                return i+1, j+1
    return None

# search
for t in targets:
    found = None
    for pg, lines in page_lines.items():
        res = find_in_page_lines(lines, t['text'])
        if res:
            found = {'page': pg, 'start_line': res[0], 'end_line': res[1], 'matched_text': t['text']}
            break
    if not found:
        # try fallback shorter search
        for pg, lines in page_lines.items():
            res = find_in_page_lines(lines, t['fallback'])
            if res:
                # capture the actual text from start to end
                matched = '\n'.join(lines[res[0]-1:res[1]])
                found = {'page': pg, 'start_line': res[0], 'end_line': res[1], 'matched_text': matched}
                break
    if not found:
        results[t['key']] = None
    else:
        results[t['key']] = found

# write pincites JSON
with open(ROOT / 'pincites.json', 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)

# build reply text with pincites inserted at quoted spots
reply = []
reply.append('Reply Section — Credibility and Record (Short, Judge‑Friendly)')
reply.append('')
reply.append('- Introduction / Issue Presented: Plaintiffs’ opposition re‑frames the record by relying on Mr. Brown’s uncorroborated testimony. The SOD confirms the relevant reality: the court found Mr. Brown’s testimony of limited weight due to pervasive cognitive impairment and contradictions. ')
# insert brown_testimony_weight pincite
bt = results.get('brown_testimony_weight')
if bt:
    reply[-1] += f" (SOD p.{bt['page']}:{bt['start_line']})"

reply.append('')
reply.append('- Court’s Core Credibility Findings: The court specifically rejected Pak’s and Plaintiffs’ versions where they relied on inconsistent witness statements:')
# pak_not_credible
pn = results.get('pak_not_credible')
if pn:
    reply.append(f"  - On Pak’s credibility: \"The court does not find Pak to be a credible witness\" (SOD p.{pn['page']}:{pn['start_line']}).")
else:
    reply.append('  - On Pak’s credibility: (SOD citation not found)')
# brown testimony again
if bt:
    reply.append(f"  - On Brown’s memory and weight of his testimony: \"{targets[0]['text']}\" (SOD p.{bt['page']}:{bt['start_line']}).")

reply.append('')
reply.append('- Controlling Record on Liability and Remedies:')
# no_conversion_norris
nc = results.get('no_conversion_norris')
if nc:
    reply.append(f"  - Theft/Conversion & Norris: {targets[2]['text']} (SOD p.{nc['page']}:{nc['start_line']}).")
else:
    reply.append('  - Theft/Conversion & Norris: (SOD citation not found)')
# pak_188500
p188 = results.get('pak_188500')
if p188:
    reply.append(f"  - Pak’s fiduciary breaches and elder‑abuse findings: {p188['matched_text']} (SOD p.{p188['page']}:{p188['start_line']}-{p188['end_line']}).")
else:
    reply.append('  - Pak’s fiduciary breaches and elder‑abuse findings: (SOD citation not found)')
# lease_24000
l24 = results.get('lease_24000')
if l24:
    reply.append(f"  - Lease/balloon damages: {l24['matched_text']} (SOD p.{l24['page']}:{l24['start_line']}).")

reply.append('')
reply.append('- Narrow Relief; Why Plaintiffs’ Overbroad Theories Fail:')
# failed_prove_norris_pak
fp = results.get('failed_prove_norris_pak')
if fp:
    reply.append(f"  - The SOD rejects broad theories of conversion, unfair‑competition, and treble damages against Norris because the court found consent/assent and insufficient proof: {fp['matched_text']} (SOD p.{fp['page']}:{fp['start_line']}).")

reply.append('')
reply.append('Closing sentence for filing slot (copy‑ready):')
reply.append('For the reasons stated and as reflected in the court’s Final Statement of Decision, Plaintiffs’ broad theories of theft and conversion are unsupported by the record and the court’s findings; the only sustainable claims are the narrow fiduciary/elder‑abuse breaches against Ms. Pak and the limited lease remedy awarded to Plaintiffs. We therefore respectfully request entry of judgment consistent with those findings and denial of any relief beyond what the SOD awarded.')

# write the text file
with open(REPLY_TEXT, 'w', encoding='utf-8') as f:
    f.write('\n'.join(reply))

print('WROTE', REPLY_TEXT)
print('WROTE pincites.json')

# Now append to docx if python-docx available
try:
    from docx import Document
    doc = None
    if DOCX_PATH.exists():
        doc = Document(str(DOCX_PATH))
    else:
        doc = Document()
    # append heading and text
    doc.add_heading('Reply Section — Credibility and Record (Short, Judge‑Friendly)', level=2)
    for para in reply[1:]:
        doc.add_paragraph(para)
    new_path = DOCX_PATH.with_name(DOCX_PATH.stem + '_with_reply.docx')
    doc.save(str(new_path))
    print('WROTE', new_path)
except Exception as e:
    print('Skipped docx insertion:', e)
