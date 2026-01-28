from docx import Document
from pathlib import Path

ROOT = Path(__file__).resolve().parent
POLISHED = ROOT / 'PakBrownReplyMtnNewTrial_Polished_Option1.txt'
IN_DOCX = ROOT / 'PakBrownReplyMtnNewTrial_rev1.docx'
OUT_DOCX = ROOT / 'PakBrownReplyMtnNewTrial_final.docx'

# read polished text
with open(POLISHED, 'r', encoding='utf-8') as f:
    text = f.read().splitlines()

# load doc
if IN_DOCX.exists():
    doc = Document(str(IN_DOCX))
else:
    doc = Document()

# remove any existing heading we used previously
heading = 'Reply Section — Credibility and Record (Short, Judge‑Friendly)'
found = None
for i,p in enumerate(doc.paragraphs):
    if heading in p.text:
        found = i
        break

if found is not None:
    # rebuild document up to found
    new = Document()
    for p in doc.paragraphs[:found]:
        new._body._element.append(p._p)
    new.add_heading('Reply Section — Credibility and Record (Polished Option 1)', level=2)
    for ln in text:
        new.add_paragraph(ln)
    new.save(str(OUT_DOCX))
    print('WROTE', OUT_DOCX)
else:
    doc.add_heading('Reply Section — Credibility and Record (Polished Option 1)', level=2)
    for ln in text:
        doc.add_paragraph(ln)
    doc.save(str(OUT_DOCX))
    print('WROTE', OUT_DOCX)
