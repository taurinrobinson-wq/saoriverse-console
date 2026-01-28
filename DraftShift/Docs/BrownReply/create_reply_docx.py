from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_paragraph()
title_run = title.add_run('REPLY TO PLAINTIFFS\' OPPOSITION TO MOTION FOR NEW TRIAL')
title_run.font.size = Pt(12)
title_run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add spacing
doc.add_paragraph()

# I. INTRODUCTION
doc.add_heading('I. INTRODUCTION', level=1)

doc.add_paragraph(
    'Plaintiffs\' Opposition misstates the purpose of Ms. Pak\'s Motion for New Trial. The motion does not dispute the Court\'s factual findings or credibility determinations; it expressly accepts the factual findings of the Statement of Decision ("SOD") as written and argues that the judgment does not follow from those findings. That is precisely the type of legal error addressed by Code of Civil Procedure section 657(7). Plaintiffs\' Opposition, focused almost entirely on credibility and factual disputes, does not engage with the legal inconsistency that the motion actually raises.'
)

# II. PLAINTIFFS' OPPOSITION RESPONDS TO A MOTION PAK DID NOT FILE
doc.add_heading('II. PLAINTIFFS\' OPPOSITION RESPONDS TO A MOTION PAK DID NOT FILE', level=1)

doc.add_paragraph(
    'The motion does not challenge the Court\'s evaluation of witness credibility, its weighing of evidence, or its resolution of factual disputes. The motion assumes all of the Court\'s findings are correct and asks whether the legal conclusions follow. This is a proper ground for relief under CCP Section 657(7). Plaintiffs\' attempt to recast the motion as a frontal attack on the SOD is inconsistent with the motion\'s actual posture and the governing legal standard.'
)

# III. THE COURT'S FINDINGS REGARDING FIDUCIARY DUTY SCOPE
doc.add_heading('III. THE COURT\'S FINDINGS REGARDING FIDUCIARY DUTY SCOPE', level=1)

doc.add_paragraph(
    'Plaintiffs dramatically overstate the fiduciary duty the Court found. They assert that Ms. Pak owed a continuous, undifferentiated fiduciary duty from 2015 through October 21, 2023. The SOD does not support this claim.'
)

doc.add_paragraph(
    'The Court adopted Judge Recana\'s ruling that any general fiduciary duty existed only from 2015 to April 1, 2020. (SOD 24:18-19.) Beyond that date, the Court recognized only a limited, case-specific duty tied to "those cases" - the few personal-injury matters Brown retained after April 2020 that Pak continued to manage. The Court expressly confined that duty to "the handling of those cases and management of [Brown\'s] finances relating to those cases." (SOD 24:21-22.) The Court then found no breach and no damages within that limited post-April 2020 duty. (SOD 28:10-16.)'
)

doc.add_paragraph(
    'Plaintiffs\' attempt to recast this as a sweeping fiduciary obligation through 2023 is inconsistent with the SOD.'
)

# IV. PLAINTIFFS MISCHARACTERIZE THE COURT'S CREDIBILITY FINDINGS
doc.add_heading('IV. PLAINTIFFS MISCHARACTERIZE THE COURT\'S CREDIBILITY FINDINGS', level=1)

# A. Subsection
doc.add_heading('A. The SOD Relied on Pak\'s Testimony Where It Was Corroborated', level=2)

doc.add_paragraph(
    'California law permits a trial court to accept portions of a witness\'s testimony where corroborated, even if other portions are rejected. (Halagan v. Ohanesian (1967) 257 Cal.App.2d 14, 21.) That is exactly what the Court did.'
)

doc.add_paragraph(
    'The SOD relied on Ms. Pak\'s testimony where it aligned with documentary evidence, third-party witnesses, or Mr. Brown\'s own admissions. It credited her description of the profit-sharing structure and her advancement of marketing and lead costs, which were corroborated by bank records and by Plaintiffs\' forensic accountant (SOD 7:3-8; 26:7-13). It accepted her account of Brown\'s retirement discussions and his desire to transfer the mass-tort practice, which was supported by Ms. Navarro\'s testimony and by Brown\'s text messages (SOD 8:14-28; 9:1-12; 10:1-12). It relied on her testimony regarding the Androgel negotiations, which was confirmed by the written agreement and by Brown\'s emails (SOD 11:1-15). It also relied on her description of office operations during Brown\'s suspension, which was consistent with the documentary record and with the testimony of the legal assistants (SOD 12:1-18; 13:1-14). These were central factual findings, not incidental references.'
)

# B. Subsection
doc.add_heading('B. Plaintiffs Overstate the Court\'s Limited Credibility Concerns', level=2)

doc.add_paragraph(
    'Plaintiffs rely heavily on the SOD\'s statement that Ms. Pak was "entirely untrustworthy," but omit the critical qualifier: the Court made that finding only "in weighing Pak\'s counter-claims." (SOD 33:16-21.) The Court did not find her globally unreliable.'
)

doc.add_paragraph(
    'The SOD identifies only narrow areas where the Court found her testimony unpersuasive, such as inconsistent dates regarding contract termination and inadequate explanations for certain 2019 transactions. (SOD 8:20-28; 27:18-28:6.) These are routine, issue-specific credibility determinations. They do not apply to the findings underlying Plaintiffs\' claims. If the Court had found Ms. Pak wholly unreliable, it would not have relied on her testimony in any respect. Yet the SOD repeatedly does so on core issues.'
)

# C. Subsection
doc.add_heading('C. The Court\'s Key Findings Were Independent of Pak\'s Testimony', level=2)

doc.add_paragraph(
    'Even if the Court had disregarded Ms. Pak\'s testimony entirely, the findings relevant to fiduciary duty and elder financial abuse would remain unchanged. The Court\'s conclusions regarding Brown\'s consent, Brown\'s knowledge, the absence of divided loyalty, and the purpose of the 2019 lead purchases were based on documents, third-party witnesses, and the Court\'s own inferences. (SOD 7-15; 20-26; 13 n.4.) Plaintiffs\' attempt to convert limited credibility concerns into a global rejection of all testimony is inconsistent with the SOD.'
)

# D. Subsection
doc.add_heading('D. Plaintiffs\' Own Credibility Problems Are Far More Serious', level=2)

doc.add_paragraph(
    'If credibility is to be part of this analysis, it must be applied consistently. The SOD contains far more serious credibility concerns about Plaintiffs\' own witnesses. The Court found that Mr. Brown\'s testimony could be given "little weight" due to severe memory impairment (SOD 3:20-28), that he had "little to no recollection" of events from 2019 through 2023 (SOD 3:24-28), and that Mrs. Brown\'s testimony was inconsistent and contradicted by objective evidence (SOD 6:10-28). Plaintiffs cannot selectively invoke credibility principles while ignoring the Court\'s far more significant concerns about their own witnesses.'
)

# V. THE COURT'S FACTUAL FINDINGS FORECLOSE PLAINTIFFS' THEORY OF DIVIDED LOYALTY
doc.add_heading('V. THE COURT\'S FACTUAL FINDINGS FORECLOSE PLAINTIFFS\' THEORY OF DIVIDED LOYALTY', level=1)

doc.add_paragraph(
    'The SOD\'s most critical finding is this: Although Pak did not admit it, the Court concludes that she "likely purchased the leads in May with the intent that they be used by Brown," and that no agreement or understanding with Mr. Norris existed before June 2019. (SOD 13 n.4.) The Court further found that Pak only began identifying the leads as Norris\'s after Mr. Brown agreed that Norris would take over the practice, a transition corroborated by Ms. Navarro\'s testimony and the July 2019 meeting between Brown and Norris.'
)

doc.add_paragraph(
    'This finding eliminates any theory of divided loyalty or self-dealing in connection with the 2019 lead purchases. Plaintiffs do not and cannot explain how a fiduciary breaches a duty of loyalty by purchasing leads for the benefit of her principal.'
)

# VI. THE LEAD PURCHASE FINDINGS
doc.add_heading('VI. THE LEAD PURCHASE FINDINGS', level=1)

doc.add_paragraph(
    'The SOD makes clear that the Court questioned only three isolated items: (1) $13,500 of a $48,500 May 31, 2019 SMG purchase, (2) a $40,000 August 19, 2019 OPL purchase, and (3) a $135,000 October 22, 2019 Broughton Partners purchase. (SOD 27:18-28:6.) Every other 2019 lead-purchase transaction was expressly supported by documentary evidence showing that Ms. Pak used her own funds.'
)

doc.add_paragraph(
    'The Court found that Pak funded four major purchases with her own deposits: $195,000 to Broughton Partners on May 14, 2019; $35,000 of the $48,500 SMG purchase on May 31, 2019; $100,000 to SMG on June 3, 2019; and $135,000 to Broughton Partners on October 21, 2019. (SOD 26:7-13; 27:1-17.) These transactions, totaling $465,000, were therefore not part of any alleged breach of fiduciary duty. The Court\'s only criticism of the remaining three items was insufficient documentation, not disloyalty, concealment, misuse of funds, or divided loyalty.'
)

doc.add_paragraph(
    'Even if the Court questioned the documentation for three transactions, those gaps cannot support liability under any theory Plaintiffs assert. A fiduciary breach requires disloyalty, concealment, or unauthorized conduct; elder financial abuse requires a "wrongful taking," undue influence, or intent to defraud. (Welf. & Inst. Code Section 15610.30.) The SOD expressly found none of these elements. The Court found that Brown had full access to the accounts, reviewed the transactions, raised no objections, and knew Pak was purchasing leads. (SOD 6:1-12; 26:7-13.) It further found that Pak did not conceal any transactions, believed she was acting within the parties\' longstanding reimbursement structure, and acted without malice, oppression, or fraud. (SOD 30:1-6.) Under these findings, a lack of documentation cannot transform transactions the Court found were intended for Brown\'s benefit into a breach of fiduciary duty or a "wrongful taking." California law does not permit liability based on bookkeeping imperfections where the Court has already found no disloyalty, no divided loyalty, no concealment, and no fraudulent intent.'
)

# VII. THE $24,000 RENT AWARD IS UNSUPPORTED AND MUST BE REVERSED
doc.add_heading('VII. THE $24,000 RENT AWARD IS UNSUPPORTED AND MUST BE REVERSED', level=1)

# A. Subsection
doc.add_heading('A. The SOD\'s Own Findings Regarding Office Use and Rent Payments Eliminate Liability', level=2)

doc.add_paragraph(
    'The SOD\'s findings regarding the office suite foreclose any rent-related liability. The Court found that Brown "gave permission to Pak and Norris to use his staff, office suite, and office phone," and that Brown "required Pak to use the suite and staff to finish up his remaining cases." (SOD 23:16-20.) The Court further found that Brown had full access to the financial books, the office, the case files, and the court records, and that he "accepted this arrangement." (SOD 23:20-23.)'
)

doc.add_paragraph(
    'Critically, the Court found that Plaintiffs presented "no evidence" quantifying any office expenses, payroll, or rent that should have been paid by the Norris firm "at least through October 21, 2023," and that Brown "continued to put no money into the firm after April 2020," while Ms. Pak deposited funds and paid office expenses herself. (SOD 23:23-25.) The Court found no unfair competition, no misuse of office resources, and no demonstrable harm. (SOD 23:25.)'
)

doc.add_paragraph(
    'These findings eliminate any basis for concluding that Ms. Pak misused rent funds, caused rent-related damages, or breached any fiduciary duty relating to the office suite.'
)

# B. Subsection
doc.add_heading('B. The Lease Placed Sole Legal Responsibility on Brown, Not Pak', level=2)

doc.add_paragraph(
    'The lease was between 3055 Wilshire, LLC and Mr. Brown personally. (Exh. 19.) Rent was paid from Brown\'s operating account. The SOD expressly found that Brown made no rent payments after April 2020, that Ms. Pak paid the vast majority of the rent, often using her own funds, and that Mr. Norris made two payments totaling $15,000. (SOD 11:5-22.)'
)

doc.add_paragraph(
    'The lease commenced April 1, 2019 and was a 60-month term, expiring April 1, 2024. Plaintiffs\' own October 21, 2023 termination letter barred Ms. Pak from the premises, the accounts, the landlord, and any involvement with the office. For at least the final six months of the lease term, Ms. Pak had no legal right or practical ability to pay rent, negotiate rent, or communicate with the landlord.'
)

doc.add_paragraph(
    'Under the lease, only Mr. Brown, who was the sole tenant, remained responsible for rent during that period. Any "balloon" or accelerated rent obligation was therefore caused by Mr. Brown\'s own decision to oust Ms. Pak and Mr. Norris, not by any act or omission by Ms. Pak.'
)

# C. Subsection
doc.add_heading('C. The Rent Award Is Unsupported and Must Be Reduced by Mandatory Offset', level=2)

doc.add_paragraph(
    'Damages for breach of fiduciary duty must be reduced by any benefit conferred, and the SOD repeatedly finds that Brown made no rent payments after April 2020, that Ms. Pak paid the vast majority of the rent using her own funds, that Mr. Norris contributed $15,000, and that Brown benefitted from Ms. Pak\'s continued use of the office to wind down his cases until October 21, 2023. (SOD 11:5-22; 23:16-25; 28:21-23.)'
)

doc.add_paragraph(
    'The Court acknowledged that the "balloon" payment was owed as of February 29, 2024, approximately one to two months before the lease termination date. The normal rent under the lease was approximately $3,000 per month, declining to $1,800 per month after the negotiated reduction. Brown, as the sole tenant, remained solely responsible for rent during the final six months when Pak was barred from the office.'
)

doc.add_paragraph(
    'Under these findings, any "balloon" rent obligation was caused by Brown\'s own decision to oust Ms. Pak. At minimum, Ms. Pak\'s substantial rent payments and contributions must be offset against the alleged $24,000, which eliminates the award entirely.'
)

# VIII. THE ELDER FINANCIAL ABUSE FINDING IS UNSUPPORTED
doc.add_heading('VIII. THE ELDER FINANCIAL ABUSE FINDING IS UNSUPPORTED BY THE COURT\'S OWN FACTUAL FINDINGS', level=1)

doc.add_paragraph(
    'The SOD\'s elder-abuse finding cannot stand because it is unsupported by the Court\'s own factual findings and rests entirely on the erroneous premise that a fiduciary-duty breach automatically constitutes elder financial abuse.'
)

doc.add_paragraph(
    'The SOD expressly finds that Brown willingly transferred his cases, understood and cooperated in the transition, had full access to the accounts and records, accepted the arrangement, and suffered no wrongful taking of clients, fees, or office resources. (SOD 29:8-30:13; 23:16-25.) The Court further found that Plaintiffs failed to prove that rent payments were made with Brown\'s money, that Brown made no contributions after April 2020, and that Pak paid the rent and funded the office while winding down Brown\'s cases. (SOD 23:23-25; 28:21-23.)'
)

doc.add_paragraph(
    'These findings negate the statutory elements of wrongful taking, undue influence, intent to defraud, and causation under Welfare & Institutions Code Section 15610.30. The Court awarded no elder-abuse damages and imposed only mandatory attorney\'s fees, confirming that no statutory harm was proven. Because the elder-abuse conclusion is unsupported by the SOD and rests on an incorrect legal standard, it is "against law" under CCP Section 657(6) and must be modified.'
)

# IX. CONCLUSION
doc.add_heading('IX. CONCLUSION', level=1)

doc.add_paragraph(
    'The motion does not ask the Court to revisit its factual findings. It asks the Court to apply the governing legal standards to those findings. When the SOD is taken as written, the elements of breach of fiduciary duty and elder financial abuse are not met. A new trial is therefore warranted under Code of Civil Procedure section 657(7).'
)

# Save document
doc.save(r'd:\saoriverse-console\DraftShift\Docs\BrownReply\PakBrownReplyMtnNewTrial_FINAL.docx')
print('Document saved successfully at: d:\\saoriverse-console\\DraftShift\\Docs\\BrownReply\\PakBrownReplyMtnNewTrial_FINAL.docx')
