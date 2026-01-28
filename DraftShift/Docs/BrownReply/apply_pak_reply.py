from docx import Document
from pathlib import Path

ROOT = Path(__file__).resolve().parent
POLISHED = ROOT / 'PakBrownReplyMtnNewTrial_Polished_For_Pak.txt'
IN_DOCX = ROOT / 'PakBrownReplyMtnNewTrial_final.docx'
OUT_DOCX = ROOT / 'PakBrownReplyMtnNewTrial_final_for_Pak.docx'

# read polished text
with open(POLISHED, 'r', encoding='utf-8') as f:
    text = f.read().splitlines()

# load doc
if IN_DOCX.exists():
    doc = Document(str(IN_DOCX))
else:
    doc = Document()

# find existing polished heading or previous reply heading
old_heading_candidates = [
    'Reply Section — Credibility and Record (Polished Option 1)',
    'Reply Section — Credibility and Record (Short, Judge‑Friendly)'
]
found_idx = None
paras = list(doc.paragraphs)
for i,p in enumerate(paras):
    for h in old_heading_candidates:
        if h in p.text:
            found_idx = i
            break
    if found_idx is not None:
        break

if found_idx is not None:
    new = Document()
    for p in paras[:found_idx]:
        new._body._element.append(p._p)
    new.add_heading('Reply Section — Focused for Ms. Pak (Polished)', level=2)
    for ln in text:
        new.add_paragraph(ln)
    new.save(str(OUT_DOCX))
    print('WROTE', OUT_DOCX)
else:
    doc.add_heading('Reply Section — Focused for Ms. Pak (Polished)', level=2)
    for ln in text:
        doc.add_paragraph(ln)
    doc.save(str(OUT_DOCX))
    print('WROTE', OUT_DOCX)
