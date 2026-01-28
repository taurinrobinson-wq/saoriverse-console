"""
DocumentBuilder — High-Level Orchestrator for California Pleadings

Acts as the director: loads YAML configs, instantiates BaseDocument,
accepts structured input, and produces fully formatted DOCX output.
"""

from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from .base import BaseDocument


class DocumentBuilder:
    """
    High-level orchestrator for generating California pleadings.
    Takes structured data and uses BaseDocument to produce a DOCX.

    Workflow:
        1. Load YAML configs (formatting + citations)
        2. Instantiate BaseDocument with configs
        3. Build components in sequence (attorney block, caption, arguments, etc.)
        4. Apply formatting rules from YAML
        5. Generate and save DOCX
    """

    def __init__(self, config_path: str, citation_path: str):
        """
        Initialize DocumentBuilder with config file paths.

        Args:
            config_path: Path to california_civil.yaml
            citation_path: Path to california_civil_citation.yaml
        """
        self.config_path = config_path
        self.citation_path = citation_path

    # ---------------------------------------------------------
    # MAIN ENTRY POINT
    # ---------------------------------------------------------
    def build_document(self, output_path: str, data: dict) -> str:
        """
        Builds a complete California pleading using the BaseDocument engine.

        Args:
            output_path: Where to save the generated DOCX
            data: Dictionary containing:
                - attorney: Attorney info (name, bar, firm, address, contact)
                - case: Case info (county, case_number, parties, case_title)
                - title: Document title (e.g., "REPLY TO PLAINTIFFS' OPPOSITION")
                - arguments: List of dicts with 'level' (heading) or 'paragraph' (text)
                - pos: Proof of Service data (server_name, date, location, etc.)

        Returns:
            Path to generated document
        """
        # Instantiate BaseDocument with YAML configs
        doc = BaseDocument(self.config_path, self.citation_path)

        # 1. Attorney block (top of first page)
        if "attorney" in data:
            self._build_attorney_block(doc, data["attorney"])

        # 2. Caption table (court header + case number)
        if "case" in data:
            self._build_caption(doc, data["case"])

        # 3. Document title (centered, bold)
        if "title" in data:
            self._build_document_title(doc, data["title"])

        # 4. Body arguments (structured headings + paragraphs)
        if "arguments" in data:
            self._build_arguments(doc, data["arguments"])

        # 5. Signature block (attorney signature line)
        if "attorney" in data:
            doc.add_signature_block(data["attorney"])

        # 6. Proof of Service (separate page)
        if "pos" in data:
            doc.add_proof_of_service(data["pos"])

        # 7. Save document
        doc.save(output_path)

        return output_path

    # ---------------------------------------------------------
    # COMPONENT BUILDERS
    # ---------------------------------------------------------
    def _build_attorney_block(self, doc: BaseDocument, attorney: dict) -> None:
        """
        Builds the attorney block at the top of the first page.

        Expected keys in attorney dict:
            - name: Full attorney name
            - title: Title/suffix (e.g., "Esq.")
            - bar_number: State bar number
            - firm: Law firm name
            - address: Street address
            - city_state_zip: City, State ZIP
            - phone: Phone number
            - email: Email address
            - party: Party represented
        """
        # Format: Name, Esq. Bar No. XXXXXX
        p = doc.doc.add_paragraph()
        run = p.add_run(f"{attorney.get('name', '')}")
        run.font.size = Pt(12)
        p.add_run(f", {attorney.get('title', 'Esq.')}\n")

        # Firm and address block
        lines = [
            attorney.get("firm", ""),
            attorney.get("address", ""),
            attorney.get("city_state_zip", ""),
            attorney.get("phone", ""),
            attorney.get("email", ""),
        ]

        for line in lines:
            if line:
                p.add_run(f"{line}\n")

        p.add_run(f"\nAttorney for {attorney.get('party', '')}\n")
        p.add_run(f"State Bar No. {attorney.get('bar_number', '')}")

        p.paragraph_format.space_after = Pt(12)

    def _build_caption(self, doc: BaseDocument, case: dict) -> None:
        """
        Builds the caption section (court header + case number).

        Expected keys in case dict:
            - county: County name
            - case_number: Case number
            - parties: Dict with 'plaintiff' and 'defendant'
        """
        doc.add_caption({
            "county": case.get("county", ""),
            "case_number": case.get("case_number", ""),
            "parties": case.get("parties", {}),
        })

    def _build_document_title(self, doc: BaseDocument, title: str) -> None:
        """
        Builds the document title (centered, bold, e.g., "REPLY TO OPPOSITION").

        Args:
            title: Title text
        """
        p = doc.doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)

    def _build_arguments(self, doc: BaseDocument, arguments: list) -> None:
        """
        Builds the body arguments from a structured list.

        Each item in arguments should be one of:
            - {"level": 1, "text": "HEADING TEXT"}  (adds heading with auto-numbering)
            - {"paragraph": "Body text content..."}  (adds body paragraph)

        Example:
            [
                {"level": 1, "text": "INTRODUCTION"},
                {"level": 2, "text": "Legal Standard"},
                {"paragraph": "California law imposes..."},
                {"level": 1, "text": "ARGUMENT"},
                {"level": 2, "text": "Defendant Breached"},
                {"paragraph": "The evidence shows..."},
            ]

        Args:
            arguments: List of heading/paragraph dicts
        """
        for item in arguments:
            if "level" in item and "text" in item:
                # Add heading with automatic numbering
                doc.add_heading(item["text"], item["level"])

            elif "paragraph" in item:
                # Add body paragraph
                doc.add_paragraph(item["paragraph"], spacing="body")

    # ---------------------------------------------------------
    # BUILDER PATTERN HELPERS (for fluent API)
    # ---------------------------------------------------------
    def with_attorney(self, attorney_data: dict) -> "DocumentBuilder":
        """
        Fluent API: store attorney data for later build.
        """
        self._attorney_data = attorney_data
        return self

    def with_case(self, case_data: dict) -> "DocumentBuilder":
        """
        Fluent API: store case data for later build.
        """
        self._case_data = case_data
        return self

    def with_arguments(self, arguments: list) -> "DocumentBuilder":
        """
        Fluent API: store arguments for later build.
        """
        self._arguments = arguments
        return self

    # ---------------------------------------------------------
    # VALIDATION
    # ---------------------------------------------------------
    def validate_data(self, data: dict) -> bool:
        """
        Validates that required fields are present in input data.

        Args:
            data: Input data dictionary

        Returns:
            True if all required fields present, raises ValueError otherwise
        """
        required_top_level = ["attorney", "case", "title", "arguments"]

        for field in required_top_level:
            if field not in data:
                raise ValueError(f"Missing required field: {field}")

        # Validate attorney fields
        attorney_required = ["name", "bar_number", "party"]
        for field in attorney_required:
            if field not in data.get("attorney", {}):
                raise ValueError(f"Missing required attorney field: {field}")

        # Validate case fields
        case_required = ["county", "case_number", "parties"]
        for field in case_required:
            if field not in data.get("case", {}):
                raise ValueError(f"Missing required case field: {field}")

        # Validate parties
        parties = data.get("case", {}).get("parties", {})
        if "plaintiff" not in parties or "defendant" not in parties:
            raise ValueError("Missing plaintiff and/or defendant in case.parties")

        return True

    # ---------------------------------------------------------
    # PRESET BUILDERS (for specific pleading types)
    # ---------------------------------------------------------
    @classmethod
    def build_reply(cls, config_path: str, citation_path: str, data: dict, output_path: str) -> str:
        """
        Convenience method for building a Reply brief.

        Args:
            config_path: Path to california_civil.yaml
            citation_path: Path to california_civil_citation.yaml
            data: Input data with attorney, case, arguments, pos
            output_path: Where to save DOCX

        Returns:
            Path to generated document
        """
        builder = cls(config_path, citation_path)
        builder.validate_data(data)
        return builder.build_document(output_path, data)

    @classmethod
    def build_motion(cls, config_path: str, citation_path: str, data: dict, output_path: str) -> str:
        """
        Convenience method for building a Motion.
        """
        builder = cls(config_path, citation_path)
        builder.validate_data(data)
        return builder.build_document(output_path, data)

    @classmethod
    def build_opposition(cls, config_path: str, citation_path: str, data: dict, output_path: str) -> str:
        """
        Convenience method for building an Opposition.
        """
        builder = cls(config_path, citation_path)
        builder.validate_data(data)
        return builder.build_document(output_path, data)
