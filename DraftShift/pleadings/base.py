"""
BaseDocument — Core Pleading Engine for Draftshift

Loads formatting + citation YAML configs and exposes helper methods
for building California-compliant pleadings.
"""

import yaml
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches


class BaseDocument:
    """
    Core pleading engine for Draftshift.
    Loads formatting + citation YAML configs and exposes
    helper methods for building California-compliant pleadings.
    """

    def __init__(self, config_path: str, citation_path: str):
        """
        Initialize BaseDocument with formatting and citation configs.
        
        Args:
            config_path: Path to california_civil.yaml
            citation_path: Path to california_civil_citation.yaml
        """
        self.config = self._load_yaml(config_path)
        self.citations = self._load_yaml(citation_path)

        # Merge citation rules into main config tree
        if "citation" not in self.config:
            self.config["citation"] = {}
        self.config["citation"].update(self.citations.get("citation", {}))

        # Initialize DOCX document
        self.doc = Document()

        # Apply base page geometry
        self._apply_page_margins()

        # Heading counters for automatic numbering
        self._heading_counters = {1: 0, 2: 0, 3: 0, 4: 0}

    # ---------------------------------------------------------
    # YAML LOADING
    # ---------------------------------------------------------
    def _load_yaml(self, path: str) -> dict:
        """Load and parse YAML configuration file."""
        with open(Path(path), "r", encoding="utf-8") as f:
            return yaml.safe_load(f)

    # ---------------------------------------------------------
    # PAGE GEOMETRY
    # ---------------------------------------------------------
    def _apply_page_margins(self):
        """Apply page margins from YAML config to document."""
        margins = self.config.get("page_geometry", {}).get("margins", {})
        section = self.doc.sections[0]

        section.top_margin = Inches(float(margins.get("top", 1.0)))
        section.bottom_margin = Inches(float(margins.get("bottom", 1.0)))
        section.left_margin = Inches(float(margins.get("left", 1.0)))
        section.right_margin = Inches(float(margins.get("right", 1.0)))

    # ---------------------------------------------------------
    # PARAGRAPH HELPERS
    # ---------------------------------------------------------
    def add_paragraph(self, text: str, style: str = None, spacing: str = "body") -> object:
        """
        Adds a paragraph using the spacing rules defined in the YAML.
        
        Args:
            text: Paragraph content
            style: Paragraph style name (optional)
            spacing: Spacing profile ('body', 'caption', 'signature')
        
        Returns:
            Paragraph object
        """
        p = self.doc.add_paragraph(text)

        if spacing == "body":
            rules = self.config.get("body_text", {})
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(float(rules.get("line_spacing_pt", 24.0)))
            p.paragraph_format.first_line_indent = Inches(0.5)

        elif spacing == "caption":
            rules = self.config.get("caption", {})
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(float(rules.get("line_spacing", 12)))

        elif spacing == "signature":
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(12)

        return p

    # ---------------------------------------------------------
    # HEADINGS
    # ---------------------------------------------------------
    def add_heading(self, text: str, level: int) -> object:
        """
        Adds a heading using the YAML-defined hierarchy.
        Automatically generates numbering (I, 1, A, i) for each level.
        
        Args:
            text: Heading text
            level: Heading level (1-4)
        
        Returns:
            Paragraph object
        """
        level_key = f"level_{level}"
        rules = self.config.get("argument_headings", {}).get(level_key, {})

        # Increment counter for this level
        self._heading_counters[level] += 1

        # Reset subordinate level counters
        for sub_level in range(level + 1, 5):
            self._heading_counters[sub_level] = 0

        # Generate heading number
        number = self._generate_heading_number(level, self._heading_counters[level])

        # Create paragraph
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(float(rules.get("indent", 0)))
        p.paragraph_format.space_before = Pt(float(rules.get("spacing_before", 6)))
        p.paragraph_format.space_after = Pt(float(rules.get("spacing_after", 3)))

        # Add numbering
        run_num = p.add_run(f"{number} ")
        run_num.bold = rules.get("bold", True)
        run_num.font.size = Pt(float(rules.get("font_size", 12)))

        # Add heading text
        run_text = p.add_run(text)
        run_text.bold = rules.get("bold", True)
        run_text.italic = rules.get("italic", False)
        run_text.font.size = Pt(float(rules.get("font_size", 12)))

        # Alignment
        alignment = rules.get("alignment", "left")
        if alignment == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Line spacing
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(float(self.config.get("body_text", {}).get("line_spacing_pt", 24)))

        return p

    def _generate_heading_number(self, level: int, counter: int) -> str:
        """
        Generate heading number based on level and counter.
        Level 1: Roman uppercase (I, II, III)
        Level 2: Numeric (1, 2, 3)
        Level 3: Alpha uppercase (A, B, C)
        Level 4: Roman lowercase (i, ii, iii)
        """
        if level == 1:
            return self._int_to_roman(counter).upper()
        elif level == 2:
            return str(counter)
        elif level == 3:
            return chr(64 + counter)  # A, B, C, ...
        elif level == 4:
            return self._int_to_roman(counter).lower()
        return ""

    @staticmethod
    def _int_to_roman(num: int) -> str:
        """Convert integer to Roman numeral."""
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        syms = [
            "M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"
        ]
        roman_num = ""
        i = 0
        while num > 0:
            for _ in range(num // val[i]):
                roman_num += syms[i]
                num -= val[i]
            i += 1
        return roman_num

    # ---------------------------------------------------------
    # CAPTION BLOCK
    # ---------------------------------------------------------
    def add_caption(self, case_data: dict) -> object:
        """
        Builds the two-column caption table with correct borders,
        margins, and formatting per California civil pleading rules.
        
        Args:
            case_data: Dictionary containing:
                - county: County name
                - case_number: Case number
                - parties: Dictionary with plaintiff and defendant
        
        Returns:
            Table object
        """
        caption_cfg = self.config.get("caption", {})
        
        # Create two-column table for caption
        table = self.doc.add_table(rows=2, cols=2)
        table.autofit = False

        # Left column: Superior court header
        left_cell = table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.add_run("SUPERIOR COURT OF CALIFORNIA\n")
        left_para.add_run(f"COUNTY OF {case_data.get('county', '').upper()}")
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Right column: Case number
        right_cell = table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.add_run(f"Case No. {case_data.get('case_number', '')}")
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Apply borders to left cell (right and bottom only)
        self._set_cell_border(left_cell, right=True, bottom=True)

        # Row 2: Case title
        title_cell = table.cell(1, 0)
        title_cell.merge(table.cell(1, 1))
        title_para = title_cell.paragraphs[0]
        
        plaintiff = case_data.get("parties", {}).get("plaintiff", "")
        defendant = case_data.get("parties", {}).get("defendant", "")
        
        run = title_para.add_run(f"{plaintiff} v. {defendant}")
        run.italic = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return table

    def _set_cell_border(self, cell, top=False, bottom=False, left=False, right=False, color="000000"):
        """Apply borders to a table cell."""
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")

        for border_name in ["top", "left", "bottom", "right"]:
            if locals()[border_name]:
                border_el = OxmlElement(f"w:{border_name}")
                border_el.set(qn("w:val"), "single")
                border_el.set(qn("w:sz"), "12")  # 1.5pt
                border_el.set(qn("w:space"), "0")
                border_el.set(qn("w:color"), color)
                tcBorders.append(border_el)

        tcPr.append(tcBorders)

    # ---------------------------------------------------------
    # ATTORNEY BLOCK
    # ---------------------------------------------------------
    def add_attorney_block(self, attorney_data: dict) -> object:
        """
        Adds attorney block at bottom right of first page.
        
        Args:
            attorney_data: Dictionary containing:
                - name: Attorney name
                - title: Attorney title (e.g., "Esq.")
                - firm: Firm name
                - address: Street address
                - city_state_zip: City, State ZIP
                - phone: Phone number
                - email: Email address
                - bar_number: State bar number
                - party: Party represented
        
        Returns:
            Paragraph object
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)

        # "Attorney for [Party]"
        p.add_run(f"Attorney for {attorney_data.get('party', '')}\n")

        # Tab to bar number column
        p.add_run("\t")
        p.add_run(f"Bar No. {attorney_data.get('bar_number', '')}\n\n")

        # Attorney details (left column)
        lines = [
            attorney_data.get("name", ""),
            attorney_data.get("firm", ""),
            attorney_data.get("address", ""),
            attorney_data.get("city_state_zip", ""),
            attorney_data.get("phone", ""),
            attorney_data.get("email", ""),
        ]

        for line in lines:
            if line:
                p.add_run(f"{line}\n")

        return p

    # ---------------------------------------------------------
    # SIGNATURE BLOCK
    # ---------------------------------------------------------
    def add_signature_block(self, attorney: dict) -> object:
        """
        Adds signature block with date and attorney name.
        
        Args:
            attorney: Dictionary containing:
                - name: Attorney name
                - title: Attorney title
        
        Returns:
            Paragraph object
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)

        # Date line
        p.add_run("Dated: _______________\n\n")

        # Signature line
        p.add_run("_______________________________\n")
        run = p.add_run(f"{attorney.get('name', '')}")
        run.bold = True
        p.add_run(f"\n{attorney.get('title', '')}")

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return p

    # ---------------------------------------------------------
    # PROOF OF SERVICE
    # ---------------------------------------------------------
    def add_proof_of_service(self, pos_data: dict) -> bool:
        """
        Adds Proof of Service section on new page.
        
        Args:
            pos_data: Dictionary containing:
                - server_name: Name of person serving
                - server_county: County of service
                - document_title: Title of document served
                - service_date: Date of service
                - service_method: How document was served
                - service_location: Where document was served
                - server_declaration: Declarant statement
        
        Returns:
            True on success
        """
        self.doc.add_page_break()

        # Heading
        heading = self.doc.add_paragraph("PROOF OF SERVICE")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].bold = True
        heading.paragraph_format.space_after = Pt(12)

        # Declaration body
        declaration_text = (
            f"I, {pos_data.get('server_name', '')}, declare under penalty of perjury "
            f"under the laws of the State of California that I am employed in the "
            f"County of {pos_data.get('server_county', '')}, State of California. "
            f"I am over the age of eighteen years and not a party to the within action. "
            f"I am familiar with this office's practice of collection and processing "
            f"correspondence for mailing with the United States Postal Service.\n\n"
            f"On {pos_data.get('service_date', '')}, I enclosed the within document(s) "
            f"described as {pos_data.get('document_title', '')} in a sealed envelope "
            f"with postage thereon fully prepaid, in the ordinary course of business, "
            f"and deposited said envelope for collection and mailing at "
            f"{pos_data.get('service_location', '')}.\n\n"
            f"I declare under penalty of perjury that the foregoing is true and correct. "
            f"Executed on {pos_data.get('execution_date', '')} at "
            f"{pos_data.get('execution_location', '')}, California."
        )

        self.add_paragraph(declaration_text, spacing="body")

        # Signature line
        sig_p = self.doc.add_paragraph()
        sig_p.paragraph_format.space_before = Pt(24)
        sig_p.add_run("_______________________________\n")
        sig_p.add_run(pos_data.get('server_name', ''))

        return True

    # ---------------------------------------------------------
    # SAVE
    # ---------------------------------------------------------
    def save(self, path: str) -> bool:
        """
        Save document to specified path.
        
        Args:
            path: Output file path
        
        Returns:
            True on success
        """
        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        return True
