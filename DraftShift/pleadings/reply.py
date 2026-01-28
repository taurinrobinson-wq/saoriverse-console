"""
Reply — Reply Class for California Civil Pleadings

Represents a Reply in support of a Motion or opposition.
Inherits formatting + citation rules from BaseDocument.
Exposes a clean build() method for assembling replies.
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from .base import BaseDocument


class Reply(BaseDocument):
    """
    Represents a Reply in support of a Motion or opposition.
    Inherits formatting + citation rules from BaseDocument.

    Structure:
        1. Attorney block (top of page)
        2. Court header
        3. Caption table (court name + case number)
        4. Document title (reply title)
        5. Arguments (structured headings + paragraphs)
        6. Signature block
        7. Proof of Service (separate page)

    Replies are structurally simpler than Motions:
        - No Notice of Motion section
        - Often shorter
        - Same caption + argument structure
    """

    def __init__(self, config_path: str, citation_path: str):
        """
        Initialize Reply with formatting and citation configs.

        Args:
            config_path: Path to california_civil.yaml
            citation_path: Path to california_civil_citation.yaml
        """
        super().__init__(config_path, citation_path)

    # ---------------------------------------------------------
    # MAIN BUILD METHOD
    # ---------------------------------------------------------
    def build(self, data: dict) -> None:
        """
        Builds a complete reply using structured data.

        Expected data structure:
        {
            "attorney": {
                "name": "...",
                "title": "Esq.",
                "bar_number": "...",
                "firm": "...",
                "address": "...",
                "city_state_zip": "...",
                "phone": "...",
                "email": "...",
                "party": "..."
            },
            "case": {
                "county": "Los Angeles",
                "case_number": "...",
                "parties": {"plaintiff": "...", "defendant": "..."}
            },
            "title": "REPLY TO PLAINTIFFS' OPPOSITION TO MOTION FOR NEW TRIAL",
            "arguments": [
                {"level": 1, "text": "PROCEDURAL POSTURE"},
                {"paragraph": "Plaintiffs filed an opposition..."},
                {"level": 1, "text": "ARGUMENT"},
                {"level": 2, "text": "I. Fiduciary Duty Scope"},
                {"paragraph": "The court's findings establish..."},
                ...
            ],
            "pos": {
                "server_name": "...",
                "service_date": "...",
                ...
            }
        }

        Args:
            data: Dictionary with attorney, case, title, arguments, pos
        """
        # 1. Attorney block
        self._build_attorney_block(data["attorney"])

        # 2. Court header
        self._build_court_header(data["case"])

        # 3. Caption table
        self.add_caption({
            "county": data["case"]["county"],
            "case_number": data["case"]["case_number"],
            "parties": data["case"]["parties"],
        })

        # 4. Document title
        self._build_document_title(data["title"])

        # 5. Arguments
        self._build_arguments(data.get("arguments", []))

        # 6. Signature block
        self.add_signature_block(data["attorney"])

        # 7. Proof of service
        if "pos" in data:
            self.add_proof_of_service(data["pos"])

    # ---------------------------------------------------------
    # COMPONENT BUILDERS
    # ---------------------------------------------------------
    def _build_attorney_block(self, attorney: dict) -> None:
        """
        Builds attorney block at top of first page.

        Args:
            attorney: Dictionary with name, bar_number, firm, address, etc.
        """
        p = self.doc.add_paragraph()

        # Name line
        run = p.add_run(f"{attorney.get('name', '')}")
        run.font.size = Pt(12)
        p.add_run(f", {attorney.get('title', 'Esq.')}\n")

        # Firm and address block
        lines = [
            attorney.get("firm", ""),
            attorney.get("address", ""),
            attorney.get("city_state_zip", ""),
            attorney.get("phone", ""),
            attorney.get("email", ""),
        ]

        for line in lines:
            if line:
                p.add_run(f"{line}\n")

        p.add_run(f"\nAttorney for {attorney.get('party', '')}\n")
        p.add_run(f"State Bar No. {attorney.get('bar_number', '')}")

        p.paragraph_format.space_after = Pt(12)

    def _build_court_header(self, case: dict) -> None:
        """
        Builds court header section.

        Args:
            case: Dictionary with county, case_number, etc.
        """
        p1 = self.doc.add_paragraph()
        p1.add_run("SUPERIOR COURT OF THE STATE OF CALIFORNIA")
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_after = Pt(0)

        p2 = self.doc.add_paragraph()
        p2.add_run(f"COUNTY OF {case.get('county', '').upper()}")
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_after = Pt(12)

    def _build_document_title(self, title: str) -> None:
        """
        Builds the document title (centered, bold).

        Args:
            title: Reply title (e.g., "REPLY TO PLAINTIFFS' OPPOSITION")
        """
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)

    def _build_arguments(self, arguments: list) -> None:
        """
        Builds argument section from structured list.

        Each item should be:
            - {"level": N, "text": "HEADING"}
            - {"paragraph": "Body text..."}

        Args:
            arguments: List of heading/paragraph dicts
        """
        for item in arguments:
            if "level" in item and "text" in item:
                self.add_heading(item["text"], item["level"])
            elif "paragraph" in item:
                self.add_paragraph(item["paragraph"], spacing="body")
