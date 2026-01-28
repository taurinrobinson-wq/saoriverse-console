"""
Declaration — Declaration Class for California Civil Pleadings

Represents a Declaration in support of a Motion or Opposition.
Inherits formatting + citation rules from BaseDocument.
Exposes a clean build() method for assembling declarations.
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from .base import BaseDocument


class Declaration(BaseDocument):
    """
    Represents a Declaration in support of a Motion or Opposition.
    Inherits formatting + citation rules from BaseDocument.

    Structure (distinct from Motion/Reply/Opposition):
        1. Attorney block (top of page)
        2. Court header
        3. Caption table (court name + case number)
        4. Document title (declaration title)
        5. Declarant intro ("I, [Name], declare as follows:")
        6. Numbered paragraphs (with mandatory attestation as para 1)
        7. Signature block (declarant signature, not attorney)
        8. Proof of Service (separate page)

    Declarations use numbered paragraphs instead of heading hierarchy,
    and automatically prepend a mandatory attestation sentence.
    """

    def __init__(self, config_path: str, citation_path: str):
        """
        Initialize Declaration with formatting and citation configs.

        Args:
            config_path: Path to california_civil.yaml
            citation_path: Path to california_civil_citation.yaml
        """
        super().__init__(config_path, citation_path)

    # ---------------------------------------------------------
    # MAIN BUILD METHOD
    # ---------------------------------------------------------
    def build(self, data: dict) -> None:
        """
        Builds a complete declaration using structured data.

        Expected data structure:
        {
            "attorney": {
                "name": "...",  # Attorney name (for caption)
                "title": "Esq.",
                "bar_number": "...",
                "firm": "...",
                "address": "...",
                "city_state_zip": "...",
                "phone": "...",
                "email": "...",
                "party": "..."
            },
            "case": {
                "county": "Los Angeles",
                "case_number": "...",
                "parties": {"plaintiff": "...", "defendant": "..."}
            },
            "title": "DECLARATION OF MICHELLE PAK IN SUPPORT OF MOTION FOR NEW TRIAL",
            "declarant": {
                "name": "Michelle Pak",
                "type": "Defendant",  # Used for mandatory attestation
                "title": "Declarant"   # For signature line
            },
            "statements": [
                "On or about [date], I...",
                "The facts are as follows...",
                ...
            ],
            "pos": {
                "server_name": "...",
                "service_date": "...",
                ...
            }
        }

        Args:
            data: Dictionary with attorney, case, title, declarant, statements, pos
        """
        # 1. Attorney block
        self._build_attorney_block(data["attorney"])

        # 2. Court header
        self._build_court_header(data["case"])

        # 3. Caption table
        self.add_caption({
            "county": data["case"]["county"],
            "case_number": data["case"]["case_number"],
            "parties": data["case"]["parties"],
        })

        # 4. Document title
        self._build_document_title(data["title"])

        # 5. Declarant intro ("I, [Name], declare as follows:")
        self._build_intro(data["declarant"])

        # 6. Numbered paragraphs (with auto-attestation as para 1)
        self._build_numbered_paragraphs(
            data.get("statements", []),
            declarant=data["declarant"]
        )

        # 7. Signature block (declarant signature)
        self.add_signature_block(data["declarant"])

        # 8. Proof of service
        if "pos" in data:
            self.add_proof_of_service(data["pos"])

    # ---------------------------------------------------------
    # DECLARATION-SPECIFIC COMPONENTS
    # ---------------------------------------------------------
    def _build_intro(self, declarant: dict) -> None:
        """
        Builds the declarant intro paragraph.
        Format: "I, [Name], declare as follows:"

        Args:
            declarant: Dictionary with 'name'
        """
        intro_text = f"I, {declarant.get('name', '')}, declare as follows:"
        self.add_paragraph(intro_text, spacing="body")

    def _build_numbered_paragraphs(self, statements: list, declarant: dict) -> None:
        """
        Builds numbered paragraphs for factual statements.
        Automatically prepends a mandatory attestation sentence as paragraph 1.

        Args:
            statements: List of factual statement strings
            declarant: Dictionary with 'type' (e.g., "Plaintiff", "Defendant")
        """
        # Mandatory attestation (always paragraph 1)
        attestation = (
            f"I am the {declarant.get('type', '')} in this action. "
            "I have personal knowledge of the facts set forth in this declaration, "
            "and if called as a witness, I could and would testify competently thereto."
        )

        numbered = [attestation] + statements

        for counter, text in enumerate(numbered, start=1):
            self.add_paragraph(f"{counter}. {text}", spacing="body")

    # ---------------------------------------------------------
    # COMPONENT BUILDERS
    # ---------------------------------------------------------
    def _build_attorney_block(self, attorney: dict) -> None:
        """
        Builds attorney block at top of first page.

        Args:
            attorney: Dictionary with name, bar_number, firm, address, etc.
        """
        p = self.doc.add_paragraph()

        # Name line
        run = p.add_run(f"{attorney.get('name', '')}")
        run.font.size = Pt(12)
        p.add_run(f", {attorney.get('title', 'Esq.')}\n")

        # Firm and address block
        lines = [
            attorney.get("firm", ""),
            attorney.get("address", ""),
            attorney.get("city_state_zip", ""),
            attorney.get("phone", ""),
            attorney.get("email", ""),
        ]

        for line in lines:
            if line:
                p.add_run(f"{line}\n")

        p.add_run(f"\nAttorney for {attorney.get('party', '')}\n")
        p.add_run(f"State Bar No. {attorney.get('bar_number', '')}")

        p.paragraph_format.space_after = Pt(12)

    def _build_court_header(self, case: dict) -> None:
        """
        Builds court header section.

        Args:
            case: Dictionary with county, case_number, etc.
        """
        p1 = self.doc.add_paragraph()
        p1.add_run("SUPERIOR COURT OF THE STATE OF CALIFORNIA")
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_after = Pt(0)

        p2 = self.doc.add_paragraph()
        p2.add_run(f"COUNTY OF {case.get('county', '').upper()}")
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_after = Pt(12)

    def _build_document_title(self, title: str) -> None:
        """
        Builds the document title (centered, bold).

        Args:
            title: Declaration title (e.g., "DECLARATION OF MICHELLE PAK...")
        """
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
