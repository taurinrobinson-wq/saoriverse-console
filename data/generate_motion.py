"""
Generates the consolidated Motion for Reconsideration re Attorney Fees
(Michelle Pak - Case No. 23STCV27439) as a properly formatted .docx file
compliant with Los Angeles County Superior Court local rules and CCP § 1008.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Motion_Reconsideration_AttyFees_DRAFT.docx")

doc = Document()

# ── Page setup ─────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.5)   # extra left margin for line-number column
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Default style ───────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after  = Pt(0)

# ── Helpers ─────────────────────────────────────────────────────────────────────
def body(text="", *, bold=False, center=False, indent=0, space_after=12):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent  = Inches(indent)
        p.paragraph_format.right_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

def heading(text, *, roman=None, center=True):
    """Section heading — all-caps bold centered (or left-aligned with roman numeral)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    label = f"{roman}. " if roman else ""
    run = p.add_run(label + text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

def subheading(text, *, letter=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    label = f"{letter}. " if letter else ""
    run = p.add_run(label + text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

def hr():
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

def cite(text):
    """Italicized case citation helper — adds inline italic run."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

def bullet(text, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

def table_row(tbl, *cells, bold_first=False):
    row = tbl.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                if bold_first and i == 0:
                    run.bold = True
    return row

# ══════════════════════════════════════════════════════════════════════════════
#  CAPTION
# ══════════════════════════════════════════════════════════════════════════════

# Attorney block (top-left)
atty_block = [
    "TAURIN J. ROBINSON, ESQ. (SBN [___])",
    "LAW OFFICE OF TAURIN J. ROBINSON",
    "3055 Wilshire Blvd., Suite 980",
    "Los Angeles, California 90010",
    "Telephone: [___-___-____]",
    "Facsimile: [___-___-____]",
    "Email: [_________@_________]",
    "",
    "Attorney for Defendant and Cross-Complainant",
    "MICHELLE PAK",
]
for line in atty_block:
    p = doc.add_paragraph(line)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

doc.add_paragraph()

# Court name
body("SUPERIOR COURT OF THE STATE OF CALIFORNIA", bold=True, center=True)
body("FOR THE COUNTY OF LOS ANGELES", bold=True, center=True)
body("CENTRAL DISTRICT", bold=True, center=True)

doc.add_paragraph()

# Caption table
cap_tbl = doc.add_table(rows=1, cols=2)
cap_tbl.style = "Table Grid"
cap_tbl.autofit = False
cap_tbl.columns[0].width = Inches(3.5)
cap_tbl.columns[1].width = Inches(3.5)

left_cell  = cap_tbl.rows[0].cells[0]
right_cell = cap_tbl.rows[0].cells[1]

left_lines = [
    "BAIRD A. BROWN, an individual; et al.,",
    "",
    "    Plaintiffs,",
    "",
    "vs.",
    "",
    "MICHELLE PAK, an individual; et al.,",
    "",
    "    Defendants.",
]
right_lines = [
    "Case No. 23STCV27439",
    "",
    "NOTICE OF MOTION AND MOTION FOR",
    "RECONSIDERATION OF THE COURT'S",
    "APRIL 30, 2026 ORDER ON PLAINTIFFS'",
    "MOTION FOR ATTORNEY'S FEES AND",
    "COSTS; MEMORANDUM OF POINTS AND",
    "AUTHORITIES IN SUPPORT THEREOF;",
    "DECLARATION OF TAURIN J. ROBINSON;",
    "EXHIBITS",
    "",
    "Hearing Date: [___________]",
    "Time: [_______]",
    "Dept.: [___]",
    "",
    "Judge: [Hon. _______________]",
    "",
    "Action Filed: November 28, 2023",
    "Trial Date: February 3, 2025",
]

for i, line in enumerate(left_lines):
    if i == 0:
        p = left_cell.paragraphs[0]
    else:
        p = left_cell.add_paragraph()
    p.text = line
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

for i, line in enumerate(right_lines):
    if i == 0:
        p = right_cell.paragraphs[0]
    else:
        p = right_cell.add_paragraph()
    p.text = line
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        if i in (0, 2, 3, 4, 5, 6, 7, 8, 9):
            run.bold = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  NOTICE OF MOTION
# ══════════════════════════════════════════════════════════════════════════════

heading("NOTICE OF MOTION AND MOTION FOR RECONSIDERATION")

body(
    "TO ALL PARTIES AND THEIR ATTORNEYS OF RECORD:"
)

body(
    "PLEASE TAKE NOTICE that on [DATE], at [TIME], in Department [___] of the above-entitled "
    "Court, located at [COURTHOUSE ADDRESS], Defendant and Cross-Complainant MICHELLE PAK "
    "('Pak') will, and hereby does, move this Court pursuant to Code of Civil Procedure section "
    "1008, subdivision (a), for reconsideration of its April 30, 2026 Order on Plaintiffs' "
    "Motion for Attorney's Fees and Costs (the 'Order'), which awarded $532,936 in attorney's "
    "fees and $9,973 in costs, for a total of $542,909."
)

body(
    "This Motion is based on new facts and circumstances arising from the Order itself — "
    "specifically, the internal methodological inconsistencies identified through a line-by-line "
    "comparison of the April 30, 2026 Order against Plaintiffs' billing records — which could "
    "not have been identified or argued prior to the issuance of the Order. (Even Zohar "
    "Construction & Remodeling, Inc. v. Bellaire Townhouses, LLC (2015) 61 Cal.4th 830, 839.)"
)

body(
    "The grounds for this Motion are:"
)

body(
    "(1) The Court's NR (not reasonably necessary) exclusion was applied to certain trial-day "
    "entries but not to substantively comparable entries by the same or other timekeepers for "
    "the identical court sessions;"
)
body(
    "(2) The Court's duplication reduction was applied inconsistently between Andrew White and "
    "Jonathan White, and the Order states two irreconcilable staffing standards;"
)
body(
    "(3) The Court's clerical/administrative exclusion was applied to some entries but not to "
    "substantively identical entries by Andrew White;"
)
body(
    "(4) Percentage reductions for block billing, duplication, and other categories were stacked "
    "against entries already subject to other reductions, generating a compounding effect "
    "inconsistent with the Court's stated methodology;"
)
body(
    "(5) Andrew White's and Jonathan White's billing entries for work on unrelated matters "
    "(Norris, Robinson, and internal BrownLaw disputes), which the Court held to be "
    "non-compensable, were not excluded in their entirety; and"
)
body(
    "(6) The award of post-judgment fees and fees-on-fees exceeds the statutory authorization "
    "under Welfare and Institutions Code section 15657.5, subdivision (a)."
)

body(
    "Pak respectfully requests that the Court modify its fee award to $139,646.80, the only "
    "amount mathematically consistent with the Court's stated methodology and governing law."
)

body(
    "This Motion is based on this Notice, the attached Memorandum of Points and Authorities, "
    "the concurrently filed Declaration of Taurin J. Robinson, the exhibits attached thereto, "
    "the pleadings and records on file in this action, and such oral argument as the Court may "
    "permit at the hearing."
)

body("DATED: May 15, 2026", bold=False, center=False)
doc.add_paragraph()
body("Respectfully submitted,", center=False)
doc.add_paragraph()
doc.add_paragraph()
body("By: _________________________________", center=False)
body("    Taurin J. Robinson, Esq.", center=False)
body("    Attorney for Defendant/Cross-Complainant Michelle Pak", center=False)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  MEMORANDUM OF POINTS AND AUTHORITIES
# ══════════════════════════════════════════════════════════════════════════════

heading("MEMORANDUM OF POINTS AND AUTHORITIES")

# ── INTRODUCTION ───────────────────────────────────────────────────────────────
heading("INTRODUCTION", center=True)

body(
    "Pak brings this Motion not to relitigate the Court's discretion to award attorney's fees in "
    "an amount less than the full amount requested, nor to challenge the existence or propriety "
    "of any category of reduction the Court adopted. This Motion accepts as its premise each of "
    "the five categories of reductions the Court identified in its April 30, 2026 Order. What "
    "Pak respectfully urges is a single, uncomplicated proposition: once a court selects a "
    "methodology, it must apply that methodology with internal consistency across the entire fee "
    "record. The Order, on its face, does not do so."
)

body(
    "Pak does not concede that Plaintiffs are entitled to attorney's fees under Welfare and "
    "Institutions Code section 15657.5 and reserves that issue for appeal. This Motion concerns "
    "only the fee award and raises issues that could not have been anticipated prior to the "
    "Court's April 30, 2026 ruling. Pak does not seek reconsideration of the Court's costs "
    "ruling."
)

body(
    "The April 30, 2026 Order applied five distinct categories of reductions: (1) exclusion of "
    "certain trial-day entries as 'not reasonably necessary' (NR); (2) percentage and line-item "
    "reductions for duplicative effort (DW); (3) discounting of paralegal and clerical billing "
    "entries (CL); (4) percentage reductions for block billing (BB); and (5) an award of fees "
    "for post-judgment work and the fee application itself (fees-on-fees). A careful, "
    "entry-by-entry review of the billing records and the Order reveals, however, that each of "
    "these categories was applied inconsistently. The relief Pak seeks, on each ground, is "
    "modest: consistent application of the Court's own Order."
)

body(
    "When the Court's stated methodology is applied uniformly across the entire billing record, "
    "the only fee award consistent with the Order is $139,646.80."
)

# ── PROCEDURAL HISTORY ─────────────────────────────────────────────────────────
heading("PROCEDURAL HISTORY", center=True)

body(
    "Plaintiffs filed this elder financial abuse action on November 28, 2023, in the Superior "
    "Court of California, County of Los Angeles, Case No. 23STCV27439. The complaint alleged "
    "financial elder abuse under Welfare and Institutions Code section 15657.5 and sought, "
    "among other relief, recovery of attorney's fees and costs pursuant to that statute."
)

body(
    "On January 12, 2026, the Court entered judgment in favor of Plaintiff Baird A. Brown on "
    "his claims for elder financial abuse and breach of fiduciary duty. On February 26, 2026, "
    "Plaintiffs filed their motion for attorney's fees, seeking $1,239,255 in attorney's fees "
    "and $199,230.43 in costs under Welfare and Institutions Code section 15657.5, subdivision "
    "(a). Defendant Pak filed her opposition on March 3, 2026, and Plaintiffs filed their reply "
    "on March 18, 2026. The Court heard oral argument on March 25, 2026."
)

body(
    "Following post-trial motions — including Pak's motion for new trial and motion to vacate — "
    "the Court entered a Final Amended Judgment on April 9, 2026, which largely conformed to "
    "the January 12, 2026 judgment."
)

body(
    "On April 30, 2026, the Court issued its Order on Attorney's Fees, awarding Plaintiffs "
    "$532,936 in attorney's fees and $9,973 in costs, for a total award of $542,909. Notice of "
    "Entry of the Order was served on May 4, 2026, but only on the following email addresses "
    "(Robinson Decl., ¶ [___]; Ex. A):"
)

body("    \"White, Andrew M.\" <amwhiteesq@aol.com>", indent=0.5)
body("    \"Welch, David Ryan\" <litigation@enso.law>", indent=0.5)
body("    \"Yee, Steven R.\" <sry@yeelaw.com>", indent=0.5)

body(
    "Pak's counsel was not included on the service list, despite having consistently provided "
    "the same service email address in all filings throughout the litigation. (Robinson Decl., "
    "¶ [___]; Ex. A.) Because Pak's counsel was not served with Notice of Entry of the Order, "
    "the statutory period for filing a motion for reconsideration under Code of Civil Procedure "
    "section 1008, subdivision (a) has not begun to run. In any event, even absent this defect, "
    "this Motion is timely filed within ten days of the clerk's service of the Order on the "
    "other parties."
)

body(
    "The concurrently filed Declaration of Pak's counsel, Taurin J. Robinson, sets forth the "
    "new circumstances and facts arising from the April 30, 2026 Order that form the basis of "
    "this Motion, in compliance with the affidavit requirement of Code of Civil Procedure "
    "section 1008, subdivision (a)."
)

# ── LEGAL STANDARD ─────────────────────────────────────────────────────────────
heading("LEGAL STANDARD", center=True)

subheading("A.", letter=None)
body(
    "A. Reconsideration Under Code of Civil Procedure Section 1008(a)",
    bold=True
)

body(
    "Code of Civil Procedure section 1008, subdivision (a), permits a party to move for "
    "reconsideration within ten days after service of written notice of entry of the order, "
    "based on 'new or different facts, circumstances, or law.' The statute allows trial courts "
    "to correct errors without requiring appellate intervention. (Gilberd v. AC Transit (1995) "
    "32 Cal.App.4th 1494, 1500; Garcia v. Hejmadi (1997) 58 Cal.App.4th 674, 690.)"
)

body(
    "'New or different facts' include circumstances that arise from the order itself, where the "
    "error could not have been identified or argued until the court issued its ruling. (Even "
    "Zohar Construction & Remodeling, Inc. v. Bellaire Townhouses, LLC (2015) 61 Cal.4th 830, "
    "839.) The inconsistencies and methodological errors identified in this Motion became "
    "apparent only after a line-by-line comparison of the April 30, 2026 Order against the "
    "billing record — an analysis that necessarily could occur only after the Order issued."
)

body(
    "Section 1008(a) also requires a declaration identifying the prior application and the new "
    "or different facts or circumstances supporting reconsideration. The concurrently filed "
    "Declaration of Taurin J. Robinson satisfies this requirement."
)

body(
    "B. Lodestar Method and the Requirement of Consistent Application",
    bold=True
)

body(
    "California courts calculate fee awards using the lodestar method: the reasonable hours "
    "expended multiplied by the reasonable hourly rate. (Ketchum v. Moses (2001) 24 Cal.4th "
    "1122, 1131-1132.) After determining the lodestar, the court may adjust it based on "
    "identified factors, but the methodology must be applied consistently and rationally across "
    "the entire fee record. (Id. at 1134-1135.)"
)

body(
    "Once a court adopts a reduction methodology — whether categorical exclusions, percentage "
    "reductions, or revaluation of time — it must apply that methodology uniformly to comparable "
    "entries and timekeepers. (Graciano v. Robinson Ford Sales, Inc. (2006) 144 Cal.App.4th "
    "140, 154; Vo v. Las Virgenes Municipal Water Dist. (2000) 79 Cal.App.4th 440, 446.) "
    "Applying a reduction to some entries but not to substantively identical entries, without "
    "articulating a principled basis for the distinction, constitutes an arbitrary lodestar "
    "analysis. Internal inconsistency in a fee award is an abuse of discretion. (PLCM Group v. "
    "Drexler (2000) 22 Cal.4th 1084, 1096.)"
)

body(
    "A motion for reconsideration is the appropriate mechanism to correct such errors "
    "efficiently before appellate review becomes necessary. Where the record does not reveal a "
    "consistent basis for materially different treatment of similar billing entries, the court "
    "may revisit the allocation on reconsideration to ensure the methodology is applied "
    "uniformly. (Gorman v. Tassajara Development Corp. (2009) 178 Cal.App.4th 44, 101.)"
)

# ── ARGUMENT ───────────────────────────────────────────────────────────────────
heading("ARGUMENT", center=True)

# ─────────────────────────────────────────────────────────────────────────────
heading(
    "THE COURT'S NR (NOT REASONABLY NECESSARY) EXCLUSION MUST BE APPLIED "
    "CONSISTENTLY ACROSS ALL TIMEKEEPERS FOR THE SAME TRIAL DATES",
    roman="I", center=False
)

body(
    "The April 30, 2026 Order identified a category of billing entries designated 'NR — not "
    "reasonably necessary' and excluded those entries entirely from the lodestar calculation. "
    "The Order applied this designation to certain trial-preparation and trial-attendance "
    "entries, explaining that the presence of multiple attorneys at the same trial session "
    "generated duplicative and unnecessary time charges that the opposing party should not be "
    "required to bear. The Court expressly found that seven specific trial dates — "
    "February 10, February 28, March 4, March 5, March 7, March 13, and March 14 — were "
    "'clearly unrelated' to the Elder Financial Abuse claim and would therefore be excluded. "
    "(Order at pp. 6-7.)"
)

body(
    "A. The NR Exclusion Was Not Applied Consistently for March 4, March 5, and March 14",
    bold=True
)

body(
    "Despite the Court's express finding that March 4, March 5, and March 14, 2025 were "
    "unrelated to the elder-abuse claim, the Court's line-item reductions did not consistently "
    "exclude both attorneys' time for those dates:"
)

bullet("For March 4, the Court reduced Jonathan White's entry from 12 hours to 2 hours rather "
       "than striking it entirely, even though the Court held this date was non-compensable.")
bullet("For March 5, the Court made no reduction to Jonathan White's time at all.")
bullet("Andrew White's entries for March 5 and March 14 remain untouched in Exhibits A and B.")

body(
    "Because the Court has already made the factual finding that these dates are "
    "non-compensable, all billing entries by both attorneys for those dates must be removed to "
    "align with the Order's stated findings."
)

body(
    "B. The Inconsistent Application of the NR Exclusion Requires Correction",
    bold=True
)

body(
    "The NR exclusion was applied to some timekeepers' trial-day billing entries but not to "
    "substantively comparable entries by the same or other timekeepers for the identical court "
    "sessions. The underlying billing records reflect that both Andrew White and Jonathan White "
    "attended the identical trial proceedings on the excluded dates and performed substantively "
    "equivalent tasks during those sessions."
)

body(
    "The Order provides no explanation for treating one timekeeper's trial-attendance entries "
    "as 'not reasonably necessary' while simultaneously retaining the other timekeeper's "
    "trial-attendance entries for the identical sessions as reasonable and compensable. If the "
    "presence of multiple attorneys at those trial sessions was not reasonably necessary — the "
    "Court's stated ground for the NR exclusion — that conclusion applies with equal logical "
    "force to all attorneys who attended those sessions. (Ketchum v. Moses, supra, 24 Cal.4th "
    "at p. 1132 ['reasonableness of hours expended is evaluated by reference to each category "
    "of work, and consistent principles must govern that evaluation'].)"
)

body(
    "Correcting this inconsistency requires striking all billing entries for both attorneys on "
    "March 4, March 5, and March 14, consistent with the Court's stated findings that those "
    "dates were non-compensable. Additional reduction required: $[_________] "
    "(see Robinson Decl., Ex. [___])."
)

# ─────────────────────────────────────────────────────────────────────────────
heading(
    "ANDREW WHITE'S BILLING ENTRIES REQUIRE ADDITIONAL REDUCTIONS UNDER THE "
    "COURT'S OWN METHODOLOGY",
    roman="II", center=False
)

body(
    "A. Additional Unrelated-Matter (NR) Reductions Are Required — Additional Reduction: "
    "$191,010",
    bold=True
)

body(
    "The Court's April 30 Order expressly held that work relating to matters other than the "
    "elder financial abuse claim against Pak — including the Norris and Robinson disputes, "
    "internal BrownLaw matters, and work on unsubstantiated theories — was not compensable "
    "under Welfare & Institutions Code section 15657.5(a). The Court found these issues were "
    "'sufficiently discreet that [they] can be segregated' and 'not compensable here.' (Order "
    "at pp. 5-6.)"
)

body(
    "A detailed review of Andrew White's billing records reveals that 310 hours, totaling "
    "$232,500, fall squarely within the categories the Court already held to be "
    "non-compensable. These hours consist of work expressly tied to:"
)

unrelated_categories = [
    "Pre-litigation BrownLaw internal disputes",
    "Termination-letter drafting and office-security events",
    "LAPD involvement and digital-forensics coordination",
    "CCP 36 trial-preference motion practice",
    "Protective-order motion practice",
    "Case-management conferences",
    "Discovery and motion practice directed at Norris, Norris Corp., and Robinson",
    "Work on the First Amended Complaint",
    "Work relating to ethics experts Buckner and Osman",
    "Work on the 'stolen clients' theory and review of Norris-firm computers and files",
]
for cat in unrelated_categories:
    bullet(cat)

body(
    "Despite the Court's finding, the Court did not exclude the majority of Andrew White's time "
    "falling within these same categories. After the Court's partial reductions, Andrew's "
    "unrelated-time subtotal remained $198,675 — a figure fundamentally inconsistent with the "
    "Court's apportionment findings. Applying the Court's own methodology consistently, the "
    "compensable subtotal for this category is $9,581.25, which after the Court's 20% lodestar "
    "reduction yields a final compensable amount of $7,665.00. Correcting this inconsistency "
    "requires a further reduction of $191,010 from Andrew White's unrelated-time category "
    "alone. (Bell v. Vista Unified School Dist. (2000) 82 Cal.App.4th 672; Christian Research "
    "Institute v. Alnor (2008) 165 Cal.App.4th 1315.)"
)

body(
    "B. Additional Duplicative-Work (DW) Reductions Are Required — Additional Reduction: "
    "$14,700",
    bold=True
)

body(
    "The Court's April 30 ruling correctly recognized substantial duplicative work between "
    "Andrew White and Jonathan White, and stated a clear rule: 'Many of the tasks for which "
    "plaintiffs are seeking compensation easily could be performed by an attorney with seven "
    "years' experience (Jonathan), and did not require two experienced attorneys to handle the "
    "same tasks. Accordingly, the Court has stricken from Andrew White's billing records all "
    "duplicative tasks which were also performed by Jonathan, except where two attorneys were "
    "reasonably required (trial or important hearings).' (Order at p. 8.)"
)

body(
    "Pak's review of the billing records shows that 20.6 hours of Andrew White's time — "
    "totaling $15,450 billed — fall squarely within the Court's duplicative-work category but "
    "were not removed. These entries include joint strategy conferences, joint document review, "
    "joint discovery review, joint drafting sessions, joint preparation for hearings and "
    "depositions, joint review of client materials, declarations, and discovery responses, and "
    "joint preparation of motions and supporting papers. Each of these is exactly the type of "
    "work the Court found Jonathan White was fully capable of performing alone."
)

body(
    "Applying the Court's stated rule consistently requires excluding 20.6 hours / $15,450 "
    "billed (/$14,700 after the Court's other reductions), yielding a proposed remaining "
    "subtotal for this category of $300 ($240 after the Court's 20% lodestar reduction). "
    "This correction is required under Gorman v. Tassajara Development Corp., supra, "
    "178 Cal.App.4th at 101 [fee order must apply methodology uniformly]."
)

body(
    "C. Additional Clerical/Administrative (CL) Reductions Are Required — Additional "
    "Reduction: $2,100",
    bold=True
)

body(
    "In its April 30 ruling, the Court expressly held that clerical and administrative tasks "
    "are not compensable at any rate, citing Christian Research Institute v. Alnor and Bell v. "
    "Vista Unified School District. The Court applied this rule to several of Jonathan White's "
    "entries, marking them CL and striking them in full. However, substantial clerical and "
    "administrative time billed by Andrew White was not marked CL, even though the tasks fall "
    "squarely within the Court's own definition of non-compensable clerical work."
)

body(
    "Pak's audit identifies 25.5 hours of Andrew White's time that fall within the Court's CL "
    "category, including reviewing invoices, assembling transcripts, preparing exhibit "
    "spreadsheets, reviewing case-management statements, reviewing amended minute orders, and "
    "coordinating expert payments. Under the Court's own rule, these entries must be excluded. "
    "Consistent application of the CL rule requires an additional $2,100 reduction to "
    "Andrew White's lodestar. (Christian Research Institute v. Alnor, supra, 165 Cal.App.4th "
    "at 1326 ['Clerical tasks are not compensable at attorney rates and are generally not "
    "compensable at all'].)"
)

body(
    "D. MSA-Related Billing Entries Must Be Reduced for Internal Consistency — Additional "
    "Reduction: $37,575",
    bold=True
)

body(
    "The Court did not find that the MSA work was unrelated to the elder-abuse claim. Rather, "
    "the Court found that the MSA was egregiously overbilled given the simplicity of the issue "
    "presented. As the Court explained, the motion 'presented a single issue — did Pak, as "
    "Baird Brown's office manager, owe the firm and Brown a fiduciary duty,' the memorandum was "
    "'only 10 pages in length,' supported by a 'straightforward' Separate Statement containing "
    "'only 16 undisputed facts,' with evidence consisting solely of two declarations and a small "
    "set of checks and interrogatory responses, and a legal question that was 'hardly novel.' "
    "Despite this, Jonathan White billed 54.2 hours and Andrew White billed 42.9 hours — a "
    "combined 97.1 hours — for a motion the Court found so overbilled that it described the "
    "$50,000 billed as 'shocking.' (Order at p. 8.) The Court imposed a top-line reduction of "
    "$26,435."
)

body(
    "However, the Court's own findings demonstrate that the overbilling was not limited to the "
    "initial MSA filing. The same excessive, duplicative, and block-billed practices continued "
    "throughout the reply briefing, evidentiary objections, separate-statement responses, and "
    "hearing preparation. Maintaining internal consistency with the Court's stated methodology "
    "— particularly its findings regarding duplication, block billing, and excessive time — "
    "requires proportionate reductions to the remaining MSA entries."
)

body(
    "Andrew White billed $64,050 for MSA-related work. After the Court's cuts of $51,825, the "
    "proposed subtotal is $37,575, yielding a final compensable amount of $16,010.80 after the "
    "Court's 20% lodestar reduction. This represents an additional reduction of $37,575."
)

body(
    "E. Additional Norris/Robinson-Related (DN/TR) Reductions Are Required — Additional "
    "Reduction: $59,006.25",
    bold=True
)

body(
    "The Court's April 30 Order expressly held that work relating to the Norris defendants and "
    "to Mr. Robinson was not compensable under Welfare & Institutions Code section 15657.5(a). "
    "The Court found that Brown did not prevail on the elder-abuse claims against Norris or "
    "Norris Corp.; that 'most of the case was devoted to trying to prove that Norris and his "
    "firm had stolen Brown's clients,' a theory the Court rejected; that substantial time was "
    "spent defending against Norris's counterclaims; that time spent litigating the claims "
    "involving Robinson was unrelated to the elder-abuse claim against Pak; and that these "
    "issues were 'sufficiently discreet that [they] can be segregated' and 'are not compensable "
    "here.' (Order at pp. 5-6.) The Court further held that time spent on the seizure and "
    "review of the Norris Law Firm computers and files was non-compensable because Brown's own "
    "nondisclosure caused that litigation activity."
)

body(
    "Despite these clear findings, a substantial volume of Andrew White's time that falls "
    "squarely within the Court's non-compensable categories was not marked NR, DW, or otherwise "
    "reduced. Pak's line-by-line review identifies additional entries expressly tied to:"
)

tr_categories = [
    "Norris-related discovery disputes",
    "Norris-related motions (including sanctions, terminating sanctions, and motions in limine)",
    "Preparation for and attendance at Norris depositions",
    "Work on the failed claims against Robinson",
    "Work on ethics-expert issues tied to the Norris/Robinson theories",
    "Trial preparation and trial work directed at Norris-related issues",
    "Closing-brief work addressing Norris-related issues",
    "Post-trial work on objections and judgment drafting involving Norris issues",
]
for cat in tr_categories:
    bullet(cat)

body(
    "Pak's audit identifies 230.6 hours billed for work that is expressly Norris- or "
    "Robinson-related. After applying the same methodology the Court used in Exhibit B "
    "(including partial reductions where appropriate), the reasonable proposed amount for "
    "exclusion is $59,006.25 in additional DN/TR-related reductions ($47,205.00 after the "
    "Court's 20% lodestar reduction). These reductions are fully consistent with the Court's "
    "own apportionment findings and simply extend the Court's methodology to entries that were "
    "not previously marked."
)

body(
    "F. Andrew White's Post-Judgment and Fees-on-Fees Entries Must Be Reduced to $0 "
    "— Reduction of $92,925",
    bold=True
)

body(
    "Andrew White billed 127.0 hours totaling $95,250.00 for post-judgment tasks, including "
    "work on the new trial motion, work on the motion to vacate, work on the fee motion itself "
    "(fees-on-fees), work on successor-in-interest stipulations, and post-judgment enforcement "
    "and procedural matters. After the Court's line-item reductions, the Court allowed "
    "$92,925.00 in post-judgment fees for Andrew White."
)

body(
    "The Court then applied the following across-the-board reductions to Andrew White's overall "
    "lodestar: a $14,144.36 reduction for the fee motion, and a $15,247.66 reduction for the "
    "new trial motion opposition. The Court made no across-the-board reduction — and no "
    "line-item reduction — for Andrew White's work on the motion to vacate, even though the "
    "motion to vacate is substantively indistinguishable from the new trial motion in purpose, "
    "posture, and timing. This selective treatment is internally inconsistent."
)

body(
    "More fundamentally, none of Andrew White's post-judgment work falls within the scope of "
    "Welfare & Institutions Code section 15657.5(a), which authorizes only those fees "
    "'reasonably incurred' in proving financial elder abuse. None of the post-judgment entries "
    "relate to proving elder abuse, to trial preparation or trial, or to any compensable claim. "
    "All of it is purely collateral post-judgment procedural work expressly outside the scope "
    "of section 15657.5(a). The only reasonable lodestar for Andrew White's post-judgment work "
    "is $0.00, and the full $92,925.00 must be excluded. (See also Argument Section VI "
    "regarding the lack of statutory authorization for fees-on-fees under section 15657.5.)"
)

# ─────────────────────────────────────────────────────────────────────────────
heading(
    "JONATHAN WHITE'S BILLING ENTRIES REQUIRE ADDITIONAL REDUCTIONS UNDER THE "
    "COURT'S OWN METHODOLOGY",
    roman="III", center=False
)

body(
    "G. Jonathan White's Post-Judgment and Fees-on-Fees Entries Must Likewise Be Reduced "
    "to $0 — Reduction of $20,880",
    bold=True
)

body(
    "Jonathan White billed 69.6 hours totaling $20,880.00 for the same categories of "
    "post-judgment work: the fee motion, the new trial motion opposition, the motion to vacate, "
    "and related post-judgment procedural matters. The Court made line-item cuts to some of "
    "these entries, and separately applied across-the-board reductions of $3,515.64 for the "
    "fee motion and $4,884.36 for the new trial motion opposition. The Court made no reduction "
    "at all — neither line-item nor across-the-board — for Jonathan White's work on the motion "
    "to vacate."
)

body(
    "As with Andrew White, none of Jonathan White's post-judgment work falls within the "
    "compensable scope of Welfare & Institutions Code section 15657.5(a). The basis for "
    "exclusion is identical: post-judgment procedural motions do not advance the elder "
    "financial abuse claim, do not relate to liability or damages, and are not authorized by "
    "the statute's text. The only reasonable lodestar for Jonathan White's post-judgment work "
    "is $0.00, and the entire $20,880.00 must be excluded."
)

body(
    "H. The Court's Inconsistent Treatment of Post-Judgment Work Across All Three Motions",
    bold=True
)

body(
    "The following table illustrates the Court's inconsistent treatment of the three "
    "post-judgment motions for both attorneys:"
)

# Comparison table for post-judgment inconsistency
pj_tbl = doc.add_table(rows=5, cols=4)
pj_tbl.style = "Table Grid"
pj_headers = ["Post-Judgment Motion", "Andrew White — Court Treatment",
               "Jonathan White — Court Treatment", "Pak's Position"]
for i, h in enumerate(pj_headers):
    cell = pj_tbl.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)

pj_rows = [
    ("Fee Motion",
     "Across-the-board reduction: $14,144.36",
     "Across-the-board reduction: $3,515.64",
     "Non-compensable; $0"),
    ("New Trial Motion Opposition",
     "Across-the-board reduction: $15,247.66",
     "Across-the-board reduction: $4,884.36",
     "Non-compensable; $0"),
    ("Motion to Vacate",
     "No reduction",
     "No reduction",
     "Non-compensable; $0"),
    ("Total Post-Judgment Billed / Allowed",
     "$95,250 billed / $92,925 allowed",
     "$20,880 billed",
     "$0 for both"),
]
for row_data in pj_rows:
    r = pj_tbl.add_row()
    for i, (cell, text) in enumerate(zip(r.cells, row_data)):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
doc.add_paragraph()

body(
    "This table demonstrates the core inconsistency: the Court treated the new trial motion "
    "and fee motion as 'excessive' and applied reductions, while leaving the motion to vacate "
    "— an indistinguishable post-judgment procedural motion — entirely untouched. Under the "
    "Court's own apportionment logic, if the new trial opposition does not advance the elder "
    "abuse claim, neither does the motion to vacate. If the Court's approach on reconsideration "
    "is to reduce rather than exclude all post-judgment work, then at minimum the motion to "
    "vacate must receive the same across-the-board reduction the Court applied to the new trial "
    "motion. But the more principled and legally correct result — consistent with section "
    "15657.5(a)'s text and Conservatorship of McQueen — is to exclude all post-judgment work "
    "in its entirety."
)

body(
    "A. Additional Norris/Robinson-Related (DN/TR) Reductions Are Required",
    bold=True
)

body(
    "A substantial portion of Jonathan White's billed time relates not to work performed for "
    "Brown but to work attributable to Donald Norris and Taurin Robinson. The billing records "
    "show that from the outset of the case, Mr. White devoted significant time to reviewing "
    "documents concerning Norris and Robinson, researching claims against them, drafting the "
    "complaint and related exhibits, preparing for settlement discussions with their counsel, "
    "and developing litigation strategy directed at their conduct. This pattern continues "
    "throughout discovery, trial preparation, trial, and post-trial briefing, where large "
    "blocks of time were spent addressing Norris-specific discovery failures, expert issues, "
    "trial briefs, objections to the proposed statement of decision, and closing arguments. "
    "Under the Court's apportionment ruling, this work must be reduced or excluded because it "
    "does not relate to Brown's claims against Pak."
)

body(
    "The Court's line-item reductions did not address this DN/TR apportionment problem, even "
    "though the billing descriptions make the non-Brown-related nature of the work "
    "unmistakable. When DN/TR apportionment is applied consistently, the total for this "
    "category drops from $79,605 (after the Court's cuts) to $38,872.50, and to $31,098 after "
    "the Court's 20% lodestar reduction."
)

body(
    "B. Trial-Day NR Exclusions for March 4 and March 5 Were Not Applied",
    bold=True
)

body(
    "The Court expressly found that seven specific trial dates — February 10, February 28, "
    "March 4, March 5, March 7, March 13, and March 14 — were 'clearly unrelated' to the "
    "Elder Financial Abuse claim and would be excluded. (Order at pp. 6-7.) However, two of "
    "those dates — March 4 and March 5 — were not actually removed in the Court's line-item "
    "reductions for Jonathan White. On March 4, the Court reduced the entry from 12 hours to "
    "2 hours rather than striking it entirely; on March 5, the Court made no reduction at all. "
    "Both dates fall squarely within the seven days the Court held were non-compensable. "
    "Correcting this internal inconsistency requires striking both entries entirely, consistent "
    "with the Court's stated methodology."
)

body(
    "C. MSA-Related Billing Entries Contain Inconsistently Applied Reductions",
    bold=True
)

body(
    "The Court's treatment of the MSA-related billing contains several internal inconsistencies "
    "that require correction. The Court stated that Jonathan White billed 54.2 hours preparing "
    "the MSA and separate statement, but the billing records reflect 109.1 hours containing "
    "MSA-related work, totaling $32,730. Outside of the $7,650 reduction for duplicate work "
    "between November 13 and November 20, the Court made no reductions to the remaining MSA "
    "entries, even though seven of those entries contain clerical or paralegal components and "
    "one entry (11/6/24) pertains to Norris discovery rather than the MSA."
)

body(
    "Applying the same clerical/paralegal reductions the Court used elsewhere reduces these "
    "entries from $7,710 to $3,855, and the Norris-related entry from $1,260 to $630. "
    "Correcting these inconsistencies results in a proposed subtotal of $10,500, and $8,400 "
    "after the Court's 20% lodestar reduction."
)

body(
    "D. Mistrial-Related Entries Must Be Excluded — Reduction: $7,770",
    bold=True
)

body(
    "The Court expressly recognized that the mistrial issue arose from Plaintiffs' counsel's "
    "own conduct and therefore fell squarely within the rule of Bell v. Vista Unified School "
    "District and Christian Research Institute v. Alnor that 'time spent litigating issues "
    "caused by counsel's own conflicts is non-compensable.' Yet the Court made no reductions "
    "to the 26.0 hours ($7,770) Jonathan White billed for drafting, revising, filing, and "
    "arguing the opposition to the mistrial motion, even though every entry in this block of "
    "time is exclusively mistrial-related. Because this work was necessitated solely by "
    "counsel's own conflict and did not advance the elder financial abuse claim, it must be "
    "excluded in full under the Court's own stated methodology."
)

body(
    "E. Unrelated-Matter (NR) Reductions Must Be Applied Consistently — Additional Reduction: "
    "$71,720",
    bold=True
)

body(
    "The billing records contain 343.0 hours ($101,730) of work wholly unrelated to the Elder "
    "Financial Abuse claim that should have been excluded under the Court's own methodology. "
    "The Court made some partial cuts, mainly under the duplicate-work category, but still left "
    "$71,720 compensable. This category includes pre-litigation internal BrownLaw disputes, "
    "termination-letter drafting, office-security and LAPD involvement, digital forensics, "
    "CCP 36 trial-preference work, protective-order motion practice, case-management "
    "conferences, and extensive discovery and motion practice directed at Norris and Pak's "
    "business disputes. None of this work advanced the EFA claim, and much of it was caused "
    "by counsel's own conflicts — categories the Court expressly held were non-compensable "
    "under Bell and Christian Research Institute. Correcting this inconsistency requires "
    "excluding this entire category, reducing the $71,720 billed for this work to $0."
)

body(
    "F. Conflict-Created Mistrial Work Must Be Excluded Consistently for Both Attorneys — "
    "Additional Reduction: $35,595 Total",
    bold=True
)

body(
    "The billing entries associated with the mistrial-related conflict work fall squarely "
    "within the category of non-compensable hours because they were generated entirely by "
    "Plaintiffs' counsel's own conduct — specifically, placing their client in a position "
    "where he testified against his own interests regarding the Pixley Forensics engagement "
    "and the October 2023 security operation. As the Court itself recognized, citing Bell and "
    "Christian Research Institute, time spent litigating issues caused by counsel's own "
    "conflicts cannot be shifted to the opposing party. The Court nevertheless awarded $27,825 "
    "to Andrew White and $7,770 to Jonathan White for this conflict-created work. Because none "
    "of these hours advanced the elder financial abuse claim — and were instead necessitated by "
    "counsel's own ethical entanglements — the full $35,595 must be deducted from the fee award."
)

# ─────────────────────────────────────────────────────────────────────────────
heading(
    "THE COURT'S DUPLICATION RULE WAS APPLIED INCONSISTENTLY AND MUST BE RECONCILED",
    roman="IV", center=False
)

body(
    "A. The Order States Two Irreconcilable Standards for Duplicative Work",
    bold=True
)

body(
    "The Court articulated a clear duplication principle in the body of the Order: 'Many of "
    "the tasks for which plaintiffs are seeking compensation easily could be performed by an "
    "attorney with seven years' experience (Jonathan), and did not require two experienced "
    "attorneys to handle the same tasks. Accordingly, the Court has stricken from Andrew "
    "White's billing records all duplicative tasks which were also performed by Jonathan, "
    "except where two attorneys were reasonably required (trial or important hearings).' "
    "(Order at p. 8.) This rule is internally coherent: Jonathan is capable of performing "
    "routine and mid-level tasks; therefore, when both attorneys perform the same task, "
    "Andrew's duplicative time is cut."
)

body(
    "However, Footnote 2 of the Order introduced a different and contradictory rule: 'Because "
    "Andrew White was lead counsel, the Court has cut all time spent at the less critical "
    "hearings by Jonathan White as unreasonably incurred.' These two standards cannot be "
    "reconciled. The main text adopts a staffing-efficiency rule (Jonathan is capable, so "
    "Andrew's duplicative time is cut); the footnote adopts a hierarchy-based rule (Andrew is "
    "lead counsel, so Jonathan's duplicative time is cut). Both cannot be true simultaneously. "
    "Because the Order does not identify when one rule applies and when the other applies, "
    "identical categories of duplicative work were treated differently without a consistent "
    "rationale. (Gorman v. Tassajara Dev. Corp., supra, 178 Cal.App.4th at 101.)"
)

body(
    "B. Three Specific NR/DW Classification Errors Require Correction",
    bold=True
)

body(
    "(a) October 3, 2024 — Opposition to Demurrer. The Court marked Jonathan White's entry "
    "for this task as both NR and DW, but Andrew White's billing for the same task remains "
    "fully compensable. If the task was both non-compensable and duplicative for Jonathan, "
    "then the same reasoning necessarily applies to Andrew's entry for the identical work. "
    "Under the Court's own stated principles, both attorneys' time should be stricken for this "
    "task (total: $4,170)."
)

body(
    "(b) December 11, 2024 — Deposition of Robinson. The Court marked Jonathan White's entry "
    "as DW but did not designate it NR, even though the task was entirely unrelated to Pak or "
    "the EFA claim. The entry includes taking Robinson's deposition, travel to and from Culver "
    "City, legal research for a motion in limine regarding reasonable accommodations, and "
    "conferring with Andrew White. None of these tasks relate to Pak or the EFA issue. Under "
    "the Court's own reasoning, this entry should have been marked NR for both attorneys "
    "(total: $8,400)."
)

body(
    "(c) December 15, 2024 — Motion in Limine No. 2. The Court marked Jonathan White's entry "
    "as NR and DW, but Andrew White's corresponding entry (performed the day before) was not "
    "reduced. Because the Court found the task both non-compensable and duplicative as to "
    "Jonathan, the same classification necessarily applies to Andrew's entry for the same work. "
    "Both attorneys' time should be stricken (total: $2,955)."
)

body(
    "These three entries require a further net reduction in attorney's fees of $10,875."
)

body(
    "C. Jonathan White's Billing Records Contain 29 Arithmetic Errors",
    bold=True
)

body(
    "The Court's identification of a single arithmetic error in Jonathan White's billing entry "
    "for March 12, 2025 (1.4 hours billed at $300/hour but charged at $3,000) prompted Pak's "
    "counsel to recalculate all of Jonathan's hours and billed amounts. Upon review, 29 entries "
    "— 6.68% of Jonathan's total time entries — contain arithmetic errors where the charged "
    "amount does not match the hours multiplied by the stated hourly rate. The Court identified "
    "only one such error."
)

body(
    "Because the Court's stated line-item reductions for Jonathan ($117,555) are derived from "
    "these billing records, and because the Court's own NR/DW markings — when applied to the "
    "same entries — total $254,675, the discrepancy appears to stem in part from reliance on "
    "arithmetically inaccurate billing data. Pak's counsel's recalculation (attached as "
    "Exhibit [___]) corrects both overcharges and undercharges, resulting in a net additional "
    "$1,020 that should have been deducted. The presence of pervasive arithmetic errors "
    "— combined with the other inconsistencies identified in this Motion — significantly "
    "undermines the credibility and reliability of Plaintiffs' fee request and the mathematical "
    "foundation of the Court's ruling. (Gorman v. Tassajara Dev. Corp., supra, 178 Cal.App.4th "
    "at 101.)"
)

# ─────────────────────────────────────────────────────────────────────────────
heading(
    "THE COURT'S CLERICAL AND PARALEGAL REDUCTION MUST BE APPLIED UNIFORMLY",
    roman="V", center=False
)

body(
    "The Court correctly recognized that clerical and administrative tasks are not compensable "
    "at attorney rates, citing Christian Research Institute v. Alnor and Bell v. Vista Unified "
    "School District. California law is unequivocal: tasks such as calendaring, scheduling, "
    "coordinating service, assembling exhibits, cataloging documents, and communicating with "
    "vendors are clerical in nature and must either be excluded or compensated at a reduced "
    "paralegal rate. (Heritage Pacific Financial v. Monroy (2013) 215 Cal.App.4th 972.)"
)

body(
    "However, the Court's application of this rule was inconsistent. While some clerical "
    "entries were struck or reduced, many others containing identical task descriptions were "
    "left untouched. Entries involving service of process, document cataloging, exhibit "
    "assembly, calendaring, and communications with ADR services or court reporters were "
    "treated differently despite being substantively indistinguishable. Under Gorman v. "
    "Tassajara, a fee award must apply its methodology uniformly; internal inconsistency is "
    "itself grounds for reconsideration."
)

body(
    "The CL-related time for Jonathan White is discrete, segregable, and easily identifiable. "
    "Applying the Court's own rule consistently yields a corrected total of $1,932, reflecting "
    "a consistent application of the Court's own methodology. This requires an additional "
    "reduction of $[_________] (see Robinson Decl., Ex. [___])."
)

body(
    "For Andrew White, as detailed in Section II.C above, consistent application of the CL "
    "rule requires an additional $2,100 reduction."
)

body(
    "Because these tasks are non-legal and not intertwined with the elder-abuse claim, they "
    "must be excluded or revalued at a clerical/paralegal rate. Reconsideration is required to "
    "eliminate the internal inconsistency and ensure that clerical work is not compensated at "
    "attorney rates."
)

# ─────────────────────────────────────────────────────────────────────────────
heading(
    "THE FEES-ON-FEES AWARD EXCEEDS THE STATUTORY AUTHORIZATION UNDER WELFARE "
    "AND INSTITUTIONS CODE SECTION 15657.5",
    roman="VI", center=False
)

body(
    "A. Statutory Framework",
    bold=True
)

body(
    "Welfare and Institutions Code section 15657.5, subdivision (a), provides that a prevailing "
    "plaintiff in a financial elder abuse action is entitled to recover 'reasonable attorney's "
    "fees and costs.' The statute's remedial purpose is served by fee-shifting on the underlying "
    "merits litigation. The question presented here is whether that purpose extends to fees "
    "incurred in litigating the fee application itself — so-called 'fees-on-fees' — and to fees "
    "incurred on post-judgment procedural motions including the new trial motion and motion to "
    "vacate."
)

body(
    "B. Section 15657.5 Does Not Authorize Post-Judgment Fees or Fees-on-Fees",
    bold=True
)

body(
    "California applies a rule of strict construction to fee-shifting statutes: courts may not "
    "expand the scope of a statutory fee award beyond what the text of the statute expressly or "
    "necessarily authorizes. (Hanna v. Mercedes-Benz USA, LLC (2019) 36 Cal.App.5th 493, 510; "
    "Graham v. DaimlerChrysler Corp. (2004) 34 Cal.4th 553, 580.) Unlike fee-shifting statutes "
    "that expressly authorize recovery of fees incurred in litigating fee applications, Welfare "
    "and Institutions Code section 15657.5 contains no express authorization for fees-on-fees "
    "or post-judgment motion fees. Under California's strict construction rule, this silence "
    "counsels against expansion of the award."
)

body(
    "The Court's reliance on Ketchum v. Moses and Serrano IV is misplaced. Both cases involved "
    "fee-shifting statutes with broad remedial purposes and expansive language: Ketchum involved "
    "the anti-SLAPP statute (Code Civ. Proc. § 425.16(c)), and Serrano IV involved the private "
    "attorney general statute (Code Civ. Proc. § 1021.5). Both statutes expressly authorize "
    "fees for enforcement and fee litigation. Section 15657.5(a) is materially narrower: it "
    "mandates an award of 'reasonable attorney's fees and costs' only upon proof of financial "
    "elder abuse. Nothing in its text extends that entitlement to post-judgment litigation."
)

body(
    "The Supreme Court's decision in Conservatorship of McQueen reinforces this conclusion. "
    "While McQueen addressed appellate and enforcement fees, its core holding is that "
    "section 15657.5(a) does not authorize post-judgment fees unless the statute expressly "
    "says so. The Court emphasized that fee-shifting statutes must be read narrowly and that "
    "post-judgment fees cannot be implied. That reasoning applies with even greater force here: "
    "motions for new trial, motions to vacate, and fee motions are not enforcement proceedings, "
    "nor are they part of proving elder abuse. They are collateral, post-judgment procedural "
    "matters."
)

body(
    "The Court's reliance on Pellegrino is also misplaced. Pellegrino held that fees-on-fees "
    "may be recoverable only when the underlying statute authorizes them. Section 15657.5(a) "
    "does not. The Court cited Pellegrino only to reject the application of multipliers but "
    "ignored its threshold rule: fees-on-fees require statutory authorization."
)

body(
    "C. The Court's Apportionment Findings Independently Require Exclusion of Post-Judgment Work",
    bold=True
)

body(
    "The Court's own apportionment findings further undermine the award. The Court held that "
    "fees are recoverable only for work necessary to prove the elder abuse claim, and that all "
    "other work — including work on unrelated claims, cross-claims, discovery disputes, and "
    "procedural motions — must be excluded. Post-judgment work is not part of proving elder "
    "abuse. It does not relate to liability, damages, or statutory elements. It is purely "
    "procedural."
)

body(
    "Yet the Court applied this apportionment principle inconsistently: it reduced the time "
    "spent on the new trial opposition and the fee motion as 'excessive,' but did not exclude "
    "them as non-compensable. And it made no reduction at all for the motion to vacate, even "
    "though it is indistinguishable from the new trial motion in purpose, posture, and timing. "
    "If the Court's apportionment rule is applied consistently, all post-judgment work — "
    "including the motion to vacate — must be excluded."
)

body(
    "D. Requested Relief",
    bold=True
)

body(
    "The Court should strike the fees-on-fees and post-judgment component of the award in its "
    "entirety on the grounds that Welfare and Institutions Code section 15657.5 does not "
    "authorize it. The total post-judgment fees that must be excluded are:"
)
bullet("Andrew White: $92,925.00 (127.0 hours / $95,250 billed; $92,925 after Court's "
       "line-item cuts; proposed $0)")
bullet("Jonathan White: $20,880.00 (69.6 hours billed; proposed $0)")
bullet("Combined post-judgment reduction required: $113,805.00")

body(
    "In the alternative, if the Court declines to exclude the post-judgment fees entirely, "
    "it must at minimum apply the same across-the-board reductions to the motion to vacate "
    "that it applied to the new trial motion — a reduction of $15,247.66 for Andrew White and "
    "$4,884.36 for Jonathan White — as the motion to vacate is substantively indistinguishable "
    "from the new trial motion and received no reduction at all."
)

# ─────────────────────────────────────────────────────────────────────────────
heading(
    "CORRECTED LODESTAR CALCULATION",
    roman="VII", center=False
)

body(
    "Pak does not ask the Court to change its methodology. She asks only that the Court's "
    "existing methodology — its line-item reductions, across-the-board category reductions, "
    "and 20% lodestar haircut — be applied consistently and with accurate arithmetic. When "
    "that is done, the corrected lodestar is reproducible and materially lower than the amount "
    "awarded."
)

body(
    "1. Jonathan White",
    bold=True
)

body(
    "The Court's line-item reductions leave a subtotal of $254,675. Applying the Court's "
    "global reductions proportionally to Jonathan's share of the work yields:"
)
bullet("MSA reduction: –$8,873.50 → $245,801.50")
bullet("Fee-motion reduction: –$3,515.64 → $242,285.86")
bullet("New-trial reduction: –$4,884.36 → $237,401.50")
bullet("20% lodestar reduction → $189,921.20")

body(
    "Applying the Court's own apportionment and non-compensability findings uniformly — "
    "including DN/TR work, unrelated matters, mistrial work, clerical tasks, and post-judgment "
    "work — yields a compensable subtotal of $59,557.50. After the Court's 20% lodestar "
    "reduction, the reasonable fee for Jonathan White is:"
)
body("$47,646.00", bold=True, center=True)

body(
    "2. Andrew White",
    bold=True
)

body(
    "The Court's line-item reductions leave a subtotal of $558,150.00. Applying the Court's "
    "global reductions proportionally to Andrew's share of the work yields:"
)
bullet("MSA reduction: –$17,561.50 → $540,588.50")
bullet("Fee-motion reduction: –$14,144.36 → $526,444.14")
bullet("New-trial reduction: –$15,247.66 → $511,196.48")
bullet("20% lodestar reduction → $408,957.18")

body(
    "Applying the Court's apportionment and non-compensability findings uniformly — including "
    "DN/TR work, unrelated matters, clerical tasks, duplicative work, and all post-judgment "
    "and fees-on-fees work — yields a compensable subtotal of $115,001.00. After the Court's "
    "20% lodestar reduction, the reasonable fee for Andrew White is:"
)
body("$92,000.80", bold=True, center=True)

body(
    "3. Post-Judgment Deductions (Both Attorneys)",
    bold=True
)

body(
    "As detailed in Sections II.F, II.G, and VI above, all post-judgment and fees-on-fees work "
    "must be excluded under Welfare & Institutions Code section 15657.5(a):"
)
bullet("Andrew White post-judgment: Court allowed $92,925.00 → Proposed $0.00 "
       "(reduction: $92,925.00)")
bullet("Jonathan White post-judgment: $20,880.00 billed → Proposed $0.00 "
       "(reduction: $20,880.00)")
bullet("Combined post-judgment reduction: $113,805.00")

body(
    "4. Remaining Time — Agreed Compensable Amount",
    bold=True
)

body(
    "The Court reduced Jonathan White's billed amount from $71,220 to $6,840 with line-item "
    "cuts. Pak agrees with this amount, provided the Court's 20% downward adjustment is "
    "applied consistently, as in all other categories, leaving a total of $5,472."
)

body(
    "4. Corrected Total Fee Award",
    bold=True
)

body(
    "The only lodestar that is mathematically consistent with the Court's methodology is:"
)

# Summary table
sum_tbl = doc.add_table(rows=4, cols=3)
sum_tbl.style = "Table Grid"
headers = ["Timekeeper", "Compensable Subtotal", "After 20% Lodestar Reduction"]
for i, h in enumerate(headers):
    cell = sum_tbl.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(11)

table_row(sum_tbl, "Jonathan White", "$59,557.50", "$47,646.00")
table_row(sum_tbl, "Andrew White", "$115,001.00", "$92,000.80")
table_row(sum_tbl, "TOTAL", "$174,558.50", "$139,646.80", bold_first=True)

doc.add_paragraph()

body(
    "This corrected total applies every reduction category the Court adopted, allocates the "
    "Court's across-the-board reductions proportionally, applies the Court's 20% lodestar "
    "haircut uniformly, and avoids the internal inconsistencies in the April 30, 2026 Order "
    "(including the stated $666,170 figure, which cannot be reproduced from the Court's own "
    "reductions). Pak therefore respectfully requests that the Court modify its fee award to "
    "$139,646.80, the only amount consistent with the Court's stated methodology and "
    "governing law."
)

# ── CONCLUSION ─────────────────────────────────────────────────────────────────
heading("CONCLUSION", center=True)

body(
    "For the foregoing reasons, Defendant Michelle Pak respectfully requests that the Court "
    "grant this Motion for Reconsideration and issue an amended order modifying the attorney's "
    "fee award as follows:"
)

relief = [
    ("I",   "Apply the NR exclusion consistently to all trial-day entries for February 10, "
             "February 28, March 4, March 5, March 7, March 13, and March 14, 2025 for both "
             "Andrew White and Jonathan White;"),
    ("II",  "Apply an additional $191,010 reduction to Andrew White's unrelated-matter "
             "entries to reflect consistent application of the Court's NR exclusion;"),
    ("III", "Apply an additional $14,700 reduction to Andrew White's duplicative entries "
             "consistent with the Court's stated staffing-efficiency rule;"),
    ("IV",  "Apply an additional $2,100 reduction to Andrew White's clerical/administrative "
             "entries consistent with the Court's CL rule applied to Jonathan White;"),
    ("V",   "Apply an additional $37,575 reduction to Andrew White's MSA-related entries "
             "for internal consistency with the Court's findings regarding excessive billing;"),
    ("VI",  "Apply an additional $59,006.25 reduction to Andrew White's Norris/Robinson-"
             "related entries consistent with the Court's apportionment findings;"),
    ("VII", "Reduce Andrew White's post-judgment and fees-on-fees award from $92,925 to $0;"),
    ("VIII","Apply a $40,732.50 additional reduction to Jonathan White's DN/TR entries;"),
    ("IX",  "Strike Jonathan White's March 4 and March 5 entries consistent with the Court's "
             "finding that those dates were non-compensable;"),
    ("X",   "Apply a reduction to Jonathan White's MSA entries consistent with the Court's "
             "methodology, yielding a compensable subtotal of $10,500 ($8,400 after lodestar "
             "reduction);"),
    ("XI",  "Exclude Jonathan White's $7,770 in mistrial-related entries as non-compensable "
             "under Bell and Christian Research Institute;"),
    ("XII", "Exclude $71,720 in Jonathan White's unrelated-matter entries;"),
    ("XIII","Reconcile the contradictory duplication standards in the Order body text and "
             "Footnote 2 and apply one consistent rule;"),
    ("XIV", "Correct the 29 arithmetic errors in Jonathan White's billing records with a net "
             "additional deduction of $1,020;"),
    ("XV",  "Strike or substantially reduce the fees-on-fees award for lack of statutory "
             "authorization under Welfare and Institutions Code section 15657.5; and"),
    ("XVI", "Modify the total fee award from $532,936 to $139,646.80, consistent with the "
             "Court's own stated methodology."),
]

for num, text in relief:
    body(f"{num}. {text}", indent=0.25)

doc.add_paragraph()
body("DATED: May 15, 2026")
doc.add_paragraph()
doc.add_paragraph()
body("By: _________________________________")
body("    Taurin J. Robinson, Esq.")
body("    Attorney for Defendant/Cross-Complainant Michelle Pak")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  DECLARATION OF TAURIN J. ROBINSON
# ══════════════════════════════════════════════════════════════════════════════

heading("DECLARATION OF TAURIN J. ROBINSON IN SUPPORT OF MOTION FOR RECONSIDERATION", center=True)

body("I, Taurin J. Robinson, declare as follows:")

decl_paras = [
    ("1.", "I am an attorney duly licensed to practice law in the State of California and "
           "am counsel of record for Defendant and Cross-Complainant Michelle Pak in the above-"
           "captioned matter. I have personal knowledge of the facts set forth in this "
           "Declaration. If called as a witness, I could and would testify competently to the "
           "following."),
    ("2.", "Prior Application. On March 3, 2026, Pak filed her Opposition to Plaintiffs' "
           "Motion for Attorney's Fees and Costs. That opposition objected to the fee request "
           "generally and challenged the reasonableness of the hours and rates claimed. "
           "The Court heard argument on March 25, 2026, and issued its Order on Attorney's "
           "Fees on April 30, 2026."),
    ("3.", "New or Different Facts and Circumstances. The following new facts and "
           "circumstances form the basis for this Motion for Reconsideration and could not "
           "have been identified, argued, or presented prior to the issuance of the April 30, "
           "2026 Order."),
    ("4.", "Defective Service of Notice of Entry. On May 4, 2026, Notice of Entry of the "
           "Order was served by email on the following addresses: "
           "amwhiteesq@aol.com (Andrew White), litigation@enso.law (David Ryan Welch), and "
           "sry@yeelaw.com (Steven R. Yee). My office was not included on the service list. "
           "I have consistently provided my service email address in all filings in this "
           "matter. A true and correct copy of the service notification and my office's "
           "service designation is attached hereto as Exhibit A."),
    ("5.", "Post-Order Line-by-Line Audit. After receiving the April 30, 2026 Order, I "
           "undertook a line-by-line comparison of the Court's reductions against Plaintiffs' "
           "billing records, which are attached to Plaintiffs' fee motion and incorporated in "
           "the Court's Exhibits A and B. That analysis could not have been conducted prior "
           "to the issuance of the Order because the specific entries the Court reduced, "
           "excluded, or retained were not known until the Order issued."),
    ("6.", "Trial-Day NR Inconsistencies. My review revealed that the Court's NR exclusion "
           "for specific trial dates was not applied consistently to both timekeepers. "
           "Specifically, March 4 and March 5, 2025 were identified by the Court as "
           "non-compensable dates, but the Court's line-item reductions did not strike all "
           "entries for both attorneys on those dates. Similarly, Andrew White's entries for "
           "March 5 and March 14, 2025 were not removed, despite the Court's express "
           "finding. The details of these inconsistencies, including the specific entries and "
           "dollar amounts, are set forth in Exhibit [___] attached hereto."),
    ("7.", "Duplication and CL Inconsistencies. My review further revealed inconsistencies "
           "in the Court's application of the duplication (DW) and clerical/administrative "
           "(CL) reduction categories, as detailed in the accompanying Memorandum of Points "
           "and Authorities. The specific entries at issue and the proposed corrections are "
           "set forth in Exhibit [___] attached hereto."),
    ("8.", "Arithmetic Errors. My recalculation of Jonathan White's billing entries revealed "
           "29 arithmetic errors — representing 6.68% of his total time entries — in which "
           "the charged amount does not match the hours multiplied by the stated hourly rate. "
           "The Court identified only one such error (the March 12, 2025 entry). A true and "
           "correct copy of my recalculation is attached hereto as Exhibit [___]."),
    ("9.", "Unrelated-Matter and Apportionment Inconsistencies. My review identified "
           "additional billing entries for Andrew White and Jonathan White that fall squarely "
           "within the categories the Court held to be non-compensable — including "
           "Norris/Robinson-related work, internal BrownLaw disputes, and conflict-generated "
           "mistrial work — that were not marked NR or otherwise reduced in the Court's "
           "Exhibits A and B. The specific entries, organized by category, are set forth in "
           "Exhibits [___] through [___] attached hereto."),
    ("10.", "Post-Judgment and Fees-on-Fees. My review of the post-judgment billing entries "
            "for Andrew White confirms that every entry in that category relates exclusively "
            "to the new trial motion, the motion to vacate, the fee application itself, "
            "successor-in-interest stipulations, or other collateral post-judgment matters, "
            "none of which advance the elder financial abuse claim. The details are set forth "
            "in Exhibit [___]."),
    ("11.", "Corrected Calculation. Based on my line-by-line audit, the only fee award "
            "consistent with the Court's stated methodology, applied uniformly, is $139,646.80. "
            "My annotated billing charts and supporting calculations are attached hereto as "
            "Exhibits [___] through [___]."),
]

for num, text in decl_paras:
    body(f"{num}  {text}")

doc.add_paragraph()
body(
    "I declare under penalty of perjury under the laws of the State of California that the "
    "foregoing is true and correct."
)
doc.add_paragraph()
body("Executed on May 15, 2026, at Los Angeles, California.")
doc.add_paragraph()
doc.add_paragraph()
body("_________________________________")
body("Taurin J. Robinson, Esq.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PROPOSED ORDER
# ══════════════════════════════════════════════════════════════════════════════

heading("[PROPOSED] ORDER GRANTING MOTION FOR RECONSIDERATION AND MODIFYING FEE AWARD", center=True)

body(
    "The Motion for Reconsideration of the Court's April 30, 2026 Order on Plaintiffs' Motion "
    "for Attorney's Fees and Costs filed by Defendant Michelle Pak came on regularly for "
    "hearing on [DATE] at [TIME] in Department [___] of the above-entitled Court, the "
    "Honorable [JUDGE] presiding. Counsel appeared as noted on the record. Having considered "
    "the moving papers, opposition, reply, oral argument, and all matters presented, and good "
    "cause appearing, the Court ORDERS as follows:"
)

body("1. The Motion for Reconsideration is GRANTED.", bold=False)
body(
    "2. The Court's April 30, 2026 Order on Plaintiffs' Motion for Attorney's Fees and Costs "
    "is hereby MODIFIED as follows:"
)
body("    a. The attorney's fees award is reduced from $532,936 to $139,646.80.", indent=0.25)
body("    b. The costs award of $9,973 remains unchanged.", indent=0.25)
body("    c. The total modified award is $149,619.80.", indent=0.25)
body(
    "3. All other provisions of the April 30, 2026 Order remain in full force and effect."
)
doc.add_paragraph()
body("IT IS SO ORDERED.")
doc.add_paragraph()
body("DATED: _______________")
doc.add_paragraph()
doc.add_paragraph()
body("_________________________________")
body("Judge of the Superior Court")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PROOF OF SERVICE
# ══════════════════════════════════════════════════════════════════════════════

heading("PROOF OF SERVICE", center=True)

body(
    "I am a resident of the State of California, over the age of eighteen years, and not a "
    "party to the within action. My business address is 3055 Wilshire Blvd., Suite 980, "
    "Los Angeles, California 90012. On May 15, 2026, I served the following document(s) on "
    "the parties listed below by the method(s) indicated:"
)

doc_list = [
    "NOTICE OF MOTION AND MOTION FOR RECONSIDERATION OF THE COURT'S APRIL 30, 2026 "
    "RULING ON PLAINTIFFS' MOTION FOR ATTORNEYS' FEES AND COSTS;",
    "MEMORANDUM OF POINTS AND AUTHORITIES IN SUPPORT THEREOF;",
    "DECLARATION OF TAURIN J. ROBINSON WITH EXHIBITS A THROUGH [___]; AND",
    "[PROPOSED] ORDER.",
]
for d in doc_list:
    bullet(d)

body("PARTIES SERVED:", bold=True)

parties = [
    ("Counsel for Plaintiffs",
     "Andrew M. White, Esq.\namwhiteesq@aol.com"),
    ("Counsel for Plaintiffs",
     "David Ryan Welch, Esq.\nlitigation@enso.law"),
    ("Counsel for Plaintiffs",
     "Steven R. Yee, Esq.\nsry@yeelaw.com"),
]
for role, info in parties:
    body(f"{role}: {info}", indent=0.25)

doc.add_paragraph()

service_methods = [
    ("☐", "BY U.S. MAIL: I enclosed said documents in a sealed envelope with postage thereon "
           "fully prepaid and placed for collection and mailing in accordance with ordinary "
           "business practices pursuant to Cal. Code Civ. Proc. § 1013(a)."),
    ("☐", "BY EXPRESS MAIL: I caused said documents to be deposited with an express service "
           "carrier providing overnight delivery pursuant to Cal. Code Civ. Proc. § 1013(c)."),
    ("☐", "BY PERSONAL SERVICE: I caused said documents to be hand-delivered to the office "
           "of the addressee pursuant to Cal. Code Civ. Proc. § 1011."),
    ("☒", "BY ELECTRONIC MAIL (E-MAIL): I caused said documents to be delivered "
           "electronically to the email addresses identified above pursuant to Cal. Code Civ. "
           "Proc. § 1010.6."),
]
for mark, text in service_methods:
    body(f"  {mark}  {text}", indent=0.25)

doc.add_paragraph()
body(
    "I declare that I am employed in the office of a member of the bar of this Court, under "
    "whose direction the service was made. I declare under penalty of perjury under the laws "
    "of the State of California that the foregoing is true and correct."
)
body("Executed on May 15, 2026, at Los Angeles, California.")
doc.add_paragraph()
doc.add_paragraph()
body("_________________________________")
body("Ivan Vilente Manriquez")

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Saved to: {OUTPUT_PATH}")
