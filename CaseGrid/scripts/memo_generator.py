"""Word legal memo generator for CaseGrid medical analysis."""
from __future__ import annotations

from pathlib import Path
from typing import Any

MEMO_OUTPUT_DIR = Path(__file__).resolve().parents[1] / "assets" / "memo_output"
MEMO_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Human-readable display names for technical rule fields
_FIELD_DISPLAY_NAMES: dict[str, str] = {
    "primary_cancer": "Ovarian cancer diagnosis",
    "age_at_diagnosis": "Age at diagnosis",
    "pathology_present": "Pathology documentation",
    "brca_result": "BRCA status",
    "metastatic": "Metastatic disease",
    "diagnosis_date": "Diagnosis date",
    "dob": "Date of birth",
}


def _lastname_first(name: str) -> str:
    """Return 'Lastname, Firstname' from any name string, else return name unchanged."""
    parts = name.strip().split()
    if len(parts) >= 2:
        return f"{parts[-1]}, {' '.join(parts[:-1])}"
    return name


def generate_memo(
    patient_name: str,
    dob: str | None,
    record_range: str,
    evaluation: dict[str, Any],
    evidence: list[dict[str, Any]],
    missing: list[str],
) -> Path:
    """Create and save a legal memorandum in .docx format."""
    from docx import Document

    doc = Document()
    display_name = _lastname_first(patient_name or "Unknown Patient")

    doc.core_properties.title = f"CaseGrid Medical Record Analysis Memorandum - {display_name}"

    doc.add_heading("THE ROBINSON LAW FIRM, PROF. CORP.", level=0)
    doc.add_heading(f"CaseGrid Medical Record Analysis Memorandum - {display_name}", level=1)

    # Executive Summary
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(f"This report analyzes medical records for patient {display_name}.")
    doc.add_paragraph(f"Patient: {display_name}")
    doc.add_paragraph(f"Date of Birth: {dob or 'Unknown'}")
    doc.add_paragraph(f"Record Range: {record_range}")
    doc.add_paragraph(f"Qualification Status: {evaluation.get('status', 'Unknown')}")

    # Qualification Summary — narrative, human-readable
    doc.add_heading("Qualification Summary", level=2)
    for item in evaluation.get("details", []):
        rule = item.get("rule", {})
        field = rule.get("field", "")
        display_field = _FIELD_DISPLAY_NAMES.get(field, field.replace("_", " ").title())
        extracted = item.get("value")
        result = item.get("result")
        operator = rule.get("operator", "")
        expected = rule.get("value", "")
        required = rule.get("required", True)

        if result is True:
            outcome = "Met"
        elif result is False:
            outcome = "Not Met"
        else:
            outcome = "Insufficient Data"

        # Build a readable sentence per criterion
        extracted_str = str(extracted) if extracted not in (None, "", [], "None") else "not found"
        if operator in ("exists", "is present"):
            detail = f"Present: {extracted_str}"
        elif operator in ("not exists", "is missing"):
            detail = f"Absent as expected: {extracted_str}"
        elif operator in ("contains",):
            detail = f"Contains '{expected}': {extracted_str}"
        elif operator in ("not contains", "does not contain"):
            detail = f"Does not contain '{expected}': {extracted_str}"
        elif operator == "<":
            detail = f"{extracted_str} (threshold: < {expected})"
        elif operator == ">":
            detail = f"{extracted_str} (threshold: > {expected})"
        elif operator in ("=", "!="):
            detail = f"{extracted_str} (expected: {expected})"
        else:
            detail = extracted_str

        req_label = "(required)" if required else "(optional)"
        doc.add_paragraph(
            f"{display_field} — {detail}   |   {outcome} {req_label}",
            style="List Bullet",
        )

    # Ovarian Cancer Evidence — grouped by page, no raw text
    doc.add_heading("Ovarian Cancer Evidence", level=2)

    evidence_by_page: dict[int, list[str]] = {}
    for item in evidence:
        page = item.get("page")
        keyword = item.get("normalized", "").strip()
        if page is not None and keyword:
            evidence_by_page.setdefault(page, []).append(keyword)

    sorted_pages = sorted(evidence_by_page.keys())

    if sorted_pages:
        page_list = ", ".join(map(str, sorted_pages))
        doc.add_paragraph(f"Keyword matches were found on the following pages:  {page_list}")

        doc.add_heading("Footnotes for Evidence Pages", level=3)
        for page in sorted_pages:
            keywords = ", ".join(dict.fromkeys(evidence_by_page[page]))  # deduplicate, preserve order
            doc.add_paragraph(
                f"Page {page}: matched keywords: {keywords}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No keyword evidence was extracted from the uploaded records.")

    # Legal conclusion + disclaimer
    doc.add_paragraph("")
    doc.add_heading("Conclusion", level=3)
    doc.add_paragraph(
        "Based on the above keyword matches, this record contains qualifying language "
        "indicating a positive ovarian cancer diagnosis."
    )
    doc.add_paragraph(
        "This automated analysis is not intended to replace human verification. "
        "Please review the referenced pages and confirm that the extracted keywords "
        "accurately reflect the medical findings."
    )

    # Missing Data
    doc.add_heading("Missing Data", level=2)
    if missing:
        for item in missing:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No missing required data identified.")

    # Recommendations
    doc.add_heading("Recommendations", level=2)
    if evaluation.get("status") == "Qualified":
        doc.add_paragraph("Proceed with litigation intake and downstream case development.")
    elif evaluation.get("status") == "Insufficient Data":
        doc.add_paragraph(
            "Request additional records and confirm pathology/diagnosis documentation."
        )
    else:
        doc.add_paragraph(
            "Review with attorney for discretionary qualification override if appropriate."
        )

    # Appendix
    doc.add_heading("Appendix", level=2)
    doc.add_paragraph(
        "Generated by CaseGrid Medical Analyzer. "
        "This memo is a draft support document and not legal advice."
    )

    # File name: Lastname_Firstname_talc_memo.docx
    safe_name = "_".join(_lastname_first(patient_name or "patient").replace(",", "").split())
    output_path = MEMO_OUTPUT_DIR / f"{safe_name}_talc_memo.docx"
    doc.save(output_path)
    return output_path
