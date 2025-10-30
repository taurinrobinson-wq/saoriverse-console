from docx import Document
from io import BytesIO

# --- Generate Word Document from Personal Log ---
def generate_doc(date, time, event, mood, reflections, insights):
    doc = Document()
    doc.add_heading("Personal Log Entry", level=1)

    doc.add_paragraph(f"Date: {date}")
    doc.add_paragraph(f"Time: {time}")
    doc.add_paragraph(f"Event: {event}")
    doc.add_paragraph(f"Mood: {mood}")
    doc.add_paragraph("Reflections:")
    doc.add_paragraph(reflections)
    doc.add_paragraph("Insights:")
    doc.add_paragraph(insights)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
# --- End Generate Word Document ---

import streamlit as st
from emotional_os.deploy.modules.auth import SaoynxAuthentication
from emotional_os.deploy.modules.doc_export import generate_doc
from emotional_os.deploy.modules.ui import render_splash_interface, render_main_app

# Page configuration
st.set_page_config(
    page_title="FirstPerson - Personal AI Companion",
    page_icon="graphics/FirstPerson-Logo.svg",
    layout="wide",
    initial_sidebar_state="expanded"
)

def main():
    auth = SaoynxAuthentication()
    if st.session_state.authenticated:
        render_main_app()
    else:
        render_splash_interface(auth)

if __name__ == "__main__":
    main()
